from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import *
from .serializers import *
from pathlib import Path
from django.conf import settings
from django.contrib.auth import authenticate, get_user_model
from rest_framework.authtoken.models import Token
from rest_framework.permissions import AllowAny, IsAuthenticated
from .permissions import IsAdmin, IsStaff
from rest_framework.authentication import TokenAuthentication
from django.shortcuts import get_object_or_404, render
from django.utils import timezone
from django.db import transaction
import logging
from django.db import IntegrityError
from django.contrib.auth.decorators import login_required
from django.db.models import Count, Q, F, Sum, Case, When, Value, IntegerField
from django.db.models.functions import ExtractMonth, ExtractYear
# Import specific models and serializers needed for attendance claims
from .models import AttendanceClaim, Attendance, AttendanceLog, AttendanceEntries, UserData, Department, Temp, Document, DocumentVersion, Log
from .serializers import AttendanceClaimSerializer, AttendanceSerializer, AttendanceLogSerializer, AttendanceEntriesSerializer, UserDataSerializer, DepartmentSerializer, TempSerializer, DocumentSerializer, DocumentVersionSerializer
from django.db.models import OuterRef, Subquery  # <-- Add this line
from django.contrib.auth.tokens import default_token_generator
from django.core.mail import send_mail
from django.db.models import Sum
from django.conf import settings
from django.contrib.auth.models import User
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes, force_str
from rest_framework.generics import GenericAPIView
from rest_framework import generics, permissions
from django.db.models import Window, F
from django.db.models.functions import RowNumber
from .models import PasswordResetOTP, AttendanceClaim
from .serializers import (
    ResetPasswordOTPRequestSerializer, 
    VerifyOTPSerializer,
    AttendanceClaimSerializer
)
from django.db.models import Count, Q
from django.db.models import Q 
from django.db.models.functions import TruncMonth
from rest_framework.decorators import api_view, permission_classes
from decimal import Decimal, InvalidOperation
from django.utils.timezone import now
from django.db.models.functions import ExtractMonth
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Q
import json
from .permissions import StaffAttendanceAccessPermission,StaffPayRollPermission,StaffAssertAccessPermission,StaffUserDataAccessPermission
import random
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from django.http import HttpResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from datetime import datetime
from .models import Task, AttendanceClaim
from .serializers import TaskCertificateSerializer, AttendanceCertificateSerializer
from .utils.email_utils import send_email_with_attachment
from docx import Document as DocxDocument
from docx2pdf import convert
import os
import tempfile
import pythoncom



logger = logging.getLogger(__name__)

def get_client_ip(request):
    """Helper function to get client IP address"""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip

def log_activity(user, action, entity_type, entity_id, description, request):
    """
    Helper function to log user activities.
    """
    try:
        # Handle cases where entity_id might be a UUID object
        if hasattr(entity_id, 'hex'):
            entity_id_str = str(entity_id)
        else:
            entity_id_str = str(entity_id) if entity_id is not None else ""
            
        # Truncate description if it's too long
        max_length = 1000  # Based on Log model's TextField
        if len(description) > max_length:
            description = description[:max_length-3] + '...'
            
        # Create log entry
        Log.objects.create(
            user_id=user,
            table_name=entity_type,
            action=action,
            old_data="",
            new_data=description,
            user_name=user.username
        )
        logger.info(f"Activity logged: User {user.username} {action} {entity_type} {entity_id_str} - {description}")
        return True
    except Exception as e:
        logger.error(f"Error logging activity: {str(e)}")
        return False


def generate_emp_id(role, department=None):
    """Generate employee ID with database locking to prevent duplicates"""
    with transaction.atomic():
        if role == 'intern' and department:
            # Use code_word if available, else first 3 letters of department
            prefix = (department.code_word or department.department[:3]).upper()
            
            # Find the last ID with this prefix (supports 001, 0001, etc.)
            last_id = Temp.objects.filter(
                emp_id__regex=rf'^{prefix}\d+$'  # Match any number of digits after prefix
            ).select_for_update().order_by('-emp_id').first()

            # Extract numeric portion and increment
            num = int(last_id.emp_id[len(prefix):]) + 1 if last_id else 1
            
            # Format with minimum 3 digits, allowing longer numbers
            return f"{prefix}{num:03d}"  # 001, 002,... 999, 1000, etc.

        else:  # Staff ID generation (unchanged)
            last_id = Temp.objects.filter(
                emp_id__startswith='STAFF'
            ).select_for_update().order_by('-emp_id').first()
            num = int(last_id.emp_id[5:]) + 1 if last_id else 1
            return f"STAFF{num:03d}"


class TempView(APIView):
    def get(self, request, emp_id=None):
        try:
            user = Temp.objects.get(user=request.user)
            if user.role == "intern":
                temp = Temp.objects.get(user=user.user, is_deleted=False)
                serializer = TempSerializer(temp)
                return Response(serializer.data, status=status.HTTP_200_OK)
                
            if emp_id:
                try:
                    temp = Temp.objects.get(emp_id=emp_id, is_deleted=False)
                    serializer = TempSerializer(temp)
                    return Response(serializer.data)
                except Temp.DoesNotExist:
                    return Response({"error": "Employee not found."}, status=status.HTTP_404_NOT_FOUND)
            else:
                temps = Temp.objects.filter(is_deleted=False)
                serializer = TempSerializer(temps, many=True)
                return Response(serializer.data)
        except Temp.DoesNotExist:
            return Response({"error": "No Temp data found for the logged-in user."}, status=status.HTTP_404_NOT_FOUND)

    @transaction.atomic 
    def post(self, request):
        # Verify admin user creating the record
        try:
            admin_temp = Temp.objects.get(user=request.user)
            if admin_temp.role.lower() not in ["admin", "hr", "staff"]:
                return Response("Only admins/HR can create employee records", 
                              status=status.HTTP_403_FORBIDDEN)
        except Temp.DoesNotExist:
            return Response({"error": "Admin user not configured"}, 
                          status=status.HTTP_403_FORBIDDEN)

        # Validate required fields
        username = request.data.get('user')
        role = request.data.get('role','intern').lower()
        department_name = request.data.get('department')

        if not username:
            return Response({"error": "Username is required"}, 
                          status=status.HTTP_400_BAD_REQUEST)
        if role == 'intern' and not department_name:
            return Response({"error": "Department required for interns"}, 
                          status=status.HTTP_400_BAD_REQUEST)

        # Get target user
        try:
            user_obj = User.objects.get(username=username)
        except User.DoesNotExist:
            return Response({"error": f"User {username} not found"}, 
                          status=status.HTTP_404_NOT_FOUND)

        # Check for existing Temp record
        if Temp.objects.filter(user=user_obj).exists():
            return Response({"error": "User already has an employee record"}, 
                          status=status.HTTP_400_BAD_REQUEST)

        # Generate emp_id with retry logic
        max_retries = 3
        for attempt in range(max_retries):
            try:
                with transaction.atomic():
                    # Generate emp_id
                    if role == 'intern':
                        department = Department.objects.get(department__iexact=department_name.strip())
                        emp_id = generate_emp_id(role, department)
                    else:
                        emp_id = generate_emp_id(role)

                    # Double-check uniqueness
                    if Temp.objects.filter(emp_id=emp_id).exists():
                        raise IntegrityError(f"emp_id {emp_id} already exists")

                    # Create Temp record
                    temp = Temp.objects.create(
                        user=user_obj,
                        emp_id=emp_id,
                        role=role
                    )

                    # Log activity
                    log_activity(request.user, "Created", "Temp", emp_id,
                               f"Created {role} record", request)

                    return Response(TempSerializer(temp).data, 
                                  status=status.HTTP_201_CREATED)

            except IntegrityError as e:
                if 'emp_id' in str(e) and attempt < max_retries - 1:
                    continue  # Retry with new emp_id
                return Response({"error": "Failed to generate unique emp_id"}, 
                              status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            except Department.DoesNotExist:
                return Response({"error": "Department not found"}, 
                              status=status.HTTP_404_NOT_FOUND)

        return Response({"error": "Failed after multiple attempts"}, 
                      status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    def put(self, request, emp_id):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" or user.role=='staff':
            return Response("Intern or staff has no access to this")
        
        try:
            temp = Temp.objects.get(emp_id=emp_id)
        except Temp.DoesNotExist:
            return Response({"error": "Employee not found."}, status=status.HTTP_404_NOT_FOUND)

        new_emp_id = request.data.get('emp_id')
        new_user_id = request.data.get('user')

        if new_emp_id and new_emp_id != emp_id:
            if Temp.objects.filter(emp_id=new_emp_id).exists():
                return Response({"error": f"emp_id {new_emp_id} already exists."},
                                status=status.HTTP_400_BAD_REQUEST)

        if new_user_id and new_user_id != temp.user.id:
            try:
                new_user = User.objects.get(id=new_user_id)
                if Temp.objects.filter(user=new_user).exists():
                    return Response({"error": f"User with id {new_user_id} already has a Temp object."},
                                    status=status.HTTP_400_BAD_REQUEST)
            except User.DoesNotExist:
                return Response({"error": f"User with id {new_user_id} not found."},
                                status=status.HTTP_404_NOT_FOUND)

        serializer = TempSerializer(temp, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()

            # Log activity
            log_activity(request.user, "Updated", "Temp", emp_id,
                         f"Updated Temp data", request)

            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, emp_id):
        try:
            user = Temp.objects.get(user=request.user)
            if user.role == "intern" or user.role=='staff':
                return Response("Intern or staff has no access to this")
        
            temp = Temp.objects.get(emp_id=emp_id, is_deleted=False)
            
            # Soft delete - set is_deleted to True instead of hard delete
            temp.is_deleted = True
            temp.save()

            # Log activity
            log_activity(request.user, "Deleted", "Temp", emp_id,
                         f"Soft deleted Temp record and related data", request)

            return Response({"message": f"Intern with emp_id {emp_id} deleted successfully"},
                            status=status.HTTP_200_OK)
        except Temp.DoesNotExist:
            return Response({"error": f"Intern with emp_id {emp_id} not found"},
                            status=status.HTTP_404_NOT_FOUND)
    
    def patch(self, request, emp_id):
        """Restore a soft-deleted intern (undo delete)"""
        try:
            user = Temp.objects.get(user=request.user)
            if user.role not in ["admin", "hr"]:
                return Response({"error": "Only admin/HR can restore deleted interns"}, 
                              status=status.HTTP_403_FORBIDDEN)
        
            temp = Temp.objects.get(emp_id=emp_id, is_deleted=True)
            
            # Restore the intern - set is_deleted to False
            temp.is_deleted = False
            temp.save()

            # Log activity
            log_activity(request.user, "Restored", "Temp", emp_id,
                         f"Restored Temp record and related data", request)

            return Response({"message": f"Intern with emp_id {emp_id} restored successfully"},
                            status=status.HTTP_200_OK)
        except Temp.DoesNotExist:
            return Response({"error": f"Deleted intern with emp_id {emp_id} not found"},
                            status=status.HTTP_404_NOT_FOUND)














class UserDataView(APIView):
    permission_classes = [IsAuthenticated]
    def check_user_role(self, username, role_type):
        if not User.objects.filter(username=username).exists():
            return False, Response({"error": f"Reporting {role_type} not found"}, status=status.HTTP_400_BAD_REQUEST)
        
        user = User.objects.get(username=username)
        try:
            temp = Temp.objects.get(user=user)
            if temp.role.lower() == 'intern':
                return False, Response({"error": f"Interns cannot be assigned as reporting {role_type}s"}, status=status.HTTP_400_BAD_REQUEST)
            return True, None
        except Temp.DoesNotExist:
            return False, Response({"error": f"Reporting {role_type} role not found"}, status=status.HTTP_400_BAD_REQUEST)

    def get(self, request, emp_id=None):
        user = Temp.objects.get(user=request.user)
        
        if user.role == "intern":
            try:
                user_data = UserData.objects.get(emp_id__emp_id=user.emp_id,is_deleted=False)
                serializer = UserDataSerializer(user_data)
                return Response(serializer.data)
            except UserData.DoesNotExist:
                return Response({"error": "UserData not found"}, status=status.HTTP_404_NOT_FOUND)

        if emp_id:
            try:
                # Get staff department if staff role
                if user.role.lower() == 'staff':
                    staff_data = UserData.objects.get(user=request.user,is_deleted=False)
                    intern_data = UserData.objects.get(emp_id__emp_id=emp_id,is_deleted=False)
                    
                    if intern_data.department != staff_data.department:
                        return Response({"error": "Intern not in your department or staff not found"}, 
                                      status=status.HTTP_403_FORBIDDEN)

                user_data = UserData.objects.get(emp_id__emp_id=emp_id,is_deleted=False)
                serializer = UserDataSerializer(user_data)
                return Response(serializer.data)
            except UserData.DoesNotExist:
                return Response({"error": "UserData not found"}, status=status.HTTP_404_NOT_FOUND)

        else:
            if user.role.lower() == 'admin':
                user_data = UserData.objects.filter(emp_id__is_deleted=False)
            elif user.role.lower() == 'staff':
                # Show all interns in same department
                staff_data = UserData.objects.get(user=request.user)
                user_data = UserData.objects.filter(
                    department=staff_data.department,
                    emp_id__role='intern',
                    emp_id__is_deleted=False
                )
            else:
                user_data = UserData.objects.none()

            serializer = UserDataSerializer(user_data, many=True)
            return Response(serializer.data)

    def post(self, request):
        # First, get the logged-in user and their role
        try:
            logged_in_temp = Temp.objects.get(user=request.user)
        except Temp.DoesNotExist:
            return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)
        
        emp_id = request.data.get("emp_id")
        reporting_supervisor = request.data.get("reporting_supervisor")
        reporting_manager = request.data.get("reporting_manager")
        
        # Check if user is an intern and enforce restrictions
        if logged_in_temp.role.lower() == "intern":
            # Interns can only create data for themselves
            if str(emp_id) != str(logged_in_temp.emp_id):
                return Response(
                    {"error": "As an intern, you can only create UserData for yourself"},
                    status=status.HTTP_403_FORBIDDEN
                )
            
            # Interns cannot assign reporting managers or supervisors
            if reporting_manager or reporting_supervisor:
                return Response(
                    {"error": "As an intern, you cannot assign reporting managers or supervisors"},
                    status=status.HTTP_403_FORBIDDEN
                )
        
        # For other roles, proceed with the current checks
        # Check reporting manager
        if reporting_manager:
            success, response = self.check_user_role(reporting_manager, "manager")
            if not success:
                return response
        
        # Check reporting supervisor
        if reporting_supervisor:
            success, response = self.check_user_role(reporting_supervisor, "supervisor")
            if not success:
                return response
        
        if not Temp.objects.filter(emp_id=emp_id).exists():
            return Response({"error": "No user found"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Check if UserData already exists but is incomplete
        existing_user_data = UserData.objects.filter(emp_id__emp_id=emp_id).first()
        if existing_user_data:
            # Check if required fields are null
            required_fields = ['domain', 'start_date', 'shift_timing', 'reporting_manager']
            incomplete = any(getattr(existing_user_data, field) is None for field in required_fields if hasattr(existing_user_data, field))
            
            if incomplete:
                # If required fields are null, update the existing record
                serializer = UserDataSerializer(existing_user_data, data=request.data, partial=True)
                if serializer.is_valid():
                    # Set status based on start date for existing records too
                    start_date_str = request.data.get('start_date')
                    if start_date_str:
                        from datetime import datetime
                        try:
                            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
                            today = datetime.now().date()
                            if start_date > today:
                                serializer.validated_data['user_status'] = 'yettojoin'
                            else:
                                serializer.validated_data['user_status'] = 'inprogress'
                        except (ValueError, TypeError):
                            serializer.validated_data['user_status'] = 'inprogress'
                    
                    serializer.save()
                    return Response(serializer.data, status=status.HTTP_200_OK)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            else:
                # If required fields are not null, return error
                return Response({"error":"Complete UserData for this emp_id already exists."},
                            status=status.HTTP_400_BAD_REQUEST)
        
        # If no UserData exists, create a new one
        temp = get_object_or_404(Temp, emp_id=emp_id)
        user = temp.user
        
        # Initialize serializer with request data
        serializer = UserDataSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Set the status after validation but before saving
        start_date_str = request.data.get('start_date')
        if start_date_str:
            from datetime import datetime
            try:
                start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
                today = datetime.now().date()
                serializer.validated_data['user_status'] = 'yettojoin' if start_date > today else 'inprogress'
            except (ValueError, TypeError):
                serializer.validated_data['user_status'] = 'inprogress'
        else:
            serializer.validated_data['user_status'] = 'inprogress'
        
        # Set required fields
        serializer.validated_data['emp_id'] = temp
        serializer.validated_data['user'] = user
        
        # Save the instance
        user_data = serializer.save()
        return Response(UserDataSerializer(user_data).data, status=status.HTTP_201_CREATED)

    def put(self, request, emp_id):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" :
            return Response("Intern has no access to this")
        
        try:
            user_data = UserData.objects.get(emp_id__emp_id=emp_id)
        except UserData.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        
        # Check reporting manager
        if request.data.get("reporting_manager"):
            success, response = self.check_user_role(request.data.get("reporting_manager"), "manager")
            if not success:
                return response
        
        # Check reporting supervisor
        if request.data.get("reporting_supervisor"):
            success, response = self.check_user_role(request.data.get("reporting_supervisor"), "supervisor")
            if not success:
                return response
        
        serializer = UserDataSerializer(user_data, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def patch(self, request, emp_id):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern":
            return Response("Intern has no access to this")
        
        try:
            user_data = UserData.objects.get(emp_id__emp_id=emp_id)
        except UserData.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        
        # Handle soft delete operation - check if this is a delete request
        if 'user_status' in request.data and str(request.data.get('user_status')).lower() == 'deleted':
            # Soft delete the UserData and related Temp record
            user_data.user_status = 'deleted'  # Set proper status
            user_data.save()
            
            # Also update the Temp record status if needed
            temp_record = user_data.emp_id
            # Don't set is_deleted=True on Temp, just update user_status in UserData
            
            # Log activity
            log_activity(request.user, "Deleted", "UserData", user_data.id,
                       f"Soft deleted intern {emp_id}", request)
            
            return Response({"message": f"Intern {emp_id} deleted successfully"}, 
                          status=status.HTTP_200_OK)
        
        # Check reporting manager
        if request.data.get("reporting_manager"):
            success, response = self.check_user_role(request.data.get("reporting_manager"), "manager")
            if not success:
                return response
        
        # Check reporting supervisor
        if request.data.get("reporting_supervisor"):
            success, response = self.check_user_role(request.data.get("reporting_supervisor"), "supervisor")
            if not success:
                return response
        
        serializer = UserDataSerializer(user_data, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    

    def delete(self, request, emp_id):
        
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" :
            return Response("Intern has no access to this")
        try:
            user_data = UserData.objects.get(emp_id__emp_id=emp_id)
        except UserData.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)

        # Log activity before deleting
        log_activity(request.user, "Deleted", "UserData", user_data.id,
                     f"Deleted UserData for emp_id {emp_id}", request)

        user_data.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

class AllUserDataView(APIView):
    """
    Simple API endpoint that returns all user data from UserData model
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            # Get all users from UserData model
            user = request.user
            user_temp = Temp.objects.get(user=user,is_deleted=False)
            
            users = UserData.objects.filter(is_deleted=False).select_related('user')
            
            # Serialize the data
            user_data = []
            for user in users:
                user_data.append({
                    'id': user.user.id,
                    'username': user.user.username,
                    'email': user.user.email,
                    'is_active': user.user.is_active,
                    'role': user_temp.role,
                    'access_rights': {
                        'is_attendance_access': getattr(user, 'is_attendance_access', False),
                        'is_payroll_access': getattr(user, 'is_payroll_access', False),
                        'is_internmanagement_access': getattr(user, 'is_internmanagement_access', False),
                        'is_assert_access': getattr(user, 'is_assert_access', False)
                    }
                })
            
            return Response({
                'status': 'success',
                'count': len(user_data),
                'users': user_data
            }, status=status.HTTP_200_OK)
            
        except Exception as e:
            return Response({
                'status': 'error',
                'message': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class UserUpdateView(APIView):
    def patch(self, request, emp_id):
        try:
            # Get the Temp instance
            temp = Temp.objects.get(emp_id=emp_id)
            user = temp.user  # Get the related User instance
            
            # Update User model fields
            if 'first_name' in request.data:
                user.first_name = request.data['first_name']
            if 'last_name' in request.data:
                user.last_name = request.data['last_name']
            if 'email' in request.data:
                user.email = request.data['email']

            user.save()
            
            return Response({
                'status': 'success',
                'message': 'User updated successfully',
                'data': {
                    'first_name': user.first_name,
                    'last_name': user.last_name,
                    'email': user.email,
                    'username': user.username
                }
            }, status=status.HTTP_200_OK)
            
        except Temp.DoesNotExist:
            return Response({
                'status': 'error',
                'message': 'User not found'
            }, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({
                'status': 'error',
                'message': str(e)
            }, status=status.HTTP_400_BAD_REQUEST)


class PersonalDataView(APIView):
    permission_classes = [IsAuthenticated]
    def get(self, request, emp_id=None):
        user = Temp.objects.get(user=request.user, is_deleted=False)
        
        # If user is an intern, they can only see their own data
        if user.role.lower() == "intern":
            try:
                personal_data = PersonalData.objects.get(emp_id__emp_id=user.emp_id, is_deleted=False)
                serializer = PersonalDataSerializer(personal_data)
                return Response(serializer.data)
            except PersonalData.DoesNotExist:
                return Response({"error": "PersonalData not found"}, status=status.HTTP_404_NOT_FOUND)

        if emp_id:
            try:
                target_user = Temp.objects.get(emp_id=emp_id, is_deleted=False)
                
                # If target is staff
                if target_user.role.lower() == 'staff':
                    # Only allow access if requesting user is admin or the same staff member
                    if user.role.lower() != 'admin' and user.emp_id != emp_id:
                        return Response(
                            {"error": "You can only access your own staff data"},
                            status=status.HTTP_403_FORBIDDEN
                        )
                # If target is intern
                else:
                    if user.role.lower() == 'staff':
                        staff_data = UserData.objects.get(user=request.user, is_deleted=False)
                        intern_userdata = UserData.objects.get(emp_id__emp_id=emp_id, is_deleted=False)
                        
                        if intern_userdata.department != staff_data.department:
                            return Response(
                                {"error": "Intern not in your department"}, 
                                status=status.HTTP_403_FORBIDDEN
                            )

                personal_data = PersonalData.objects.get(emp_id__emp_id=emp_id, is_deleted=False)
                serializer = PersonalDataSerializer(personal_data)
                return Response(serializer.data)
                
            except (Temp.DoesNotExist, PersonalData.DoesNotExist):
                return Response({"error": "PersonalData not found"},status=status.HTTP_404_NOT_FOUND)

        else:
            # Rest of the method remains the same
            if user.role.lower() == 'admin':
                personal_data = PersonalData.objects.filter(is_deleted=False)
            elif user.role.lower() == 'staff':
                staff_data = UserData.objects.get(user=request.user, is_deleted=False)
                interns_in_dept = UserData.objects.filter(
                    department=staff_data.department, is_deleted=False
                ).values_list('emp_id', flat=True)
                personal_data = PersonalData.objects.filter(emp_id__in=interns_in_dept, is_deleted=False)
            else:
                personal_data = PersonalData.objects.none()

            serializer = PersonalDataSerializer(personal_data, many=True)
            return Response(serializer.data)


    def post(self, request):
    # First, get the logged-in user and their role
        try:
            logged_in_temp = Temp.objects.get(user=request.user)
        except Temp.DoesNotExist:
            return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)
        
        emp_id = request.data.get('emp_id')
        if not emp_id:
            return Response({"error": "emp_id is required"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Check if user is an intern and enforce restrictions
        if logged_in_temp.role.lower() == "intern":
            # Interns can only create data for themselves
            if str(emp_id) != str(logged_in_temp.emp_id):
                return Response(
                    {"error": "As an intern, you can only create PersonalData for yourself"},
                    status=status.HTTP_403_FORBIDDEN
                )
        
        if PersonalData.objects.filter(emp_id=emp_id,is_deleted=False).exists():
            return Response({"error": "PersonalData already exists for this emp_id"},
                        status=status.HTTP_400_BAD_REQUEST)
        
        temp = get_object_or_404(Temp, emp_id=emp_id)
        user = temp.user
        
        serializer = PersonalDataSerializer(data=request.data)
        if serializer.is_valid():
            serializer.validated_data['emp_id'] = temp
            serializer.validated_data['user'] = user
            serializer.save()
            
            # Log activity
            log_activity(request.user, "Created", "PersonalData", serializer.data['id'],
                        f"Created PersonalData for emp_id {emp_id}", request)
            
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


    def put(self, request, emp_id):  # Changed parameter to emp_id
        
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" :
            return Response("Intern has no access to this")
        try:
            personal_data = PersonalData.objects.get(emp_id__emp_id=emp_id)  # Fetch by emp_id
        except PersonalData.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)

        serializer = PersonalDataSerializer(personal_data, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()

            # Log activity
            log_activity(request.user, "Updated", "PersonalData", personal_data.id,
                         f"Updated PersonalData for emp_id {emp_id}", request)

            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, emp_id):  # Changed parameter to emp_id
        
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" :
            return Response("Intern has no access to this")
        try:
            personal_data = PersonalData.objects.get(emp_id__emp_id=emp_id)  # Fetch by emp_id

            # Log activity before deleting
            log_activity(request.user, "Deleted", "PersonalData", personal_data.id,
                         f"Deleted PersonalData for emp_id {emp_id}", request)

            personal_data.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except PersonalData.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)






class CollegeDetailsView(APIView):
    permission_classes = [IsAuthenticated]   
    def get(self, request, emp_id=None):
        try:
            # Get the logged-in user's Temp record
            user_temp = Temp.objects.get(user=request.user,is_deleted=False)
            role = user_temp.role.lower()

            # Initialize response data
            response_data = []

            # Determine employees to fetch college details for based on role
            if emp_id:
                # Single employee request
                try:
                    target_temp = Temp.objects.get(emp_id=emp_id,is_deleted=False)

                    # Permission checks
                    if role == "intern" and user_temp.emp_id != target_temp.emp_id:
                        return Response(
                            {"error": "You can only view your own college details."},
                            status=status.HTTP_403_FORBIDDEN
                        )

                    if role == "staff":
                        # Check department alignment
                        staff_data = UserData.objects.get(user=request.user,is_deleted=False)
                        intern_data = UserData.objects.get(emp_id=target_temp,is_deleted=False)
                        if intern_data.department != staff_data.department:
                            return Response(
                                {"error": "This intern is not in your department."},
                                status=status.HTTP_403_FORBIDDEN
                            )

                    employees = [target_temp]

                except Temp.DoesNotExist:
                    return Response(
                        {"error": "Employee not found."}, 
                        status=status.HTTP_404_NOT_FOUND
                    )
            else:
                # Multiple employees based on role
                if role == "intern":
                    employees = [user_temp]
                elif role == "staff":
                    try:
                        staff_data = UserData.objects.get(user=request.user,is_deleted=False)
                        department = staff_data.department
                        interns_in_dept = UserData.objects.filter(
                            department=department,
                            emp_id__role='intern',is_deleted=False
                        ).values_list('emp_id', flat=True)
                        employees = Temp.objects.filter(emp_id__in=interns_in_dept,is_deleted=False)
                    except UserData.DoesNotExist:
                        return Response(
                            {"error": "Staff department not configured."},
                            status=status.HTTP_400_BAD_REQUEST
                        )
                elif role in ["admin", "hr"]:
                    employees = Temp.objects.filter(is_deleted=False)
                else:
                    return Response(
                        {"error": "Invalid user role."},
                        status=status.HTTP_400_BAD_REQUEST
                    )

            # Process each employee's college details
            for employee in employees:
                try:
                    college_details = CollegeDetails.objects.get(emp_id=employee,is_deleted=False)
                    serializer = CollegeDetailsSerializer(college_details)
                    
                    employee_data = {
                        "emp_id": employee.emp_id,
                        "name": employee.user.username,
                        "college_details": serializer.data
                    }
                    response_data.append(employee_data)
                
                except CollegeDetails.DoesNotExist:
                    continue  # Skip employees without college details

            # Return appropriate response format
            if emp_id:
                return Response(response_data[0] if response_data else {}, 
                              status=status.HTTP_200_OK)
            
            return Response(response_data, status=status.HTTP_200_OK)

        except Temp.DoesNotExist:
            return Response(
                {"error": "User profile not found."}, 
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            return Response(
                {"error": str(e)}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def post(self, request):
    # First, get the logged-in user and their role
        try:
            logged_in_temp = Temp.objects.get(user=request.user)
        except Temp.DoesNotExist:
            return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)
        
        emp_id = request.data.get('emp_id')
        if not emp_id:
            return Response({"error": "emp_id is required"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Check if user is an intern and enforce restrictions
        if logged_in_temp.role.lower() == "intern":
            # Interns can only create data for themselves
            if str(emp_id) != str(logged_in_temp.emp_id):
                return Response(
                    {"error": "As an intern, you can only create CollegeDetails for yourself"},
                    status=status.HTTP_403_FORBIDDEN
                )
        
        if CollegeDetails.objects.filter(emp_id__emp_id=emp_id , is_deleted=False).exists():
            return Response({"error": "The College Details of this user already exists"}, 
                        status=status.HTTP_400_BAD_REQUEST)
        
        temp = get_object_or_404(Temp, emp_id=emp_id)
        
        serializer = CollegeDetailsSerializer(data=request.data)
        if serializer.is_valid():
            serializer.validated_data['emp_id'] = temp
            serializer.save()
            
            # Log activity
            log_activity(request.user, "Created", "CollegeDetails", serializer.data['id'],
                        f"Created CollegeDetails for emp_id {emp_id}", request)
            
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def put(self, request, emp_id):  # Changed parameter to emp_id
        
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" :
            return Response("Intern has no access to this")
        try:
            college_details = CollegeDetails.objects.get(emp_id__emp_id=emp_id)  # Fetch by emp_id
        except CollegeDetails.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)

        serializer = CollegeDetailsSerializer(college_details, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()

            # Log activity
            log_activity(request.user, "Updated", "CollegeDetails", college_details.id,
                         f"Updated CollegeDetails for emp_id {emp_id}", request)

            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, emp_id):  # Changed parameter to emp_id
        
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" :
            return Response("Intern has no access to this")
        try:
            college_details = CollegeDetails.objects.get(emp_id__emp_id=emp_id)  # Fetch by emp_id

            # Log activity before deleting
            log_activity(request.user, "Deleted", "CollegeDetails", college_details.id,
                         f"Deleted CollegeDetails for emp_id {emp_id}", request)

            college_details.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except CollegeDetails.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)

# class AssertStockView(viewsets.ModelViewSet):
#     queryset = AssertStock.objects.all()
#     serializer_class = AssertStockSerializer
#     lookup_field = 'assert_id'

class AssertStockView(APIView):
    permission_classes = [IsAuthenticated, StaffAssertAccessPermission]

    def get(self, request, pk=None):
        user = request.user
        try:
            user_temp = Temp.objects.get(user=user)
            
            if pk:
                assert_stock = AssertStock.objects.get(assert_id=pk, is_deleted=False)
                if user_temp.role == 'staff':
                    staff_dept = UserData.objects.get(user=user).department
                    if assert_stock.department != staff_dept:
                        return Response({"error": "Unauthorized access to asset"}, status=status.HTTP_403_FORBIDDEN)
                serializer = AssertStockSerializer(assert_stock)
                return Response(serializer.data)
            
            if user_temp.role == 'staff':
                staff_dept = UserData.objects.get(user=user).department
                assert_stock = AssertStock.objects.filter(department=staff_dept, is_deleted=False)
            elif user_temp.role in ['admin', 'hr']:
                assert_stock = AssertStock.objects.filter(is_deleted=False)
            else:
                return Response({"error": "Unauthorized access"}, status=status.HTTP_403_FORBIDDEN)
            
            serializer = AssertStockSerializer(assert_stock, many=True)
            return Response(serializer.data)

        except Temp.DoesNotExist:
            return Response({"error": "User profile not found"}, status=status.HTTP_404_NOT_FOUND)
        except AssertStock.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)

    @transaction.atomic
    def post(self, request):
        user = request.user
        try:
            user_temp = Temp.objects.get(user=user)
            request_data = request.data.copy()

            serializer = AssertStockSerializer(data=request_data)
            
            if serializer.is_valid():
                assert_stock = serializer.save()
                
                # Handle intern assignment/unassignment
                emp_id = request_data.get('emp_id')
                if emp_id:  # Assign to intern
                    try:
                        temp = Temp.objects.get(emp_id=emp_id)
                        assert_stock.emp_id = temp
                        assert_stock.user = temp.user
                        assert_stock.inhand = False
                        assert_stock.save()
                        
                        # Update UserData with asset assignment
                        UserData.objects.filter(emp_id=temp).update(asset_code=assert_stock)
                        
                    except Temp.DoesNotExist:
                        pass  # Continue without assignment if intern not found
                else:  # Unassign from intern (emp_id is empty/None)
                    # Clear previous assignment in UserData
                    if assert_stock.emp_id:
                        UserData.objects.filter(asset_code=assert_stock).update(asset_code=None)
                    
                    # Clear asset assignment
                    assert_stock.emp_id = None
                    assert_stock.user = None
                    assert_stock.inhand = True
                    assert_stock.save()

                # Return updated data
                updated_serializer = AssertStockSerializer(assert_stock)
                return Response(updated_serializer.data, status=status.HTTP_201_CREATED)
            
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        except Temp.DoesNotExist:
            return Response({"error": "User profile not found"}, status=status.HTTP_404_NOT_FOUND)

    @transaction.atomic
    def put(self, request, pk):
        user = request.user
        try:
            user_temp = Temp.objects.get(user=user)
            assert_stock = AssertStock.objects.get(assert_id=pk)
            request_data = request.data.copy()

            # Staff department check
            if user_temp.role == 'staff':
                staff_dept = UserData.objects.get(user=user).department
                if assert_stock.department != staff_dept:
                    return Response(
                        {"error": "Cannot modify assets in other departments"},
                        status=status.HTTP_403_FORBIDDEN
                    )
                if 'department' in request_data:
                    del request_data['department']

            # First update the allocated_type directly if it's in the request
            if 'allocated_type' in request_data:
                assert_stock.allocated_type = request_data['allocated_type']
                assert_stock.save(update_fields=['allocated_type'])
                
            serializer = AssertStockSerializer(assert_stock, data=request_data, partial=True)
            
            if serializer.is_valid():
                emp_id = request_data.get('emp_id')
                if emp_id in ['null', '']:
                    emp_id = None
                
                # Handle assignment/unassignment
                if 'emp_id' in request_data:
                    error = self.handle_asset_assignment(assert_stock, emp_id, user_temp)
                    if error:
                        return error

                updated_asset = serializer.save()

                # Log activity
                log_activity(user, "Updated", "AssertStock", pk, 
                           f"Updated asset details", request)
                return Response(serializer.data)
            
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        except Temp.DoesNotExist:
            return Response({"error": "User profile not found"}, status=status.HTTP_404_NOT_FOUND)
        except AssertStock.DoesNotExist:
            return Response({"error": "Asset not found"}, status=status.HTTP_404_NOT_FOUND)

    def handle_asset_assignment(self, asset, emp_id, user_temp):
        try:
            if emp_id:  # Assignment
                temp = Temp.objects.get(emp_id=emp_id)
                
                # Validate intern department for staff
                if user_temp.role == 'staff':
                    intern_data = UserData.objects.get(emp_id=temp)
                    staff_dept = UserData.objects.get(user=user_temp.user).department
                    if intern_data.department != staff_dept:
                        return Response(
                            {"error": "Cannot assign asset to intern from different department"},
                            status=status.HTTP_400_BAD_REQUEST
                        )

                # Update assignment history
                if emp_id:  # Assignment
                    AssertAssignmentLog.objects.create(
                        asset=asset,
                        event_type='ASSIGNED',
                        user=user_temp.user,
                        emp=temp,
                        notes=f"Assigned to {temp.emp_id}"
                    )
                asset.update_assignment_history(emp_id)
                
                # Handle previous assignments
                old_assign = AssertStock.objects.filter(emp_id=temp).first()
                if old_assign:
                    old_assign.emp_id = None
                    old_assign.inhand = True
                    old_assign.save()
                    UserData.objects.filter(asset_code=old_assign).update(asset_code=None)

                # Assign to new user
                asset.emp_id = temp
                asset.user = temp.user
                asset.inhand = False
                asset.save()

                # Update UserData
                UserData.objects.update_or_create(
                    emp_id=temp,
                    defaults={'asset_code': asset}
                )

            else:  # Unassignment
                # Clear assignment history
                asset.update_assignment_history(None)
                AssertAssignmentLog.objects.create(
                    asset=asset,
                    event_type='RETURNED',
                    user=user_temp.user,
                    emp=asset.emp_id,
                    notes="Asset returned to inventory"
                )
                # Remove from previous user's UserData
                UserData.objects.filter(asset_code=asset).update(asset_code=None)
                
                # Clear asset fields
                asset.emp_id = None
                asset.user = None
                asset.inhand = True
                asset.save()

            return Response(AssertStockSerializer(asset).data)

        except Temp.DoesNotExist:
            return Response({"error": "Employee not found"}, status=status.HTTP_404_NOT_FOUND)
        except UserData.DoesNotExist:
            return Response({"error": "Intern data not found"}, status=status.HTTP_404_NOT_FOUND)
        

    def delete(self, request, pk):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" or user.role=='staff':
            return Response("Intern or staff has no access to this")
        
        try:
            assert_stock = AssertStock.objects.get(assert_id=pk)

            # Log activity before deleting
            log_activity(request.user, "Deleted", "AssertStock", pk,
                         f"Deleted AssertStock record", request)

            if assert_stock.emp_id:
                UserData.objects.filter(emp_id=assert_stock.emp_id, asset_code=assert_stock).update(asset_code=None)
            assignment_history = assert_stock.assignment_history
            assert_stock.delete()

            return Response({
                "message": f"Assert stock with ID {pk} deleted successfully.",
                "assignment_history": assignment_history
            }, status=status.HTTP_204_NO_CONTENT)
        except AssertStock.DoesNotExist:
            return Response({"error": f"Assert stock with ID {pk} not found"},
                            status=status.HTTP_404_NOT_FOUND)

class UserPermissionsView(APIView):
    """
    API endpoint that returns the current user's permissions.
    """
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        from .permissions import get_user_permissions
        permissions = get_user_permissions(request.user)
        return Response(permissions)


class AssertIssueView(APIView):
    permission_classes = [IsAuthenticated, StaffAssertAccessPermission]
    def get(self, request, pk=None):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern":
            return Response("intern has no access to this")
        if pk:
            try:
                assert_issue = AssertIssue.objects.get(pk=pk,is_deleted=False)
                serializer = AssertIssueSerializer(assert_issue)
                return Response(serializer.data)
            except AssertIssue.DoesNotExist:
                return Response(status=status.HTTP_404_NOT_FOUND)
        else:
            assert_issues = AssertIssue.objects.filter(is_deleted=False)
            serializer = AssertIssueSerializer(assert_issues, many=True)
            return Response(serializer.data)

    @transaction.atomic
    def post(self, request):
        user = Temp.objects.get(user=request.user)
        

        serializer = AssertIssueSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        data = serializer.validated_data
        assert_id = data['assert_id']
        condition = data['condition']
        alternate_laptop = request.data.get('alternate_laptop')
        # if AssertIssue.objects.filter(assert_id=assert_id , is_deleted=False).exists():
        #     return Response("Already there is a entry for this assert",status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Validate if alternate laptop is required
            if condition == 'Not Usable':
                # Check if asset is assigned to a user
                if assert_id.emp_id:
                    if not alternate_laptop:
                        return Response(
                            {"error": "Alternate laptop is required when marking an assigned asset as Not Usable"},
                            status=status.HTTP_400_BAD_REQUEST
                        )

            # Main transaction block
            with transaction.atomic():
                # Handle asset swap and user data updates
                if condition == 'Not Usable' and alternate_laptop:
                    # Validate alternate laptop
                    try:
                        alt_assert = AssertStock.objects.get(assert_id=alternate_laptop,is_deleted=False)
                        if alt_assert.emp_id:
                            return Response(
                                {"error": "Alternate laptop is already assigned"},
                                status=status.HTTP_400_BAD_REQUEST
                            )
                        
                        if AssertIssue.objects.filter(assert_id=alt_assert, condition='Not Usable',is_deleted=False).exists():
                            return Response(
                                {"error": "Alternate laptop is already under repair"},
                                status=status.HTTP_400_BAD_REQUEST
                            )

                        # Get original asset details
                        original_assert = assert_id
                        original_user = original_assert.emp_id

                        # Swap assets in UserData
                        if original_user:
                            # Update original user's data to new asset
                            UserData.objects.filter(emp_id=original_user).update(asset_code=alt_assert)
                            
                            # Clear original asset assignment
                            UserData.objects.filter(asset_code=original_assert).update(asset_code=None)

                        # Update assignment history for both assets
                        original_assert.update_assignment_history(None)
                        alt_assert.update_assignment_history(original_user.emp_id)
                        print(original_user.emp_id)

                        AssertAssignmentLog.objects.create(
                            asset=original_assert,
                            event_type='REPAIR_RETURNED',
                            user=original_user.user,
                            emp=original_assert.emp_id,
                            notes="Asset returned to inventory and assigned to Null "
                        )



                        # Update asset records
                        alt_assert.emp_id = original_user
                        alt_assert.user = original_user.user if original_user else None
                        alt_assert.inhand = False
                        alt_assert.save()

                        AssertAssignmentLog.objects.create(
                            asset=alt_assert,
                            event_type='ASSIGNED',
                            user=alt_assert.user,
                            emp=alt_assert.emp_id,
                            notes="Asset is assigned "
                        )

                        original_assert.emp_id = None
                        original_assert.user = None
                        original_assert.inhand = True
                        original_assert.save()

                    except AssertStock.DoesNotExist:
                        return Response(
                            {"error": "Alternate laptop not found"},
                            status=status.HTTP_404_NOT_FOUND
                        )

                # Create the AssertIssue record
                assert_issue = serializer.save()
                # AssertAssignmentLog.objects.create(
                #     asset=assert_issue.assert_id,
                #     event_type='REPAIR',
                #     user=request.user,
                #     notes=f"Repair initiated: {assert_issue.issue}"
                # )
                # Update assignment history with Temp object
                print(assert_id.emp_id)
                if assert_id.emp_id:
                    assert_id.update_assignment_history(assert_id.emp_id)
                    

                
                # Log the activity with proper temp reference
                log_activity(
                    request.user,
                    "Created",
                    "AssertIssue",
                    assert_issue.id,
                    f"Created issue for {assert_id.assert_id} affecting user {assert_id.emp_id.emp_id if assert_id.emp_id else 'none'}",
                    request
                )

                return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Exception as e:
            logger.error(f"Error creating AssertIssue: {str(e)}")
            return Response(
                {"error": "Internal server error"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @transaction.atomic
    def put(self, request, pk):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern":
            return Response("Intern has no access to this", status=status.HTTP_403_FORBIDDEN)

        try:
            # Get AssertIssue using string assert_id from URL
            assert_issue = AssertIssue.objects.get(pk=pk,is_deleted=False)
        except AssertIssue.DoesNotExist:
            return Response({"error": "AssertIssue not found"}, status=status.HTTP_404_NOT_FOUND)

        serializer = AssertIssueSerializer(assert_issue, data=request.data, partial=True)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        data = serializer.validated_data
        condition = data.get('condition', assert_issue.condition)
        
        alternate_laptop = request.data.get('alternate_laptop', assert_issue.alternate_laptop)
        original_assert = assert_issue.assert_id

        if condition == 'Usable':
              
            AssertAssignmentLog.objects.create(
                                asset=original_assert,
                                event_type='ALRIGHT',
                                user=None,
                                emp=original_assert.emp_id,
                                notes="The asset repair is completed and now alright"
                            )

        try:
            with transaction.atomic():
                # Handle condition changes
                if condition == 'Not Usable':
                    # Validate if asset is assigned and alternate provided
                    if original_assert.emp_id and not alternate_laptop:
                        return Response(
                            {"error": "Alternate laptop required for assigned assets marked Not Usable"},
                            status=status.HTTP_400_BAD_REQUEST
                        )

                    # Handle asset swap if alternate provided
                    if alternate_laptop:
                        try:
                            alt_assert = AssertStock.objects.get(assert_id=alternate_laptop)
                            
                            # Validate alternate laptop
                            if alt_assert.emp_id:
                                return Response(
                                    {"error": "Alternate laptop already assigned"},
                                    status=status.HTTP_400_BAD_REQUEST
                                )
                            if AssertIssue.objects.filter(assert_id=alt_assert, condition='Not Usable').exists():
                                return Response(
                                    {"error": "Alternate laptop is under repair"},
                                    status=status.HTTP_400_BAD_REQUEST
                                )

                            # Get original assignment details
                            original_user = original_assert.emp_id

                            # Update UserData for both assets
                            if original_user:
                                # Update user's asset code to alternate
                                UserData.objects.filter(emp_id=original_user).update(asset_code=alt_assert)
                                # Clear original asset assignment
                                UserData.objects.filter(asset_code=original_assert).update(asset_code=None)

                            # Update assignment history
                            original_assert.update_assignment_history(None)
                            alt_assert.update_assignment_history(original_user.emp_id if original_user else None)

                            # Swap assets in AssertStock
                            alt_assert.emp_id = original_user
                            alt_assert.user = original_user.user if original_user else None
                            alt_assert.inhand = False
                            alt_assert.save()

                            AssertAssignmentLog.objects.create(
                            asset=alt_assert,
                            event_type='ASSIGNED',
                            user=alt_assert.user,
                            emp=alt_assert.emp_id,
                            notes="Asset is assigned "
                            )

                            original_assert.emp_id = None
                            original_assert.user = None
                            original_assert.inhand = True
                            original_assert.save()

                            AssertAssignmentLog.objects.create(
                            asset=original_assert,
                            event_type='REPAIR_RETURNED',
                            user=original_user.user,
                            emp=original_assert.emp_id,
                            notes="Asset returned to inventory and assigned to Null "
                        )

                        except AssertStock.DoesNotExist:
                            return Response(
                                {"error": "Alternate laptop not found"},
                                status=status.HTTP_404_NOT_FOUND
                            )

                # Update the AssertIssue record
                updated_issue = serializer.save()
                
                # Update assignment history if emp_id changed
                if updated_issue.assert_id.emp_id:
                    updated_issue.assert_id.update_assignment_history(updated_issue.assert_id.emp_id.emp_id)

                # Log the activity
                log_activity(
                    request.user,
                    "Updated",
                    "AssertIssue",
                    updated_issue.id,
                    f"Updated issue for {original_assert.assert_id}",
                    request
                )

                return Response(serializer.data)

        except Exception as e:
            logger.error(f"Error updating AssertIssue: {str(e)}")
            return Response(
                {"error": "Internal server error"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


    @transaction.atomic
    def patch(self, request, pk):
        """
        Partially update an AssertIssue instance, typically is_deleted.
        """
        try:
            assert_issue = AssertIssue.objects.get(pk=pk)
        except AssertIssue.DoesNotExist:
            return Response({"error": "AssertIssue not found"}, status=status.HTTP_404_NOT_FOUND)

        serializer = AssertIssueSerializer(assert_issue, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" :
            return Response("Intern has no access to this")
        
        try:
            assert_issue = AssertIssue.objects.get(pk=pk)

            # Log activity before deleting
            log_activity(request.user, "Deleted", "AssertIssue", assert_issue.id,
                         f"Deleted AssertIssue record", request)

            assert_issue.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except AssertIssue.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)

class AssertIssueView(APIView):
    permission_classes = [IsAuthenticated, StaffAssertAccessPermission]
    def get(self, request, pk=None):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern":
            return Response("intern has no access to this")
        if pk:
            try:
                assert_issue = AssertIssue.objects.get(pk=pk,is_deleted=False)
                serializer = AssertIssueSerializer(assert_issue)
                return Response(serializer.data)
            except AssertIssue.DoesNotExist:
                return Response(status=status.HTTP_404_NOT_FOUND)
        else:
            assert_issues = AssertIssue.objects.filter(is_deleted=False)
            serializer = AssertIssueSerializer(assert_issues, many=True)
            return Response(serializer.data)

    @transaction.atomic
    def post(self, request):
        user = Temp.objects.get(user=request.user)
        

        serializer = AssertIssueSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        data = serializer.validated_data
        assert_id = data['assert_id']
        condition = data['condition']
        alternate_laptop = request.data.get('alternate_laptop')
        # if AssertIssue.objects.filter(assert_id=assert_id , is_deleted=False).exists():
        #     return Response("Already there is a entry for this assert",status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Validate if alternate laptop is required
            if condition == 'Not Usable':
                # Check if asset is assigned to a user
                if assert_id.emp_id:
                    if not alternate_laptop:
                        return Response(
                            {"error": "Alternate laptop is required when marking an assigned asset as Not Usable"},
                            status=status.HTTP_400_BAD_REQUEST
                        )

            # Main transaction block
            with transaction.atomic():
                # Handle asset swap and user data updates
                if condition == 'Not Usable' and alternate_laptop:
                    # Validate alternate laptop
                    try:
                        alt_assert = AssertStock.objects.get(assert_id=alternate_laptop,is_deleted=False)
                        if alt_assert.emp_id:
                            return Response(
                                {"error": "Alternate laptop is already assigned"},
                                status=status.HTTP_400_BAD_REQUEST
                            )
                        
                        if AssertIssue.objects.filter(assert_id=alt_assert, condition='Not Usable',is_deleted=False).exists():
                            return Response(
                                {"error": "Alternate laptop is already under repair"},
                                status=status.HTTP_400_BAD_REQUEST
                            )

                        # Get original asset details
                        original_assert = assert_id
                        original_user = original_assert.emp_id

                        # Swap assets in UserData
                        if original_user:
                            # Update original user's data to new asset
                            UserData.objects.filter(emp_id=original_user).update(asset_code=alt_assert)
                            
                            # Clear original asset assignment
                            UserData.objects.filter(asset_code=original_assert).update(asset_code=None)

                        # Update assignment history for both assets
                        original_assert.update_assignment_history(None)
                        alt_assert.update_assignment_history(original_user.emp_id)
                        print(original_user.emp_id)

                        AssertAssignmentLog.objects.create(
                            asset=original_assert,
                            event_type='REPAIR_RETURNED',
                            user=original_user.user,
                            emp=original_assert.emp_id,
                            notes="Asset returned to inventory and assigned to Null "
                        )



                        # Update asset records
                        alt_assert.emp_id = original_user
                        alt_assert.user = original_user.user if original_user else None
                        alt_assert.inhand = False
                        alt_assert.save()

                        AssertAssignmentLog.objects.create(
                            asset=alt_assert,
                            event_type='ASSIGNED',
                            user=alt_assert.user,
                            emp=alt_assert.emp_id,
                            notes="Asset is assigned "
                        )

                        original_assert.emp_id = None
                        original_assert.user = None
                        original_assert.inhand = True
                        original_assert.save()

                    except AssertStock.DoesNotExist:
                        return Response(
                            {"error": "Alternate laptop not found"},
                            status=status.HTTP_404_NOT_FOUND
                        )

                # Create the AssertIssue record
                assert_issue = serializer.save()
                # AssertAssignmentLog.objects.create(
                #     asset=assert_issue.assert_id,
                #     event_type='REPAIR',
                #     user=request.user,
                #     notes=f"Repair initiated: {assert_issue.issue}"
                # )
                # Update assignment history with Temp object
                print(assert_id.emp_id)
                if assert_id.emp_id:
                    assert_id.update_assignment_history(assert_id.emp_id)
                    

                
                # Log the activity with proper temp reference
                log_activity(
                    request.user,
                    "Created",
                    "AssertIssue",
                    assert_issue.id,
                    f"Created issue for {assert_id.assert_id} affecting user {assert_id.emp_id.emp_id if assert_id.emp_id else 'none'}",
                    request
                )

                return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Exception as e:
            logger.error(f"Error creating AssertIssue: {str(e)}")
            return Response(
                {"error": "Internal server error"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @transaction.atomic
    def put(self, request, pk):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern":
            return Response("Intern has no access to this", status=status.HTTP_403_FORBIDDEN)

        try:
            # Get AssertIssue using string assert_id from URL
            assert_issue = AssertIssue.objects.get(pk=pk,is_deleted=False)
        except AssertIssue.DoesNotExist:
            return Response({"error": "AssertIssue not found"}, status=status.HTTP_404_NOT_FOUND)

        serializer = AssertIssueSerializer(assert_issue, data=request.data, partial=True)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        data = serializer.validated_data
        condition = data.get('condition', assert_issue.condition)
        
        alternate_laptop = request.data.get('alternate_laptop', assert_issue.alternate_laptop)
        original_assert = assert_issue.assert_id

        if condition == 'Usable':
              
            AssertAssignmentLog.objects.create(
                                asset=original_assert,
                                event_type='ALRIGHT',
                                user=None,
                                emp=original_assert.emp_id,
                                notes="The asset repair is completed and now alright"
                            )

        try:
            with transaction.atomic():
                # Handle condition changes
                if condition == 'Not Usable':
                    # Validate if asset is assigned and alternate provided
                    if original_assert.emp_id and not alternate_laptop:
                        return Response(
                            {"error": "Alternate laptop required for assigned assets marked Not Usable"},
                            status=status.HTTP_400_BAD_REQUEST
                        )

                    # Handle asset swap if alternate provided
                    if alternate_laptop:
                        try:
                            alt_assert = AssertStock.objects.get(assert_id=alternate_laptop)
                            
                            # Validate alternate laptop
                            if alt_assert.emp_id:
                                return Response(
                                    {"error": "Alternate laptop already assigned"},
                                    status=status.HTTP_400_BAD_REQUEST
                                )
                            if AssertIssue.objects.filter(assert_id=alt_assert, condition='Not Usable').exists():
                                return Response(
                                    {"error": "Alternate laptop is under repair"},
                                    status=status.HTTP_400_BAD_REQUEST
                                )

                            # Get original assignment details
                            original_user = original_assert.emp_id

                            # Update UserData for both assets
                            if original_user:
                                # Update user's asset code to alternate
                                UserData.objects.filter(emp_id=original_user).update(asset_code=alt_assert)
                                # Clear original asset assignment
                                UserData.objects.filter(asset_code=original_assert).update(asset_code=None)

                            # Update assignment history
                            original_assert.update_assignment_history(None)
                            alt_assert.update_assignment_history(original_user.emp_id if original_user else None)

                            # Swap assets in AssertStock
                            alt_assert.emp_id = original_user
                            alt_assert.user = original_user.user if original_user else None
                            alt_assert.inhand = False
                            alt_assert.save()

                            AssertAssignmentLog.objects.create(
                            asset=alt_assert,
                            event_type='ASSIGNED',
                            user=alt_assert.user,
                            emp=alt_assert.emp_id,
                            notes="Asset is assigned "
                            )

                            original_assert.emp_id = None
                            original_assert.user = None
                            original_assert.inhand = True
                            original_assert.save()

                            AssertAssignmentLog.objects.create(
                            asset=original_assert,
                            event_type='REPAIR_RETURNED',
                            user=original_user.user,
                            emp=original_assert.emp_id,
                            notes="Asset returned to inventory and assigned to Null "
                        )

                        except AssertStock.DoesNotExist:
                            return Response(
                                {"error": "Alternate laptop not found"},
                                status=status.HTTP_404_NOT_FOUND
                            )

                # Update the AssertIssue record
                updated_issue = serializer.save()
                
                # Update assignment history if emp_id changed
                if updated_issue.assert_id.emp_id:
                    updated_issue.assert_id.update_assignment_history(updated_issue.assert_id.emp_id.emp_id)

                # Log the activity
                log_activity(
                    request.user,
                    "Updated",
                    "AssertIssue",
                    updated_issue.id,
                    f"Updated issue for {original_assert.assert_id}",
                    request
                )

                return Response(serializer.data)

        except Exception as e:
            logger.error(f"Error updating AssertIssue: {str(e)}")
            return Response(
                {"error": "Internal server error"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


    @transaction.atomic
    def patch(self, request, pk):
        """
        Partially update an AssertIssue instance, typically is_deleted.
        """
        try:
            assert_issue = AssertIssue.objects.get(pk=pk)
        except AssertIssue.DoesNotExist:
            return Response({"error": "AssertIssue not found"}, status=status.HTTP_404_NOT_FOUND)

        serializer = AssertIssueSerializer(assert_issue, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        user = Temp.objects.get(user=request.user)
        if user.role == "intern" :
            return Response("Intern has no access to this")
        
        try:
            assert_issue = AssertIssue.objects.get(pk=pk)

            # Log activity before deleting
            log_activity(request.user, "Deleted", "AssertIssue", assert_issue.id,
                         f"Deleted AssertIssue record", request)

            assert_issue.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except AssertIssue.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)



class FeesView(APIView):
    permission_classes = [IsAuthenticated,StaffPayRollPermission]

    def get(self, request, emp_id=None):
        try:
            # Get the logged-in user's role
            # Get the logged-in user's role
            user_temp = Temp.objects.get(user=request.user)
            role = user_temp.role.lower()

            # Determine which employees' fees to show based on role
            if role == 'admin':
                if emp_id:
                    employees = Temp.objects.filter(emp_id=emp_id, role='intern',is_deleted=False)  # Filter for interns only
                else:
                    employees = Temp.objects.filter(role__iexact='intern',is_deleted=False)  # Filter for interns only
            elif role == 'staff':
                staff_data = UserData.objects.get(user=request.user,is_deleted=False)
                if emp_id:
                    # Verify the employee is assigned to this staff member and is an intern
                    if not UserData.objects.filter(
                        emp_id__emp_id=emp_id,
                        department=staff_data.department,
                        emp_id__role='intern',
                          is_deleted=False  # Ensure the role is intern
                    ).exists():
                        return Response(
                            {"error": "You can only view fees for interns assigned to you."},
                            status=status.HTTP_403_FORBIDDEN
                        )
                    employees = Temp.objects.filter(emp_id=emp_id, role='intern',is_deleted=False)  # Filter for interns only
                else:
                    # Get all interns assigned to this staff member
                    assigned_interns = UserData.objects.filter(
                        department=staff_data.department,
                        emp_id__role='intern',is_deleted=False  # Ensure the role is intern
                    ).values_list('emp_id', flat=True)
                    employees = Temp.objects.filter(emp_id__in=assigned_interns,is_deleted=False)
            elif role == 'intern':
                if emp_id and str(emp_id) != str(user_temp.emp_id):
                    return Response(
                        {"error": "You can only view your own fee records."},
                        status=status.HTTP_403_FORBIDDEN
                    )
                employees = Temp.objects.filter(emp_id=user_temp.emp_id, role='intern',is_deleted=False)  # Filter for interns only
            else:
                return Response(
                    {"error": "You don't have permission to view fee records."},
                    status=status.HTTP_403_FORBIDDEN
                )

            response_data = []
            
            for employee in employees:
                # Get user data to determine domain and scheme
                # Get user data to determine domain and scheme
                try:
                    user_data = UserData.objects.get(emp_id=employee)
                except UserData.DoesNotExist:
                    continue  # Skip employees without user data

                # Get fee structure
                # Get fee structure
                try:
                    fee_structure = FeeStructure.objects.get(
                        domain=user_data.domain,
                        scheme=user_data.scheme,
                        is_deleted=False
                    )
                    total_price = fee_structure.price
                except FeeStructure.DoesNotExist:
                    fee_structure = None  # Handle missing fee structure gracefully
                    total_price = 0

                # Get all payments for this employee
                # Get all payments for this employee
                payments = Fees.objects.filter(emp_id=employee,is_deleted=False)
                print(f"🔍 Payments found for {employee.emp_id}: {payments.count()}")

                # Calculate summary
                total_paid = payments.aggregate(total=Sum('amount'))['total'] or 0
                remaining_amount = max(total_price - total_paid, 0)
                total_installments = payments.count()

                # Serialize payment details
                payment_serializer = FeesSerializer(payments, many=True)

                # Add to response
                # Serialize payment details
                payment_serializer = FeesSerializer(payments, many=True)

                # Add to response
                response_data.append({
                    "id": payments.first().id if payments.exists() else None,
                    "employee_id": employee.emp_id,
                    "employee_name": employee.user.username,
                    "domain": user_data.domain.domain if user_data.domain else None,
                    "scheme": user_data.scheme,
                    "summary": {
                        "total_fee_amount": float(total_price),
                        "total_paid": float(total_paid),
                        "remaining_amount": float(remaining_amount),
                        "total_installments": total_installments,
                        "fee_structure_id": str(fee_structure.id) if fee_structure else None
                    },
                    "payment_details": payment_serializer.data
                })

            return Response(response_data, status=status.HTTP_200_OK)
            
        except Temp.DoesNotExist:
            return Response(
                {"error": "User profile not found."},
                status=status.HTTP_404_NOT_FOUND
            )
            
        except Exception as e:
            return Response(
                {"error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


    def post(self, request):
        try:
            # Get the current user's Temp record
            user_temp = Temp.objects.get(user=request.user)
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        request_data = request.data.copy()
        role = user_temp.role.lower()

        if role == 'admin':
            # Admin can pay for any user
            emp_id = request_data.get('emp_id')
            if not emp_id:
                return Response({"error": "Employee ID is required when admin pays for a user."}, status=status.HTTP_400_BAD_REQUEST)
            
            try:
                target_temp = Temp.objects.get(emp_id=emp_id)
                target_user_data = UserData.objects.get(emp_id=target_temp)
                request_data['domain'] = target_user_data.domain.id
                request_data['scheme'] = target_user_data.scheme
                request_data['user'] = target_temp.user.id
            except Temp.DoesNotExist:
                return Response({"error": "User not found with the provided Employee ID."}, status=status.HTTP_404_NOT_FOUND)
            except UserData.DoesNotExist:
                return Response({"error": "User profile data not found."}, status=status.HTTP_404_NOT_FOUND)

        elif role == 'staff':
            emp_id = request_data.get('emp_id')
            if not emp_id:
                return Response(
                    {"error": "Employee ID is required when staff pays for an intern."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            try:
                # Get staff's department
                staff_data = UserData.objects.get(user=request.user)
                if not staff_data.department:
                    return Response(
                        {"error": "Staff department not configured. Please contact admin."},
                        status=status.HTTP_400_BAD_REQUEST
                    )

                # Get intern details
                intern_temp = Temp.objects.get(emp_id=emp_id)
                if intern_temp.role.lower() != 'intern':
                    return Response(
                        {"error": "Staff can only pay for interns"},
                        status=status.HTTP_403_FORBIDDEN
                    )

                # Verify intern is in staff's department
                try:
                    intern_data = UserData.objects.get(emp_id=intern_temp)
                    if intern_data.department != staff_data.department:
                        return Response(
                            {"error": "You can only pay for interns in your department"},
                            status=status.HTTP_403_FORBIDDEN
                        )
                except UserData.DoesNotExist:
                    return Response(
                        {"error": "Intern profile data not found"},
                        status=status.HTTP_404_NOT_FOUND
                    )

                # Set domain/scheme from intern's profile
                request_data['domain'] = intern_data.domain.id
                request_data['scheme'] = intern_data.scheme
                request_data['user'] = intern_temp.user.id
                request._full_data = request_data

            except Temp.DoesNotExist:
                return Response(
                    {"error": "Intern not found with provided Employee ID"},
                    status=status.HTTP_404_NOT_FOUND
                )
                
        elif role == 'intern':
            # Interns can only pay for themselves
            try:
                user_data = UserData.objects.get(user=request.user)
                request_data['emp_id'] = user_temp.emp_id
                request_data['domain'] = user_data.domain.id
                request_data['scheme'] = user_data.scheme
                request_data['user'] = request.user.id
            except UserData.DoesNotExist:
                return Response({"error": "User profile data not found."}, status=status.HTTP_404_NOT_FOUND)

        else:
            return Response({"error": "You don't have permission to create fee records."}, status=status.HTTP_403_FORBIDDEN)

        # Get domain and scheme from the updated request data
        domain_id = request_data.get('domain')
        scheme = request_data.get('scheme')

        # Validate fee structure based on domain and scheme
        try:
            fee_structure = FeeStructure.objects.get(domain=domain_id, scheme=scheme)
            request_data['fee_structure'] = str(fee_structure.id)
        except FeeStructure.DoesNotExist:
            return Response({"error": f"No fee structure found for domain {domain_id} and scheme {scheme}."}, status=status.HTTP_400_BAD_REQUEST)

        # Validate payment amount
        amount = request_data.get('amount')
        if amount is not None:
            try:
                amount = float(amount)
                if amount <= 0:
                    return Response({"error": "Amount must be greater than zero."}, status=status.HTTP_400_BAD_REQUEST)
            except ValueError:
                return Response({"error": "Amount must be a valid number."}, status=status.HTTP_400_BAD_REQUEST)

        serializer = FeesSerializer(data=request_data)
        if serializer.is_valid():
            domain = serializer.validated_data.get('domain')
            fee_structure = serializer.validated_data.get('fee_structure')
            user = serializer.validated_data.get('user')

            # Calculate total price and already paid amount
            total_price = fee_structure.price
            paid_amount = Fees.objects.filter(user=user, domain=domain, fee_structure=fee_structure).aggregate(total_paid=models.Sum('amount'))['total_paid'] or 0

            # Calculate remaining amount and validate payment does not exceed it
            remaining_amount = total_price - paid_amount
            current_amount = serializer.validated_data.get('amount')

            # if current_amount > remaining_amount:
            #     return Response({"error": f"Cannot pay excess amount. Remaining amount is {remaining_amount}."}, status=status.HTTP_400_BAD_REQUEST)

            # Save the payment record after validation passes
            fee_record = serializer.save()

            # Log activity
            log_activity(request.user, "Created", "Fees", fee_record.id, f"Created fee payment of {current_amount} for {user}", request)

            # Return response with payment summary
            response_data = serializer.data
            response_data.update({
                'total_price': float(total_price),
                'total_paid': float(paid_amount + current_amount),
                'remaining_amount': float(remaining_amount - current_amount)
            })
            
            return Response(response_data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


    def put(self, request, pk):
        try:
            user_temp = Temp.objects.get(user=request.user)
            if user_temp.role.lower() != 'admin':
                return Response({"error": "Only admins can update fee records."},
                            status=status.HTTP_403_FORBIDDEN)
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        try:
            fee = Fees.objects.get(pk=pk)
        except Fees.DoesNotExist:
            return Response({"error": "Fee record not found."}, status=status.HTTP_404_NOT_FOUND)

        # Validate amount is positive
        amount = request.data.get('amount')
        if amount is not None:
            try:
                amount = float(amount)
                if amount <= 0:
                    return Response(
                        {"error": "Amount must be greater than zero."},
                        status=status.HTTP_400_BAD_REQUEST
                    )
            except ValueError:
                return Response(
                    {"error": "Amount must be a valid number."},
                    status=status.HTTP_400_BAD_REQUEST
                )

        # If updating fee_structure or domain, recalculate total price
        fee_structure_id = request.data.get('fee_structure')
        domain_id = request.data.get('domain')
        
        try:
            # Get the current or new fee structure
            fee_structure = FeeStructure.objects.get(pk=fee_structure_id) if fee_structure_id else fee.fee_structure
        except FeeStructure.DoesNotExist:
            return Response(
                {"error": f"Fee structure with ID {fee_structure_id} not found."},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # Get the current or new domain
            domain = Domain.objects.get(domain=domain_id) if domain_id else fee.domain
        except Domain.DoesNotExist:
            return Response(
                {"error": f"Domain with ID {domain_id} not found."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Calculate how much has been paid already (excluding this fee)
        paid_amount = Fees.objects.filter(
            user=fee.user,
            domain=domain,
            fee_structure=fee_structure
        ).exclude(id=fee.id).aggregate(total_paid=models.Sum('amount'))['total_paid'] or 0

        # Calculate remaining amount
        total_price = fee_structure.price
        remaining_amount = total_price - paid_amount

        # Check if new amount exceeds remaining amount
        new_amount = float(amount) if amount is not None else fee.amount
        if new_amount > remaining_amount:
            return Response(
                {"error": f"Cannot pay excess amount. Remaining amount is {remaining_amount}."},
                status=status.HTTP_400_BAD_REQUEST
            )

        serializer = FeesSerializer(fee, data=request.data, partial=True)
        if serializer.is_valid():
            updated_fee = serializer.save()
            
            # Log activity
            log_activity(request.user, "Updated", "Fees", fee.id,
                        f"Updated Fees record", request)
            
            # Provide comprehensive response with payment summary
            response_data = serializer.data
            response_data.update({
                'total_price': float(total_price),
                'total_paid': float(paid_amount + new_amount),
                'remaining_amount': float(remaining_amount - new_amount)
            })
            
            return Response(response_data)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    
    def delete(self, request, pk):
        try:
            user_temp = Temp.objects.get(user=request.user)
            if user_temp.role.lower() != 'admin':
                return Response({"error": "Only admins can delete fee records."},
                            status=status.HTTP_403_FORBIDDEN)
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        try:
            fee = Fees.objects.get(pk=pk)
            # Log activity before deleting
            log_activity(request.user, "Deleted", "Fees", fee.id,
                    f"Deleted Fees record", request)
            fee.delete()
            return Response({"message": "Fee record successfully deleted."},
                        status=status.HTTP_204_NO_CONTENT)
        except Fees.DoesNotExist:
            return Response({"error": "Fee record not found."},
                        status=status.HTTP_404_NOT_FOUND)
        

class FeeStructureView(APIView):
    permission_classes = [IsAuthenticated,StaffPayRollPermission]    
    def get(self, request, pk=None):
        # All authenticated users can access the GET method
        # No role checking needed for GET
        
        # Retrieve fee structure(s)
        if pk:
            try:
                fee_structure = FeeStructure.objects.get(pk=pk,is_deleted=False)
                serializer = FeeStructureSerializer(fee_structure)
                return Response(serializer.data)
            except FeeStructure.DoesNotExist:
                return Response({"error": "Fee structure not found."}, 
                               status=status.HTTP_404_NOT_FOUND)
        else:
            fee_structures = FeeStructure.objects.filter(is_deleted=False)
            serializer = FeeStructureSerializer(fee_structures, many=True)
            return Response(serializer.data)
    
    def post(self, request):
    # Check if user is admin
        if (FeeStructure.objects.filter(domain=request.data.get("domain")).exists()) and (FeeStructure.objects.filter(domain=request.data.get("scheme")).exists()):
            return Response({"Entry already exists"})
        try:
            user = Temp.objects.get(user=request.user)
            if user.role.lower() != "admin":
                return Response({"error": "Only admins can create fee structures."}, 
                            status=status.HTTP_403_FORBIDDEN)
        except Temp.DoesNotExist:
            return Response({"error": "User role not found."}, 
                        status=status.HTTP_404_NOT_FOUND)
        
        # Check if domain is provided
        if 'domain' not in request.data or not request.data.get('domain'):
            return Response({"error": "Domain is required for fee structures."}, 
                        status=status.HTTP_400_BAD_REQUEST)
        
        # Set price to 0 if scheme is FREE
        data = request.data.copy()
        if data.get('scheme') == 'FREE':
            data['price'] = '0.00'
        
        # Create fee structure
        serializer = FeeStructureSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            # Log activity
            log_activity(request.user, "Created", "FeeStructure", serializer.data['id'],
                        f"Created FeeStructure record", request)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    # def put(self, request, pk):
    #     # Check if user is admin
    #     try:
    #         user = Temp.objects.get(user=request.user)
    #         if user.role.lower() != "admin":
    #             return Response({"error": "Only admins can update fee structures."}, 
    #                            status=status.HTTP_403_FORBIDDEN)
    #     except Temp.DoesNotExist:
    #         return Response({"error": "User role not found."}, 
    #                        status=status.HTTP_404_NOT_FOUND)
        
    #     # Update fee structure
    #     try:
    #         fee_structure = FeeStructure.objects.get(pk=pk)
    #     except FeeStructure.DoesNotExist:
    #         return Response({"error": "Fee structure not found."}, 
    #                        status=status.HTTP_404_NOT_FOUND)
        
    #     # Set price to 0 if scheme is FREE
    #     data = request.data.copy()
    #     if data.get('scheme') == 'FREE':
    #         data['price'] = '0.00'
        
    #     serializer = FeeStructureSerializer(fee_structure, data=data)
    #     if serializer.is_valid():
    #         serializer.save()
    #         # Log activity
    #         log_activity(request.user, "Updated", "FeeStructure", fee_structure.id,
    #                     f"Updated FeeStructure record", request)
    #         return Response(serializer.data)
    #     return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def patch(self, request, pk):
        # Check if user is admin
        try:
            user = Temp.objects.get(user=request.user)
            if user.role.lower() != "admin":
                return Response({"error": "Only admins can update fee structures."}, 
                               status=status.HTTP_403_FORBIDDEN)
        except Temp.DoesNotExist:
            return Response({"error": "User role not found."}, 
                           status=status.HTTP_404_NOT_FOUND)
        
        # Partially update fee structure
        try:
            fee_structure = FeeStructure.objects.get(pk=pk)
        except FeeStructure.DoesNotExist:
            return Response({"error": "Fee structure not found."}, 
                           status=status.HTTP_404_NOT_FOUND)
        
        # Create a mutable copy of the request data
        data = request.data.copy()
        
        # Handling scheme changes
        if 'scheme' in data:
            if data['scheme'] == 'FREE':
                data['price'] = '0.00'
        # Handle cases where existing fee structure has scheme FREE but price is being updated
        elif fee_structure.scheme == 'FREE' and 'price' in data:
            data['price'] = '0.00'  # Force price to 0 if scheme is already FREE
        
        serializer = FeeStructureSerializer(fee_structure, data=data, partial=True)
        if serializer.is_valid():
            serializer.save()
            # Log activity
            log_activity(request.user, "Updated", "FeeStructure", fee_structure.id,
                        f"Partially updated FeeStructure record", request)
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, pk):
        # Check if user is admin
        try:
            user = Temp.objects.get(user=request.user)
            if user.role.lower() != "admin":
                return Response({"error": "Only admins can delete fee structures."}, 
                               status=status.HTTP_403_FORBIDDEN)
        except Temp.DoesNotExist:
            return Response({"error": "User role not found."}, 
                           status=status.HTTP_404_NOT_FOUND)
        
        # Delete fee structure
        try:
            fee_structure = FeeStructure.objects.get(pk=pk)
            # Log activity before deleting
            log_activity(request.user, "Deleted", "FeeStructure", fee_structure.id,
                        f"Deleted FeeStructure record", request)
            fee_structure.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except FeeStructure.DoesNotExist:
            return Response({"error": "Fee structure not found."}, 
                           status=status.HTTP_404_NOT_FOUND)







# views.py
class LogView(APIView):
    def get(self, request, emp_id=None):
        user = request.user
        try:
            user_data = Temp.objects.get(user=user,is_deleted=False)
        except Temp.DoesNotExist:
            return Response(
                {"message": "User does not exist."},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        role = user_data.role.lower()
        
        # Interns cannot access logs
        if role == "intern":
            return Response(
                {"message": "Interns do not have access to logs."},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # If emp_id is provided, validate access based on role
        if emp_id:
            try:
                staff_data = UserData.objects.get(user=request.user,is_deleted=False)
                     
                target_temp = Temp.objects.get(emp_id=emp_id,is_deleted=False)
                
                # Staff can only access logs of interns assigned to them
                if role == "staff":
                    # Check if the intern is assigned to this staff
                    is_assigned = UserData.objects.filter(
                        emp_id=target_temp, 
                        department=staff_data.department,
                        is_deleted=False
                    ).exists()
                    
                    if not is_assigned:
                        return Response(
                            {"message": "This intern is not assigned to you."},
                            status=status.HTTP_403_FORBIDDEN
                        )
                
                # Get logs for the specified emp_id
                logs = Log.objects.filter(user_id=target_temp.user,is_deleted=False)
                
            except Temp.DoesNotExist:
                return Response(
                    {"message": "Employee ID not found."},
                    status=status.HTTP_404_NOT_FOUND
                )
        else:
            # No emp_id provided, return logs based on role
            if role == "staff":
                staff_data = UserData.objects.get(user=request.user,is_deleted=False)
                     
                # Staff can see logs of all assigned interns
                assigned_interns = UserData.objects.filter(department=staff_data.department,is_deleted=False).values_list('emp_id__user', flat=True)
                logs = Log.objects.filter(user_id__in=assigned_interns,is_deleted=False)
            elif role in ["admin", "hr"]:
                # Admin and HR can see all logs
                logs = Log.objects.filter(is_deleted=False)
            else:
                # Unknown role
                return Response(
                    {"message": "Invalid user role."},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        serializer = LogSerializer(logs, many=True)
        return Response(serializer.data)



class TaskView(APIView):
    permission_classes = [IsAuthenticated,StaffUserDataAccessPermission]    
    def get(self, request, pk=None):
        user = request.user

        try:
            temp = Temp.objects.get(user=user,is_deleted=False)  # Get the user's Temp record
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        role = temp.role.lower()

        if pk:
            # Fetch a specific task by ID
            try:
                task = Task.objects.get(pk=pk,is_deleted=False)
                if role == "admin":
                    # Admin can view any task
                    serializer = TaskSerializer(task)
                    return Response(serializer.data)
                elif role == "staff":
                    # Staff can view tasks only for interns in their department
                    staff_data = UserData.objects.get(user=user,is_deleted=False)
                    assignee_data = UserData.objects.get(user=task.assigned_to,is_deleted=False)
                    if assignee_data.department == staff_data.department:
                        serializer = TaskSerializer(task)
                        return Response(serializer.data)
                    else:
                        return Response({"error": "This task is not for an intern in your department."},
                                        status=status.HTTP_403_FORBIDDEN)
                elif role == "intern":
                    # Interns can view tasks assigned to them or tasks they assigned
                    if task.assigned_to == user | task.assigned_by == user:
                        serializer = TaskSerializer(task)
                        return Response(serializer.data)
                    else:
                        return Response({"error": "You don't have permission to view this task."},
                                        status=status.HTTP_403_FORBIDDEN)
                else:
                    return Response({"error": "Invalid role."}, status=status.HTTP_403_FORBIDDEN)
            except Task.DoesNotExist:
                return Response({"error": "Task not found."}, status=status.HTTP_404_NOT_FOUND)

        # Fetch tasks based on the user's role
        if role == "admin":
            tasks = Task.objects.filter(is_deleted=False)  # Admins can view all tasks
        elif role == "staff":
            try:
                staff_data = UserData.objects.get(user=user,is_deleted=False)
                interns_in_department = UserData.objects.filter(
                    department=staff_data.department,
                    emp_id__role__iexact='intern',
                    is_deleted=False
                ).values_list('emp_id__user', flat=True)
                tasks = Task.objects.filter(assigned_to__in=interns_in_department,is_deleted=False)  # Tasks for interns in the department
            except UserData.DoesNotExist:
                return Response({"error": "Staff department not found."}, status=status.HTTP_404_NOT_FOUND)
        elif role == "intern":
            # Interns can see tasks assigned to them or tasks they assigned
            tasks_assigned_to_me = Task.objects.filter(assigned_to=user,is_deleted=False)
            tasks_assigned_by_me = Task.objects.filter(assigned_by=user,is_deleted=False)
            tasks = tasks_assigned_to_me | tasks_assigned_by_me  # Combine both querysets
        else:
            return Response({"error": "Invalid role."}, status=status.HTTP_403_FORBIDDEN)

        serializer = TaskSerializer(tasks, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)

    # Rest of the class methods (post, put, delete) remain unchanged...    

    def post(self, request):
            user = request.user
            try:
                user_temp = Temp.objects.get(user=user)
                user_role = user_temp.role.lower()
            except Temp.DoesNotExist:
                return Response({"error": "User role information not found"}, 
                            status=status.HTTP_404_NOT_FOUND)

            request_data = request.data.copy()
            assigned_to_username = request_data.get('assigned_to')

            # Common validation for all roles
            if not assigned_to_username:
                return Response({"error": "assigned_to field is required"},
                            status=status.HTTP_400_BAD_REQUEST)

            try:
                assigned_to_user = User.objects.get(username=assigned_to_username)
                assigned_to_temp = Temp.objects.get(user=assigned_to_user)
            except (User.DoesNotExist, Temp.DoesNotExist):
                return Response({"error": "Assigned user not found"},
                            status=status.HTTP_404_NOT_FOUND)

            # Get department information for both users
            try:
                assigned_to_data = UserData.objects.get(user=assigned_to_user)
                requestor_data = UserData.objects.get(user=user)
            except UserData.DoesNotExist:
                return Response({"error": "Department information missing"},
                            status=status.HTTP_400_BAD_REQUEST)

            # Department validation for interns and staff
            if user_role in ['intern', 'staff']:
                if requestor_data.department != assigned_to_data.department:
                    return Response(
                        {"error": "Can only assign tasks within your department"},
                        status=status.HTTP_403_FORBIDDEN
                    )

            # Set assigned_by to current user
            request_data['assigned_by'] = user.username

            # Additional checks for interns
            if user_role == 'intern':
                # Interns can only create tasks for other interns
                if assigned_to_temp.role.lower() != 'intern':
                    return Response(
                        {"error": "Interns can only assign tasks to other interns"},
                        status=status.HTTP_403_FORBIDDEN
                    )

            # Proceed with task creation
            serializer = TaskSerializer(data=request_data)
            if serializer.is_valid():
                task = serializer.save()
                log_activity(user, "Created", "Task", task.id,
                        f"Created task: {task.task_title}", request)
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


    def put(self, request, pk=None):
        user = request.user
        try:
            task = Task.objects.get(pk=pk)
            user_temp = Temp.objects.get(user=user)
        except (Task.DoesNotExist, Temp.DoesNotExist):
            return Response({"error": "Task or user not found"},
                          status=status.HTTP_404_NOT_FOUND)

        # Admin can update any task
        if user_temp.role.lower() == 'admin':
            serializer = TaskSerializer(task, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        # Staff/Intern validation
        try:
            requestor_data = UserData.objects.get(user=user)
            assigned_to_data = UserData.objects.get(user=task.assigned_to)
        except UserData.DoesNotExist:
            return Response({"error": "Department information missing"},
                          status=status.HTTP_400_BAD_REQUEST)

        # Verify department match
        if requestor_data.department != assigned_to_data.department:
            return Response(
                {"error": "Cannot update tasks outside your department"},
                status=status.HTTP_403_FORBIDDEN
            )

        # Intern-specific checks
        if user_temp.role.lower() == 'intern':
            # Interns can only update tasks they created
            if task.assigned_to != user:
            # | task.assigned_by != user:
                return Response(
                    {"error": "Can only update tasks created for you "},
                    status=status.HTTP_403_FORBIDDEN
                )
            
            # Prevent changing assigned_to to different department
            new_assigned_to = request.data.get('assigned_to')
            if new_assigned_to:
                try:
                    new_user = User.objects.get(username=new_assigned_to)
                    new_user_data = UserData.objects.get(user=new_user)
                    if new_user_data.department != requestor_data.department:
                        return Response(
                            {"error": "Can only assign within your department"},
                            status=status.HTTP_403_FORBIDDEN
                        )
                except (User.DoesNotExist, UserData.DoesNotExist):
                    return Response({"error": "Invalid assigned_to user"},
                                  status=status.HTTP_400_BAD_REQUEST)

        # Proceed with update
        serializer = TaskSerializer(task, data=request.data, partial=True)
        if serializer.is_valid():
            updated_task = serializer.save()
            log_activity(user, "Updated", "Task", task.id,
                       f"Updated task: {task.task_title}", request)
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


    def delete(self, request, pk=None):
        try:
            task = Task.objects.get(pk=pk)
        except Task.DoesNotExist:
            return Response({"error": "Task not found."}, status=status.HTTP_404_NOT_FOUND)

        try:
            user_temp = Temp.objects.get(user=request.user)
        except Temp.DoesNotExist:
            return Response({"error": "User role information not found."}, 
                        status=status.HTTP_404_NOT_FOUND)

        # Allow deletion based on different roles
        if user_temp.role.lower() == 'admin':
            # Admin can delete any task
            pass
        elif user_temp.role.lower() == 'staff':
            # Staff can only delete tasks in their department
            staff_data = UserData.objects.get(user=request.user)
            assignee_data = UserData.objects.get(user=task.assigned_to)
            if staff_data.department != assignee_data.department:
                return Response(
                    {"error": "You can only delete tasks in your department"},
                    status=status.HTTP_403_FORBIDDEN
                )
        elif user_temp.role.lower() == 'intern':
            # Interns can only delete tasks they created
            if task.assigned_by != request.user:
                return Response(
                    {"error": "You can only delete tasks you created"},
                    status=status.HTTP_403_FORBIDDEN
                )
        else:
            return Response({"error": "Invalid user role"}, 
                        status=status.HTTP_403_FORBIDDEN)

        # Log activity before deletion
        log_activity(
            request.user,
            "Deleted", 
            "Task",
            task.id,
            f"Deleted task: {task.task_title} (Created by {task.assigned_by.username})",
            request
        )
        
        task.delete()
        return Response(
            {"message": "Task deleted successfully"},
            status=status.HTTP_204_NO_CONTENT
        )


# class LeaveRequestView(APIView):
# permission_classes = [IsAuthenticated] # Ensure user is logged in
# def post(self, request):
# """Intern applies for leave"""
# user = request.user
# temp = Temp.objects.filter(user=user).first()
# if not temp:
# return Response({"error": "User is not associated with an employee ID."}, status=status.HTTP_400_BAD_REQUEST)
# if temp.role.lower() != "intern":
# return Response({"error": "Only interns can apply for leave."}, status=status.HTTP_403_FORBIDDEN)
# # Auto-fetch emp_id based on logged-in user
# leave_request_data = {
# "name": user.username,
# "emp_id": temp.emp_id,
# "apply_date": timezone.now().date(),
# "number_of_days": request.data.get("number_of_days"),
# "from_date": request.data.get("from_date"),
# "to_date": request.data.get("to_date"),
# "request": request.data.get("request"),
# "leave_request": True # Mark as leave request
# }
# serializer = AttendanceSerializer(data=leave_request_data)
# if serializer.is_valid():
# serializer.save()
# return Response(serializer.data, status=status.HTTP_201_CREATED)
# return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
# def patch(self, request, emp_id=None):
# """admin/staff approve or reject leave"""
# if emp_id:
# try:
# leave_request = Attendance.objects.filter(emp_id=emp_id, leave_request=True).latest('apply_date')
# except Attendance.DoesNotExist:
# return Response({"error": "No leave request found for this employee."}, status=status.HTTP_404_NOT_FOUND)
# user = request.user
# temp = Temp.objects.filter(user=user).first()
# if not temp:
# return Response({"error": "User is not associated with an employee ID."}, status=status.HTTP_400_BAD_REQUEST)
# # Only admin/staff can approve/reject leave requests
# if temp.role.lower() in ['admin', 'staff']:
# if "status" in request.data:
# leave_request.status = request.data.get("status")
# # If leave is approved and the employee is an intern, mark as Absent
# if leave_request.status == "Approved" and leave_request.emp_id.role.lower() == "intern":
# leave_request.status = "Absent"
# leave_request.save()
# serializer = AttendanceSerializer(leave_request)
# return Response(serializer.data, status=status.HTTP_200_OK)
# return Response({"error": "Only admin/staff can approve/reject leave requests."}, status=status.HTTP_403_FORBIDDEN)
# return Response({"error": "Employee ID is required for updates."}, status=status.HTTP_400_BAD_REQUEST)


# class LeaveRequestView(APIView):
#     permission_classes = [IsAuthenticated]  # Ensure user is logged in

#     def post(self, request):
#         """Intern applies for leave"""
#         user = request.user
#         temp = Temp.objects.filter(user=user).first()
#         if not temp:
#             return Response({"error": "User is not associated with an employee ID."}, status=status.HTTP_400_BAD_REQUEST)
#         if temp.role.lower() != "intern":
#             return Response({"error": "Only interns can apply for leave."}, status=status.HTTP_403_FORBIDDEN)

#         # Auto-fetch emp_id based on logged-in user
#         leave_request_data = {
#             "name": user.username,
#             "emp_id": temp.emp_id,
#             "apply_date": timezone.now().date(),
#             "number_of_days": request.data.get("number_of_days"),
#             "from_date": request.data.get("from_date"),
#             "to_date": request.data.get("to_date"),
#             "request": request.data.get("request"),
#             "status": "Pending",
#             "leave_request": True  # Mark as leave request
#         }

#         serializer = AttendanceSerializer(data=leave_request_data)
#         if serializer.is_valid():
#             serializer.save()
#             # Log activity
#             log_activity(request.user, "Created", "Attendance", serializer.data['id'],
#                          f"Created Leave Request for emp_id {temp.emp_id}", request)

#             return Response(serializer.data, status=status.HTTP_201_CREATED)
#         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

#     def patch(self, request, emp_id):
#         """admin/staff approve or reject leave"""
#         try:
#             leave_request = Attendance.objects.filter(emp_id=emp_id, leave_request=True).latest('apply_date')
#         except Attendance.DoesNotExist:
#             return Response({"error": "No leave request found for this employee."}, status=status.HTTP_404_NOT_FOUND)

#         user = request.user
#         temp = Temp.objects.filter(user=user).first()
#         if not temp:
#             return Response({"error": "User is not associated with an employee ID."}, status=status.HTTP_400_BAD_REQUEST)

#         # Only admin/staff can approve/reject leave requests
#         if temp.role.lower() in ['admin', 'staff']:
#             if "status" in request.data:
#                 leave_request.status = request.data.get("status")

#                 # If leave is approved and employee is an intern, mark as Absent
#                 if leave_request.status == "Approved" and leave_request.emp_id.role.lower() == "intern":
#                     leave_request.present_status = "Absent"
#                 elif leave_request.status == "Rejected":
#                     leave_request.present_status = "Null"  # Reset status if leave is rejected

#                 leave_request.save()

#                 # Log activity
#                 log_activity(request.user, "Updated", "Attendance", leave_request.id,
#                              f"Leave Request {leave_request.status} for emp_id {emp_id}", request)

#                 serializer = AttendanceSerializer(leave_request)
#                 return Response(serializer.data, status=status.HTTP_200_OK)

#         return Response({"error": "Only admin/staff can approve/reject leave requests."}, status=status.HTTP_403_FORBIDDEN)


class LoginView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        username = request.data.get("username")
        password = request.data.get("password")

        if not username or not password:
            return Response(
                {"error": "Please provide both username and password."},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            user = User.objects.get(username=username)
        except User.DoesNotExist:
            return Response(
                {"error": "Invalid username or password."},
                status=status.HTTP_401_UNAUTHORIZED
            )

        if user.check_password(password):
            token, created = Token.objects.get_or_create(user=user)
            
            # Check if user is admin and send certificates automatically
            try:
                user_temp = Temp.objects.get(user=user)
                if user_temp.role.lower() == 'admin':
                    from Sims.tasks import check_and_send_certificates
                    check_and_send_certificates()
            except Temp.DoesNotExist:
                pass
            
            return Response(
                {"token": token.key, "message": "Login successful"},
                status=status.HTTP_200_OK
            )
        else:
            return Response(
                {"error": "Invalid credentials."},
                status=status.HTTP_401_UNAUTHORIZED
            )

class LogoutView(APIView):
    permission_classes = [IsAuthenticated]  # Only authenticated users can log out

    def post(self, request):
        try:
            # Log logout activity
            log_activity(request.user, "Logged Out", "User", request.user.id,
                         f"User logged out successfully", request)

            # Delete the user's token to log them out
            Token.objects.get(user=request.user).delete()
            return Response({"message": "Logged out successfully."}, status=status.HTTP_200_OK)
        except Token.DoesNotExist:
            return Response({"error": "User is not logged in."}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class RegisterView(APIView):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]  # Ensure only authenticated users (admins) can create accounts

    def post(self, request):
        try:
            username = request.data.get("username")
            email = request.data.get("email")
            password = request.data.get("password")
            first_name = request.data.get("first_name", "")
            last_name = request.data.get("last_name", "")
            requested_role = request.data.get("role", "intern").lower()
            department_name = request.data.get("department")


            try:
                creator_temp = Temp.objects.get(user=request.user)
            except Temp.DoesNotExist:
                return Response({"error": "Your user profile is not configured."}, status=403)

            if creator_temp.role.lower() != "admin":
                if requested_role == "staff":
                    return Response({"error":"Only admin can create a staff"},status=400)

            if not username or not password:
                return Response({"error": "Username and password are required."}, status=status.HTTP_400_BAD_REQUEST)

            if User.objects.filter(username=username).exists():
                return Response({"error": "Username already exists."}, status=status.HTTP_400_BAD_REQUEST)

            # Check department requirement for interns first
            if requested_role == "intern" and (not department_name or not department_name.strip()):
                department_name = "General"  # Default department if not provided

            # If staff is registering an intern, enforce department match
            if creator_temp.role.lower() == "staff" and requested_role == "intern":
                try:
                    staff_userdata = UserData.objects.get(user=request.user)
                    staff_department = staff_userdata.department.department  # department is a Department object
                except UserData.DoesNotExist:
                    return Response({"error": "Staff department not configured."}, status=400)

                if department_name.strip().lower() != staff_department.strip().lower():
                    return Response(
                        {"error": f"Staff can only register interns in their own department ({staff_department})"},
                        status=403
                    )

            with transaction.atomic():
                # Create User
                user = User.objects.create_user(
                    username=username,
                    email=email,
                    password=password,
                    first_name=first_name,
                    last_name=last_name
                )

                # Create UserProfile
                UserProfile.objects.create(user=user)

                # Generate emp_id
                if requested_role == "intern":
                    department, created = Department.objects.get_or_create(
                        department__iexact=department_name.strip(),
                        defaults={'department': department_name.strip()}
                    )
                    emp_id = generate_emp_id(requested_role, department)
                else:
                    emp_id = generate_emp_id(requested_role)

                # Create Temp record with emp_id
                temp_obj = Temp.objects.create(user=user, emp_id=emp_id, role=requested_role)
                
                # Create UserData for interns
                if requested_role == "intern":
                    UserData.objects.create(
                        emp_id=temp_obj,
                        user=user,
                        department=department
                    )
                elif requested_role == "staff":
                    # By default, grant all staff access fields
                    UserData.objects.create(
                        emp_id=temp_obj,
                        user=user,
                        is_internmanagement_access=True,
                        is_attendance_access=True,
                        is_assert_access=True,
                        is_payroll_access=True
                    )
                
                # Generate auth token
                token = Token.objects.create(user=user)
                
                # Log activity
                log_activity(request.user, "Created", "User", user.id,
                           f"Created user {username} with role {requested_role}", request)

                return Response({
                    "message": "User registered successfully.",
                    "id": user.id,
                    "username": user.username,
                    "emp_id": emp_id,
                    "role": requested_role,
                    "token": token.key
                }, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


    def put(self, request):
            serializer = UpdatePasswordSerializer(data=request.data)
            
            if serializer.is_valid():
                user = request.user
                new_password = serializer.validated_data["new_password"]
                
                # Update password
                user.set_password(new_password)
                user.save()
                
                return Response({"message": "Password updated successfully."}, status=status.HTTP_200_OK)
            
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


    def get(self, request):
        try:
            user = request.user
            user_temp = Temp.objects.get(user=user,is_deleted=False)
            
            if user_temp.role.lower() == 'admin':
                users = User.objects.all()
            elif user_temp.role.lower() == 'staff':
                # Get staff's department
                staff_data = UserData.objects.get(user=user,is_deleted=False)
                # Show all interns in the same department
                interns_in_dept = UserData.objects.filter(
                    department=staff_data.department,
                    emp_id__role='intern',
                    is_deleted=False
                ).values_list('emp_id__user', flat=True)
                users = User.objects.filter(id__in=interns_in_dept)
            else:
                users = User.objects.none()

            serializer = UserSerialize(users, many=True)
            return Response(serializer.data)
        except Temp.DoesNotExist:
            return Response({"error": "User role not found."}, status=status.HTTP_404_NOT_FOUND)


class DashboardView(APIView):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated, IsAdmin]
    
    def get(self, request):
        # Total Users Count
        total_users = Temp.objects.filter(is_deleted=False).count()

        # Users Count by Month
        users_by_month = (
            UserData.objects.annotate(month=ExtractMonth("created_date"))
            .values("month")
            .annotate(count=Count("id"))
            .order_by("month")
        )

        # Users Count by Year
        users_by_year = (
            UserData.objects.annotate(year=ExtractYear("created_date"))
            .values("year")
            .annotate(count=Count("id"))
            .order_by("year")
        )

        # Users Count by Role (Using Temp table)
        users_by_role = (
            Temp.objects.values("role")
            .annotate(count=Count("emp_id"))
            .order_by("role")
        )

        # Users Count by User Status
        users_by_status = (
            UserData.objects.values("user_status")
            .annotate(count=Count("id"))
            .order_by("user_status")
        )

        # Users Count by Domain
        users_by_domain = (
            UserData.objects.values("domain")
            .annotate(count=Count("id"))
            .order_by("domain")
        )

        return Response(
            {
                "total_users": total_users,
                "users_by_month": list(users_by_month),
                "users_by_year": list(users_by_year),
                "users_by_role": list(users_by_role),
                "users_by_status": list(users_by_status),
                "users_by_domain": list(users_by_domain),
            },
            status=status.HTTP_200_OK)



class ResetPasswordOTPRequestView(APIView):
    permission_classes = [AllowAny]  # Allow public access

    def post(self, request):
        serializer = ResetPasswordOTPRequestSerializer(data=request.data)
        if serializer.is_valid():
            email = serializer.validated_data["email"]
            try:
                user = User.objects.get(email=email)

                # Resend OTP (Invalidate old OTPs and create a new one)
                otp_entry = PasswordResetOTP.resend_otp(user)
                
                # Send OTP via email
                send_mail(
                    "Password Reset OTP",
                    f"Your OTP for password reset is: {otp_entry.otp}",
                    settings.DEFAULT_FROM_EMAIL,
                    [email],
                    fail_silently=False,
                )

                return Response({"message": "OTP sent to your email."}, status=status.HTTP_200_OK)
            
            except User.DoesNotExist:
                return Response({"error": "No user found with this email."}, status=status.HTTP_404_NOT_FOUND)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class VerifyOTPView(APIView):
    permission_classes = [AllowAny]
    def post(self, request):
        serializer = VerifyOTPSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        email = serializer.validated_data['email']
        otp = serializer.validated_data['otp']

        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)

        # Get the most recent OTP for the user
        user_otp = PasswordResetOTP.objects.filter(user=user, otp=otp).order_by('-created_at').first()

        if not user_otp:
            return Response({"error": "Invalid OTP"}, status=status.HTTP_400_BAD_REQUEST)

        if user_otp.is_expired():
            user_otp.delete()
            return Response({"error": "OTP has expired"}, status=status.HTTP_400_BAD_REQUEST)

        # Delete OTP after successful verification
        user_otp.delete()
        
        # Generate token for password reset (if needed)
        token, _ = Token.objects.get_or_create(user=user)
        return Response({
            "message": "OTP verified. Proceed to password reset.",
            "token": token.key
        }, status=status.HTTP_200_OK)



class UpdatePasswordView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = UpdatePasswordSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            user = User.objects.get(auth_token__key=serializer.validated_data['token'])
        except User.DoesNotExist:
            return Response({"error": "Invalid token"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Set new password
        user.set_password(serializer.validated_data['new_password'])
        user.save()
        
        return Response({"message": "Password updated successfully"})




# In views.py

class DomainView(APIView):
    permission_classes = [IsAuthenticated, StaffUserDataAccessPermission]

    def get_staff_department(self, user):
        try:
            staff_data = UserData.objects.get(user=user)
            return staff_data.department
        except UserData.DoesNotExist:
            return None

    def get(self, request, pk=None):
        user = request.user
        try:
            user_temp = Temp.objects.get(user=user)
            
            if user_temp.role.lower() == 'admin':
                if pk:
                    domain = Domain.objects.get(pk=pk, is_deleted=False)
                    serializer = DomainSerializer(domain)
                    return Response(serializer.data)
                domains = Domain.objects.filter(is_deleted=False)
            
            elif user_temp.role.lower() == 'staff':
                department = self.get_staff_department(user)
                if not department:
                    return Response({"error": "Staff department not configured"}, 
                                  status=status.HTTP_400_BAD_REQUEST)
                if pk:
                    domain = Domain.objects.get(pk=pk, department=department, is_deleted=False)
                    serializer = DomainSerializer(domain)
                    return Response(serializer.data)
                domains = Domain.objects.filter(department=department, is_deleted=False)
            
            else:
                return Response({"error": "Unauthorized access"}, 
                              status=status.HTTP_403_FORBIDDEN)

            serializer = DomainSerializer(domains, many=True)
            return Response(serializer.data)

        except Domain.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)

    def post(self, request):
        user = request.user
        user_temp = Temp.objects.get(user=user)
        data = request.data.copy()

        if user_temp.role.lower() == 'staff':
            department = self.get_staff_department(user)
            if not department:
                return Response({"error": "Staff department not configured"}, 
                              status=status.HTTP_400_BAD_REQUEST)
            data['department'] = department.department  # Set department from staff's UserData

        elif user_temp.role.lower() != 'admin':
            return Response({"error": "Unauthorized access"}, 
                          status=status.HTTP_403_FORBIDDEN)

        serializer = DomainSerializer(data=data)
        if serializer.is_valid():
            # For admin users, allow manual department selection
            if user_temp.role.lower() == 'admin' and 'department' not in data:
                return Response({"error": "Department is required for admin users"},
                              status=status.HTTP_400_BAD_REQUEST)
            
            serializer.save()
            log_activity(user, "Created", "Domain", serializer.data['id'],
                       f"Created domain {serializer.data['domain']}", request)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def put(self, request, pk):
        user = request.user
        user_temp = Temp.objects.get(user=user)
        
        try:
            if user_temp.role.lower() == 'staff':
                department = self.get_staff_department(user)
                domain = Domain.objects.get(pk=pk, department=department, is_deleted=False)
            elif user_temp.role.lower() == 'admin':
                domain = Domain.objects.get(pk=pk, is_deleted=False)
            else:
                return Response({"error": "Unauthorized access"},
                              status=status.HTTP_403_FORBIDDEN)
        except Domain.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)

        serializer = DomainSerializer(domain, data=request.data, partial=True)
        if serializer.is_valid():
            # Staff can't modify department field
            if user_temp.role.lower() == 'staff' and 'department' in request.data:
                return Response({"error": "Staff cannot modify department field"},
                              status=status.HTTP_403_FORBIDDEN)
            
            serializer.save()
            log_activity(user, "Updated", "Domain", pk,
                       f"Updated domain {domain.domain}", request)
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        user = request.user
        user_temp = Temp.objects.get(user=user)
        
        try:
            if user_temp.role.lower() == 'staff':
                department = self.get_staff_department(user)
                domain = Domain.objects.get(pk=pk, department=department, is_deleted=False)
            elif user_temp.role.lower() == 'admin':
                domain = Domain.objects.get(pk=pk, is_deleted=False)
            else:
                return Response({"error": "Unauthorized access"},
                              status=status.HTTP_403_FORBIDDEN)
        except Domain.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)

        domain.is_deleted = True
        domain.save()
        log_activity(user, "Deleted", "Domain", pk,
                   f"Soft-deleted domain {domain.domain}", request)
        return Response(status=status.HTTP_204_NO_CONTENT)




class PerformanceView(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def get(self, request):
        # Automatically update tasks to 'Missing' if they are not completed within the committed date
        today = timezone.now().date()
        tasks_to_update = Task.objects.filter(Q(is_deleted=False) &
            Q(status__in=['Under_Review', 'In_Progress', 'Not_Started']) & Q(committed_date__lt=today)
        )
        tasks_to_update.update(status='Missing')

        # Fetch tasks after updating their statuses
        tasks = Task.objects.filter(is_deleted=False)
        task_status_count = tasks.values('status').annotate(count=Count('status'))

        completed_tasks = tasks.filter(status='Completed',is_deleted=False)
        completed_tasks_data = [
            {
                'task_id': task.id,
                'start_date': task.start_date,
                'end_date': task.end_date
            }
            for task in completed_tasks
        ]

        response_data = {
            'total_tasks': tasks.count(),
            'status_counts': {item['status']: item['count'] for item in task_status_count},
            'completed_tasks': completed_tasks_data
        }

        return Response(response_data)








class MonthlyTaskCountView(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]
    def get(self, request):
        # Get logged-in intern
        try:
            temp_user = Temp.objects.get(user=request.user,is_deleted=False)
            if temp_user.role.lower() != "intern":
                return Response({"error": "Only interns can access their task count."}, status=status.HTTP_403_FORBIDDEN)
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)
        # Get the current year
        current_year = now().year
        # Get completed tasks for this intern, grouped by month
        task_counts = (
            Task.objects.filter(assigned_to=request.user, status="Completed", end_date__year=current_year)
            .annotate(month=ExtractMonth("end_date"))
            .values("month")
            .annotate(count=Count("id"))
            .order_by("month")
        )
        # Convert month numbers to names
        month_names = {
            1: "January", 2: "February", 3: "March", 4: "April",
            5: "May", 6: "June", 7: "July", 8: "August",
            9: "September", 10: "October", 11: "November", 12: "December"
        }
        monthly_task_data = {month_names[t["month"]]: t["count"] for t in task_counts}
        return Response({"intern": request.user.username, "task_counts": monthly_task_data}, status=status.HTTP_200_OK)
    
class InternCountByDomainView(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def get(self, request):
        try:
            # Get current user's role
            user_temp = Temp.objects.get(user=request.user,is_deleted=False)
            user_role = user_temp.role.lower()
            
            # Deny access to interns
            if user_role == 'intern':
                return Response({"error": "Interns cannot access this resource"}, 
                              status=status.HTTP_403_FORBIDDEN)

            # Initialize queryset
            queryset = UserData.objects.filter(emp_id__role__iexact='intern')

            # For staff users, filter by department
            if user_role == 'staff':
                try:
                    staff_data = UserData.objects.get(user=request.user,is_deleted=False)
                    if not staff_data.department:
                        return Response({"error": "Staff department not configured"}, 
                                      status=status.HTTP_400_BAD_REQUEST)
                    
                    # Filter interns by staff's department
                    queryset = queryset.filter(department=staff_data.department,is_deleted=False)
                
                except UserData.DoesNotExist:
                    return Response({"error": "Staff department information not found"}, 
                                  status=status.HTTP_404_NOT_FOUND)

            # Aggregate counts by domain
            domain_counts = (
                queryset
                .values('domain__domain')
                .annotate(total=Count('id'))
                .order_by('-total')
            )

            # Format response
            results = [
                {"domain": item['domain__domain'], "count": item['total']} 
                for item in domain_counts
            ]

            # Log activity
            log_activity(request.user, "Viewed", "Intern Count", "", 
                       f"Viewed intern counts by domain", request)

            return Response(results)

        except Temp.DoesNotExist:
            return Response({"error": "User role information not found"}, 
                          status=status.HTTP_404_NOT_FOUND)










class StaffDetailsView(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def get(self, request):
        try:
            current_user = Temp.objects.get(user=request.user,is_deleted=False)
        except Temp.DoesNotExist:
            return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)

        # Enforce role-based access
        if current_user.role.lower() not in ['staff', 'admin']:
            return Response({"error": "Access denied"}, status=status.HTTP_403_FORBIDDEN)

        # If the user is staff, return only their details
        if current_user.role.lower() == 'staff':
            return self._get_staff_details(current_user)

        # If the user is admin, return all staff details
        if current_user.role.lower() == 'admin':
            return self._get_all_staff_details()

        return Response({"error": "Invalid user role"}, status=status.HTTP_400_BAD_REQUEST)

    def _get_staff_details(self, staff_temp):
        """Helper method to get details for a single staff member"""
        staff_data = {
            "emp_id": staff_temp.emp_id,
            "username": staff_temp.user.username,
            "details": {
                "personal_data": self._get_personal_data(staff_temp),
                "user_data": self._get_user_data(staff_temp),
                "college_details": self._get_college_details(staff_temp)
            }
        }
        # Returning a list with a single staff member's details
        return Response([staff_data])  # Ensure it returns a list of one object

    def _get_all_staff_details(self):
        """Helper method to get details for all staff members (admin only)"""
        all_staff = Temp.objects.filter(role__iexact='staff',is_deleted=False)
        response_data = []
        
        for staff in all_staff:
            staff_entry = {
                "emp_id": staff.emp_id,
                "username": staff.user.username,
                "details": {
                    "personal_data": self._get_personal_data(staff),
                    "user_data": self._get_user_data(staff),
                    "college_details": self._get_college_details(staff)
                }
            }
            response_data.append(staff_entry)
        return Response(response_data)

    def _get_personal_data(self, temp):
        try:
            return PersonalDataSerializer(PersonalData.objects.get(emp_id=temp,is_deleted=False)).data
        except PersonalData.DoesNotExist:
            return {"error": "Personal data not found"}

    def _get_user_data(self, temp):
        try:
            return UserDataSerializer(UserData.objects.get(emp_id=temp,is_deleted=False)).data
        except UserData.DoesNotExist:
            return {"error": "User data not found"}

    def _get_college_details(self, temp):
        try:
            return CollegeDetailsSerializer(CollegeDetails.objects.get(emp_id=temp,is_deleted=False)).data
        except CollegeDetails.DoesNotExist:
            return {"error": "College details not found"}
        




class AssertStockCountView(APIView):
    permission_classes = [IsAuthenticated,StaffAssertAccessPermission]
    def get(self, request):
        try:
            # Get logged-in user role
            user_temp = Temp.objects.get(user=request.user,is_deleted=False)
            if user_temp.role.lower() not in ['admin', 'staff']:
                return Response({"error": "Access denied. Only admin and staff can view asset count."}, status=403)

            # Count total assets
            total_assets = AssertStock.objects.filter(is_deleted=False).count()

            # Count assigned assets (where emp_id is NOT NULL)
            assigned_assets = AssertStock.objects.filter(emp_id__isnull=False,is_deleted=False).count()

            # Count available assets (where emp_id is NULL)
            available_assets = AssertStock.objects.filter(emp_id__isnull=True,is_deleted=False).count()

            return Response({
                "total_assets": total_assets,
                "assigned_assets": assigned_assets,
                "available_assets": available_assets
            }, status=200)
        except Temp.DoesNotExist:
            return Response({"error": "User role not found."}, status=404)

#-------------------------------------------------------------------------------------------
class AssetTrendView(APIView):
    permission_classes = [IsAuthenticated,StaffAssertAccessPermission]
    def get(self, request):
        try:
            # Check if the user has permission
            user_temp = Temp.objects.get(user=request.user,is_deleted=False)
            if user_temp.role.lower() not in ['admin', 'staff']:
                return Response({"error": "Access denied."}, status=403)

            # Aggregate data by month
            asset_trends = (
                AssertStock.objects
                .annotate(month=TruncMonth("updated_date"))  
                .values("month")
                .annotate(
                    total_assets=Count("assert_id"),  # ✅ Use assert_id
                    assigned_assets=Count("assert_id", filter=~Q(emp_id=None))  # ✅ Use assert_id
                )
                .order_by("month")
            )

            return Response(list(asset_trends), status=200)

        except Temp.DoesNotExist:
            return Response({"error": "User role not found."}, status=404)
        except Exception as e:
            return Response({"error": str(e)}, status=500)

class AvailableAssetCountView(APIView):
    permission_classes = [IsAuthenticated,StaffAssertAccessPermission]
    def get(self, request):
        try:
            # Count available assets where IT support is 'Hand Over'
            available_assets = AssertIssue.objects.filter(it_support="Hand Over",is_deleted=False).count()

            return Response({"available_assets": available_assets}, status=200)

        except Exception as e:
            return Response({"error": str(e)}, status=500)
        



# views.py (FeedbackView updates)
class FeedbackView(APIView):
    permission_classes = [IsAuthenticated,StaffUserDataAccessPermission]    
    authentication_classes = [TokenAuthentication]

    def get_object(self, feedback_id):
        try:
            return Feedback.objects.get(id=feedback_id,is_deleted=False)
        except Feedback.DoesNotExist:
            return None

    def get(self, request, feedback_id=None):
        user = request.user
        try:
            user_temp = Temp.objects.get(user=user,is_deleted=False)
            role = user_temp.role.lower()
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        if feedback_id:
            feedback = self.get_object(feedback_id)
            if not feedback:
                return Response({"error": "Feedback not found."}, status=status.HTTP_404_NOT_FOUND)

            # Permission checks
            if role == 'intern' and feedback.feedback_to != user_temp:
                return Response({"error": "You can only view feedback addressed to you."}, 
                              status=status.HTTP_403_FORBIDDEN)
            
            if role == 'staff':
                is_related = (user_temp == feedback.feedback_by or 
                            user_temp == feedback.feedback_to or
                            UserData.objects.filter(emp_id=feedback.feedback_to, reporting_manager=user,is_deleted=False).exists())
                if not is_related:
                    return Response({"error": "You don't have permission to view this feedback."}, 
                                  status=status.HTTP_403_FORBIDDEN)

            serializer = FeedbackSerializer(feedback)
            return Response(serializer.data)

        # Get all feedback based on role
        feedbacks = Feedback.objects.filter(is_deleted=False)
        if role == 'staff':
            feedbacks = Feedback.objects.filter(
                Q(feedback_by=user_temp) |
                Q(feedback_to=user_temp) |
                Q(is_deleted=False) |
                Q(feedback_to__in=Temp.objects.filter(
                    emp_id__in=UserData.objects.filter(reporting_manager=user,is_deleted=False).values('emp_id'),is_deleted=False
                ))
            )
        elif role == 'intern':
            feedbacks = Feedback.objects.filter(feedback_to=user_temp)

        serializer = FeedbackSerializer(feedbacks, many=True)
        return Response(serializer.data)

    def post(self, request):
        try:
            # Get logged-in user's Temp object
            feedback_by_temp = Temp.objects.get(user=request.user)
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        request_data = request.data.copy()
        request_data['feedback_by'] = feedback_by_temp.emp_id  # Auto-set sender

        # Prevent self-feedback
        feedback_to_emp_id = request_data.get('feedback_to')
        if str(feedback_by_temp.emp_id) == str(feedback_to_emp_id):
            return Response(
                {"error": "You cannot submit feedback to yourself."},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            # Get recipient's Temp object
            feedback_to_temp = Temp.objects.get(emp_id=feedback_to_emp_id)
        except Temp.DoesNotExist:
            return Response({"error": "Recipient not found."}, status=status.HTTP_404_NOT_FOUND)

        # Validate permissions based on sender's role
        if feedback_by_temp.role.lower() == 'staff':
            # Staff can only feedback to their assigned interns
            if feedback_to_temp.role.lower() != 'intern':
                return Response(
                    {"error": "You can only provide feedback to interns."},
                    status=status.HTTP_403_FORBIDDEN
                )

            # Check if staff is assigned as manager/supervisor for this intern
            try:
                user_data = UserData.objects.get(emp_id=feedback_to_temp)
                is_assigned = (user_data.reporting_manager == request.user) or \
                             (user_data.reporting_supervisor == request.user)
                
                if not is_assigned:
                    return Response(
                        {"error": "You can only provide feedback to interns assigned to you."},
                        status=status.HTTP_403_FORBIDDEN
                    )
            except UserData.DoesNotExist:
                return Response(
                    {"error": "Recipient's user data not found."},
                    status=status.HTTP_404_NOT_FOUND
                )

        # Proceed with feedback creation
        serializer = FeedbackSerializer(data=request_data)
        if serializer.is_valid():
            feedback = serializer.save()
            log_activity(
                user=request.user,
                action="Created",
                entity_type="Feedback",
                entity_id=feedback.id,
                description=f"Feedback from {feedback.feedback_by} to {feedback.feedback_to}",
                request=request
            )
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    

    
    def patch(self, request, feedback_id):
        user = request.user
        feedback = self.get_object(feedback_id)
        if not feedback:
            return Response({"error": "Feedback not found."}, status=status.HTTP_404_NOT_FOUND)

        try:
            user_temp = Temp.objects.get(user=user)
            role = user_temp.role.lower()
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        # Permission checks
        if role == 'intern' and feedback.feedback_by != user_temp:
            return Response({"error": "You can only update feedback you've created."},
                          status=status.HTTP_403_FORBIDDEN)
        
        if role == 'staff':
            is_creator = (feedback.feedback_by == user_temp)
            is_recipient = (feedback.feedback_to == user_temp)
            if not (is_creator or is_recipient):
                return Response({"error": "You can only update feedback you've created or received."},
                              status=status.HTTP_403_FORBIDDEN)
            
            if not is_creator and request.data.keys() != {'status'}:
                return Response({"error": "You can only update the status of feedback you've received."},
                              status=status.HTTP_403_FORBIDDEN)

        serializer = FeedbackSerializer(feedback, data=request.data, partial=True)
        if serializer.is_valid():
            updated_feedback = serializer.save()
            log_activity(
                user=user,
                action="Updated",
                entity_type="Feedback",
                entity_id=feedback.id,
                description=f"Updated feedback from {feedback.feedback_by} to {feedback.feedback_to}",
                request=request
            )
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, feedback_id):
        user = request.user
        feedback = self.get_object(feedback_id)
        if not feedback:
            return Response({"error": "Feedback not found."}, status=status.HTTP_404_NOT_FOUND)

        try:
            user_temp = Temp.objects.get(user=user)
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        # Only feedback creator or admin can delete
        if feedback.feedback_by != user_temp and user_temp.role.lower() != 'admin':
            return Response({"error": "You can only delete feedback you've created."},
                          status=status.HTTP_403_FORBIDDEN)

        feedback.delete()
        log_activity(
            user=user,
            action="Deleted",
            entity_type="Feedback",
            entity_id=feedback_id,
            description=f"Deleted feedback from {feedback.feedback_by} to {feedback.feedback_to}",
            request=request
        )
        return Response(status=status.HTTP_204_NO_CONTENT)





class DepartmentView(APIView):
    def get(self, request, pk=None):
        if pk:
            try:
                department = Department.objects.get(pk=pk,is_deleted=False)
                serializer = DepartmentSerializer(department)
                return Response(serializer.data)
            except Department.DoesNotExist:
                return Response(status=status.HTTP_404_NOT_FOUND)
        else:
            departments = Department.objects.filter(is_deleted=False)
            serializer = DepartmentSerializer(departments, many=True)
            return Response(serializer.data)
    
    def post(self, request):
        # Check if user has a Temp record
        try:
            user_temp = Temp.objects.get(user=request.user)
        except Temp.DoesNotExist:
            return Response(
                {"error": "User role not configured. Contact administrator."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Validate user role
        if user_temp.role.lower() not in ["admin", "hr"]:
            return Response(
                {"error": "Only admins/HR can create departments"},
                status=status.HTTP_403_FORBIDDEN
            )

        # Proceed with department creation
        serializer = DepartmentSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    

    def patch(self, request, pk):
        try:
            # Check if user is admin
            user_temp = Temp.objects.get(user=request.user)
            if user_temp.role.lower() != 'admin':
                return Response({"error": "Only admins can update departments."},
                              status=status.HTTP_403_FORBIDDEN)
            
            # Get department instance
            department = Department.objects.get(pk=pk)
            
            # Partial update
            serializer = DepartmentSerializer(department, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                # Log activity
                log_activity(request.user, "Updated", "Department", pk,
                           f"Updated department to {serializer.data['department']}", request)
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            
        except Temp.DoesNotExist:
            return Response({"error": "User role information not found."},
                          status=status.HTTP_404_NOT_FOUND)
        except Department.DoesNotExist:
            return Response({"error": "Department not found."},
                          status=status.HTTP_404_NOT_FOUND)

    def delete(self, request, pk):
        try:
            # Check if user is admin
            user_temp = Temp.objects.get(user=request.user)
            if user_temp.role.lower() != 'admin':
                return Response({"error": "Only admins can delete departments."},
                              status=status.HTTP_403_FORBIDDEN)
            
            # Get and delete department
            department = Department.objects.get(pk=pk)
            # Log activity before deletion
            log_activity(request.user, "Deleted", "Department", pk,
                       f"Deleted department: {department.department}", request)
            department.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
            
        except Temp.DoesNotExist:
            return Response({"error": "User role information not found."},
                          status=status.HTTP_404_NOT_FOUND)
        except Department.DoesNotExist:
            return Response({"error": "Department not found."},
                          status=status.HTTP_404_NOT_FOUND)
    

class UserAttendanceView(APIView):
    permission_classes = [IsAuthenticated,StaffAttendanceAccessPermission]

    def get(self, request):
        try:
            # Get the logged-in user's Temp object
            user_temp = Temp.objects.get(user=request.user,is_deleted=False)
            emp_id = user_temp.emp_id

            # Fetch attendance records for the logged-in user
            attendance_records = Attendance.objects.filter(emp_id=emp_id,is_deleted=False)

            # Calculate total absences
            total_absences = attendance_records.filter(present_status='Absent',is_deleted=False).count()

            # Serialize attendance records
            serializer = AttendanceSerializer(attendance_records, many=True)

            return Response({
                "emp_id": emp_id,
                "username": request.user.username,
                "total_absences": total_absences,
                "attendance_records": serializer.data
            }, status=200)
        except Temp.DoesNotExist:
            return Response({"error": "User is not associated with an employee ID."}, status=404)
        except Exception as e:
            return Response({"error": str(e)}, status=500)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_leave_balance(request):
    try:
        temp = Temp.objects.get(user=request.user,is_deleted=False)

        # Get the latest attendance entry — assuming it's updated regularly
        attendance = Attendance.objects.filter(emp_id=temp,is_deleted=False).order_by('-date').first()

        if not attendance:
            return Response({"error": "No attendance record found."}, status=404)

        total = attendance.max_leave_count or 0
        remaining = attendance.remaining_leave_count or 5
        used = total - remaining

        return Response({
            "used": used,
            "remaining": remaining,
            "total": total,
            "leaves_taken": attendance.leaves_taken or 0
        })

    except Temp.DoesNotExist:
        return Response({"error": "User not found"}, status=404)




@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_leave_balance(request):
    user = request.user
    total = 5 # Default total leaves
    # Sum all approved leaves
    approved_leaves = LeaveRequest.objects.filter(user=user, status="APPROVED",is_deleted=False)
    leaves_taken = approved_leaves.aggregate(Sum('number_of_days'))['number_of_days__sum'] or 0
    used = leaves_taken
    remaining = max(total - used, 0)
    return Response({
        "used": used,
        "remaining": remaining,
        "total": total,
        "leaves_taken": leaves_taken
    })





class LeaveRequestView(APIView):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated,StaffAttendanceAccessPermission]    

    def get(self, request):
        """Fetch leave requests of the logged-in user."""
        leave_requests = LeaveRequest.objects.filter(user=request.user)
        serializer = LeaveRequestSerializer(leave_requests, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
    def post(self, request):
        # Pass request context to the serializer
        serializer = LeaveRequestSerializer(
            data=request.data,
            context={'request': request} # <-- Add this line
        )
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class LeaveApprovalView(APIView):
    """
    View to handle leave approval or rejection by the reporting manager.
    Includes leave balance deduction and attendance marking logic.
    """
    permission_classes = [IsAuthenticated,StaffAttendanceAccessPermission]    

    # permission_classes = [IsAuthenticated]
    def get(self, request):
        """Get all leave requests for interns managed/supervised by current user"""
        user = request.user
        # Get all interns where user is reporting manager OR supervisor
        staff_data = UserData.objects.get(user=request.user,is_deleted=False)
        interns = UserData.objects.filter(
            department=staff_data.department,
        ).values_list('user', flat=True)
        # Get leave requests for these interns
        leave_requests = LeaveRequest.objects.filter(
            user__in=interns
        ).order_by('-created_at')
        serializer = LeaveRequestSerializer(leave_requests, many=True)
        return Response(serializer.data)
    
    def patch(self, request, leave_id):
        user = request.user
        try:
            leave_request = LeaveRequest.objects.get(id=leave_id)
            # Get the intern's UserData
            user_data = UserData.objects.get(user=leave_request.user)
            staff_data = UserData.objects.get(user=request.user,is_deleted=False)
            # Check if current user is authorized (manager OR supervisor)
            if staff_data.department != user_data.department:
                return Response({"error": "Not authorized"})
            new_status = request.data.get("status")
            if new_status == 'APPROVED':
                with transaction.atomic():
                    # Calculate total leave days
                    delta = leave_request.to_date - leave_request.from_date
                    total_days = delta.days + 1 # Inclusive of both dates
                    # Handle half-day adjustments
                    if leave_request.half_day_start and leave_request.half_day_end:
                        if leave_request.from_date == leave_request.to_date:
                            total_days = 0.5
                        else:
                            total_days -= 1.0
                    elif leave_request.half_day_start or leave_request.half_day_end:
                        total_days -= 0.5
                    # Get the Temp object for the user
                    temp = Temp.objects.get(user=leave_request.user)
                    # Get the latest attendance record for this employee
                    latest_attendance = Attendance.objects.filter(emp_id=temp).order_by('-date').first()
                    if not latest_attendance:
                        return Response({"error": "No attendance record found for this employee."}, status=404)
                    # Update remaining leaves and leaves_taken
                    remaining_leaves = latest_attendance.remaining_leave_count or 5
                    new_remaining = max(0, remaining_leaves - total_days)
                    leaves_taken = latest_attendance.leaves_taken + total_days
                    # Update all attendance records with new values
                    Attendance.objects.filter(emp_id=temp).update(
                        remaining_leave_count=new_remaining,
                        leaves_taken=leaves_taken
                    )
                    # Mark dates as absent in attendance records
                    dates_to_mark = [leave_request.from_date + timedelta(days=i)
                                    for i in range((leave_request.to_date - leave_request.from_date).days + 1)]
                    for date in dates_to_mark:
                        Attendance.objects.update_or_create(
                            emp_id=temp,
                            date=date,
                            defaults={
                                'present_status': 'Absent',
                                'remaining_leave_count': new_remaining,
                                'leaves_taken': leaves_taken,
                            }
                        )
                    # Update the leave request status
                    leave_request.status = new_status
                    leave_request.save()
                return Response({"message": "Leave approved successfully!"}, status=200)
            elif new_status == 'REJECTED':
                leave_request.status = new_status
                leave_request.save()
                return Response({"message": "Leave rejected successfully!"}, status=200)
            return Response({"error": "Invalid status"}, status=400)
        except Exception as e:
            return Response({"error": str(e)}, status=500)




class LeaveStatusView(APIView):
    permission_classes = [IsAuthenticated,StaffAttendanceAccessPermission]    

    def get(self, request):
        try:
            # Get the logged-in user's Temp profile
            temp = Temp.objects.get(user=request.user)
            
            # Ensure the user is an intern
            if temp.role.lower() != 'intern':
                return Response({"error": "Only interns can view leave statuses."}, status=403)
            
            # Fetch leave requests for the logged-in intern
            leave_requests = LeaveRequest.objects.filter(user=request.user).order_by('-created_at')
            
            # Serialize the data
            serializer = LeaveRequestSerializer(leave_requests, many=True)
            
            return Response(serializer.data, status=200)
        except Temp.DoesNotExist:
            return Response({"error": "User profile not found."}, status=404)









class AttendanceView(APIView):

    permission_classes = [IsAuthenticated,StaffAttendanceAccessPermission]    

    def get(self, request, emp_id=None):
        try:
            # Get the logged-in user's Temp record
            user_temp = Temp.objects.get(user=request.user)
            role = user_temp.role.lower()

            # Initialize response data
            response_data = []

            # Determine employees to fetch attendance for based on role
            if emp_id:
                # Single employee request
                try:
                    target_temp = Temp.objects.get(emp_id=emp_id)

                    # Permission checks for interns and staff
                    if role == "intern" and user_temp.emp_id != target_temp.emp_id:
                        return Response(
                            {"error": "You can only view your own attendance records."},
                            status=status.HTTP_403_FORBIDDEN,
                        )
                    if role == "staff":
                        staff_data = UserData.objects.get(user=request.user)
                        intern_data = UserData.objects.get(emp_id=target_temp)
                        if intern_data.department != staff_data.department:
                            return Response(
                                {"error": "This intern is not in your department."},
                                status=status.HTTP_403_FORBIDDEN,
                            )

                    employees = [target_temp]
                except Temp.DoesNotExist:
                    return Response(
                        {"error": "Employee not found."}, status=status.HTTP_404_NOT_FOUND
                    )
            else:
                # Multiple employees based on role
                if role == "intern":
                    employees = [user_temp]
                elif role == "staff":
                    staff_data = UserData.objects.get(user=request.user)
                    interns_in_department = UserData.objects.filter(
                        department=staff_data.department,
                        emp_id__role__iexact="intern",
                    ).values_list("emp_id", flat=True)
                    employees = Temp.objects.filter(emp_id__in=interns_in_department)
                elif role in ["admin", "hr"]:
                    employees = Temp.objects.all()
                else:
                    return Response(
                        {"error": "Invalid user role."},
                        status=status.HTTP_400_BAD_REQUEST,
                    )

            # Process each employee's attendance records
            for employee in employees:
                attendance_records = Attendance.objects.filter(emp_id=employee)

                # Calculate summary metrics
                present_days = attendance_records.filter(present_status="Present").count()
                absent_days = attendance_records.filter(present_status="Absent").count()
                total_days = attendance_records.count()
                pending_leave_requests = attendance_records.filter(
                    leave_request__status="Pending"
                ).count()

                # Get the most recent record for leave counts
                latest_record = attendance_records.order_by("-date").first()
                remaining_leave = (
                    latest_record.remaining_leave_count if latest_record else 5
                )
                leaves_taken = (
                    latest_record.leaves_taken if latest_record else 0
                )

                # Create employee data structure
                employee_data = {
                    "emp_id": employee.emp_id,
                    "name": employee.user.username,
                    "summary": {
                        "total_days": total_days,
                        "present_days": present_days,
                        "absent_days": absent_days,
                        "remaining_leave": remaining_leave,
                        "leaves_taken": leaves_taken,
                        "pending_leave_requests": pending_leave_requests,
                    },
                    "records": AttendanceSerializer(attendance_records, many=True).data,
                }

                response_data.append(employee_data)

            # If it was a single employee request, return just that employee's data
            if emp_id:
                return Response(response_data[0] if response_data else {}, status=status.HTTP_200_OK)

            return Response(response_data, status=status.HTTP_200_OK)

        except Temp.DoesNotExist:
            return Response(
                {"error": "User profile not found."}, status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)    

    def post(self, request):
        def get_client_ip(request):
            """Helper function to get client IP address"""
            x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
            if x_forwarded_for:
                ip = x_forwarded_for.split(',')[0]
            else:
                ip = request.META.get('REMOTE_ADDR')
            return ip

        try:
            # Get the logged-in user and their Temp record
            user_temp = Temp.objects.get(user=request.user)
        except Temp.DoesNotExist:
            return Response({"error": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        # Handle permissions and set emp_id based on role
        if user_temp.role.lower() == 'intern':
            # Interns can only enter their own attendance
            emp_id = user_temp.emp_id
            if 'emp_id' in request.data and str(request.data['emp_id']) != str(emp_id):
                return Response(
                    {"error": "You can only enter attendance for yourself."},
                    status=status.HTTP_403_FORBIDDEN,
                )
            # Update request data with intern's info
            request.data['emp_id'] = emp_id
            request.data['name'] = request.user.username

        elif user_temp.role.lower() == 'staff':
            # Staff can enter attendance for interns in the same department
            emp_id = request.data.get('emp_id')
            if not emp_id:
                return Response(
                    {"error": "Employee ID is required."},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            try:
                # Get staff's department
                staff_data = UserData.objects.get(user=request.user)
                if not staff_data.department:
                    return Response(
                        {"error": "Your department is not set. Please set your department first."},
                        status=status.HTTP_400_BAD_REQUEST,
                    )
                # Get target intern's data
                target_temp = Temp.objects.get(emp_id=emp_id)
                intern_data = UserData.objects.get(emp_id=target_temp)
                # Check if intern is in the same department
                if intern_data.department != staff_data.department:
                    return Response(
                        {"error": "You can only enter attendance for interns in your department."},
                        status=status.HTTP_403_FORBIDDEN,
                    )
                # Get the employee's name
                request.data['name'] = target_temp.user.username
            except (Temp.DoesNotExist, UserData.DoesNotExist):
                return Response(
                    {"error": "Employee not found or missing profile data."},
                    status=status.HTTP_404_NOT_FOUND,
                )

        elif user_temp.role.lower() == 'admin':
            # Admins can enter attendance for anyone
            emp_id = request.data.get('emp_id')
            if not emp_id:
                return Response(
                    {"error": "Employee ID is required."},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            try:
                temp = Temp.objects.get(emp_id=emp_id)
                request.data['name'] = temp.user.username
            except Temp.DoesNotExist:
                return Response(
                    {"error": "Employee not found."},
                    status=status.HTTP_404_NOT_FOUND,
                )
        else:
            return Response(
                {"error": "You do not have permission to create attendance records."},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Set the date for today if not provided
        date = request.data.get('date', timezone.now().date())
        request.data['date'] = date

        # Check if there's an open attendance entry for this employee on this date
        try:
            temp_obj = Temp.objects.get(emp_id=emp_id)
            attendance = Attendance.objects.filter(emp_id=temp_obj, date=date).first()
        except Temp.DoesNotExist:
            return Response({"error": "Employee not found."}, status=status.HTTP_404_NOT_FOUND)

        current_time = timezone.now()

        # If there's no existing attendance for today, create one
        if not attendance:
            # Get the employee's current remaining leave count
            latest_attendance = Attendance.objects.filter(emp_id=temp_obj).order_by('-date').first()
            remaining_leave = latest_attendance.remaining_leave_count if latest_attendance else 5

            # Handle leave requests and present_status logic
            is_leave_request = request.data.get('leave_request', False)
            present_status = request.data.get('present_status', 'Present')

            if is_leave_request or present_status == 'Absent':
                present_status = 'Absent'
                if remaining_leave > 0:
                    remaining_leave -= 1  # Decrement remaining leaves by 1
                else:
                    remaining_leave = 0  # No remaining leaves left

            else:
                present_status = 'Present'

            # Update request data with calculated values
            request.data['present_status'] = present_status
            request.data['remaining_leave_count'] = remaining_leave

            # Create the attendance record
            serializer = AttendanceSerializer(data=request.data)
            if serializer.is_valid():
                attendance = serializer.save()

                
                reason = "Live on"

                # Create initial log entry
                temp_object=Temp.objects.get(emp_id=temp_obj.emp_id)
            
                logdetail=AttendanceLog.objects.create(
                    emp_id=temp_object,
                    attendance=attendance,
                    time=timezone.now(),
                    is_in=True,
                    reason=reason,
                    ip_address= get_client_ip(request),
                )
                log_serializer1 = AttendanceLogSerializer(logdetail)
                # Create an attendance log entry for check-in
                log_data = {
                    'attendance': attendance.id,
                    'emp_id': temp_obj.emp_id,
                    'check_in': current_time,
                    'ip_address': get_client_ip(request),
                }
                log_serializer = AttendanceEntriesSerializer(data=log_data)
                if log_serializer.is_valid():
                    log_entry = log_serializer.save()
                    return Response(
                        {
                            "attendance": serializer.data,
                            
                            
                            "message": "Attendance record created successfully.",
                        },
                        status=status.HTTP_201_CREATED,
                    )
                else:
                    attendance.delete()  # Rollback attendance creation if log fails
                    return Response(log_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        else:
            # Attendance already exists for today, just create a new log entry
            last_log = AttendanceEntries.objects.filter(
                check_out__isnull=False
            ).order_by('-check_out').first()
            
            reason = "Live on"
            if last_log:
                if last_log.reason == "End of Shift":
                    reason = "Live on"
                elif last_log.reason == "Break":
                    reason = "End Break"
                elif last_log.reason == "Lunch":
                    reason = "End Lunch"
                elif last_log.reason in ["Meeting", "Others"]:
                    reason = "End Others"

            # Create initial log entry
            temp_object=Temp.objects.get(emp_id=temp_obj.emp_id)
            print(temp_object)
            logdetail=AttendanceLog.objects.create(
                emp_id=temp_object,
                attendance=attendance,
                time=timezone.now(),
                is_in=True,
                reason=reason,
                ip_address= get_client_ip(request)
            )
            log_serializer1 = AttendanceLogSerializer(logdetail)
            log_data = {
                'attendance': attendance.id,
                'emp_id': temp_obj.emp_id,
                'check_in': current_time,  # or use a different key like 'check_out' depending on your use case
                'ip_address': get_client_ip(request),
            }
            log_serializer = AttendanceEntriesSerializer(data=log_data)
            if log_serializer.is_valid():
                log_entry = log_serializer.save()
                return Response(
                    {
                        "message": "Entry added to existing attendance.",
                        "entries": log_serializer.data,
                        "log":log_serializer1.data,
                    },
                    status=status.HTTP_201_CREATED,
                )
            else:
                return Response(log_serializer.errors, status=status.HTTP_400_BAD_REQUEST)


    def patch(self, request, attendance_id=None):
        """
        Update attendance record to set check-out time and calculate total hours:
        - Now captures reason for check-out (break, lunch, end of shift, etc.)
        - Staff can only update for assigned interns
        - Admin can update any attendance record
        """
        user = request.user
        try:
            user_temp = Temp.objects.get(user=user)
        except Temp.DoesNotExist:
            return Response(
                {"error": "User not found."},
                status=status.HTTP_404_NOT_FOUND
            )
        
        role = user_temp.role.lower()
        
        # Get the attendance record to update
        try:
            # If attendance_id is provided, use it to find the attendance record
            if attendance_id:
                attendance = Attendance.objects.get(id=attendance_id)
            else:
                # Otherwise, try to find today's attendance for the current user or specified emp_id
                emp_id = request.query_params.get('emp_id') or user_temp.emp_id
                date = request.query_params.get('date') or timezone.now().date()
                try:
                    temp_obj = Temp.objects.get(emp_id=emp_id)
                    attendance = Attendance.objects.get(emp_id=temp_obj, date=date)
                except Temp.DoesNotExist:
                    return Response(
                        {"error": "Employee not found."},
                        status=status.HTTP_404_NOT_FOUND
                    )
        except Attendance.DoesNotExist:
            return Response(
                {"error": "No attendance record found."},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Check permissions
        if role == "intern":
            # Interns can only update their own attendance
            if str(attendance.emp_id.emp_id) != str(user_temp.emp_id):
                return Response(
                    {"error": "You can only update your own attendance."},
                    status=status.HTTP_403_FORBIDDEN
                )
        if role == 'staff':
            try:
                # Get staff's department
                staff_data = UserData.objects.get(user=user)
                if not staff_data.department:
                    return Response({"error": "Staff department not configured"}, status=status.HTTP_400_BAD_REQUEST)

                # Verify intern is in staff's department
                intern_data = UserData.objects.get(emp_id=attendance.emp_id)
                
                if intern_data.department != staff_data.department:
                    return Response(
                        {"error": "You can only update attendance for interns in your department"},
                        status=status.HTTP_403_FORBIDDEN
                    )

            except UserData.DoesNotExist:
                return Response({"error": "Department information not found"}, status=status.HTTP_404_NOT_FOUND)
        elif role not in ["admin", "hr", "intern","staff"]:
            # Unknown role
            return Response(
                {"error": "Invalid user role."},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Find the most recent open log entry
        open_log = AttendanceEntries.objects.filter(
            attendance=attendance,
            check_out__isnull=True
        ).order_by('-check_in').first()
        
        if not open_log:
            return Response(
                {"message": "This entry is over."},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Set the check-out time to the current time
        current_time = timezone.now()
        
        # Calculate total hours for this log entry
        if open_log.check_in:
            # Calculate the time difference for this log entry
            time_diff = current_time - open_log.check_in
            
            # Extract reason from request data
            reason = request.data.get('reason')
            
            # Update the log entry
            open_log.check_out = current_time
            open_log.total_hours = time_diff
            if reason:
                open_log.reason = reason
            open_log.save()

            
            reason = "Live on"
            in_out=True
            if open_log:
                if open_log.reason == "End of Shift":
                    in_out=False
                    reason = "Live Off"
                    
                elif open_log.reason == "Break":
                    reason = "Start Break"
                elif open_log.reason == "Lunch":
                    reason = "Start Lunch"
                elif open_log.reason in ["Meeting", "Others"]:
                    reason = "Start Others"

            # Create initial log entry
            print("Hello")
            print(open_log.emp_id.emp_id)
            temp_object=Temp.objects.get(emp_id=open_log.emp_id.emp_id)
            
            
            logdetail=AttendanceLog.objects.create(
                emp_id=temp_object,
                attendance=attendance,
                time=timezone.now(),
                is_in=in_out,
                reason=reason,
                ip_address = get_client_ip(request)
            )
            log_serializer1 = AttendanceLogSerializer(logdetail)
            
            # Now update the attendance record with accumulated hours
            all_logs = AttendanceEntries.objects.filter(attendance=attendance)
            total_hours = timedelta()
            for log in all_logs:
                if log.total_hours:
                    total_hours += log.total_hours
            
            # Update the attendance record
            attendance.check_out = current_time # Last checkout time
            attendance.total_hours = total_hours
            attendance.save()
            
            # Log the activity
            log_activity(
                user,
                "Updated",
                "AttendanceEntries",
                open_log.id,
                f"Updated attendance entries with check-out time. Reason: {open_log.get_reason_display() if open_log.reason else 'Not specified'}. Total hours: {time_diff}",
                request
            )
            
            log_activity(
                user,
                "Updated",
                "Attendance",
                attendance.id,
                f"Updated attendance record with accumulated hours. Total hours: {total_hours}",
                request
            )
            
            # Serialize and return the updated records
            attendance_serializer = AttendanceSerializer(attendance)
            log_serializer = AttendanceEntriesSerializer(open_log)
            response_data = {
                "attendance": attendance_serializer.data,
                
                "message": "Checkout completed successfully."
            }
            
            return Response(response_data)
        else:
            return Response(
                {"error": "Cannot update check-out time: No check-in time recorded in the log."},
                status=status.HTTP_400_BAD_REQUEST
            )


    def delete(self, request, attendance_id=None):
        """
        Delete an attendance record:
        - Admin can delete any attendance record
        - Staff can delete attendance records only for interns assigned to them
        """
        user = request.user
        try:
            user_temp = Temp.objects.get(user=user)
        except Temp.DoesNotExist:
            return Response(
                {"error": "User not found."},
                status=status.HTTP_404_NOT_FOUND
            )
        
        role = user_temp.role.lower()
        
        # Check if user has permission to delete attendance records
        if role not in ["admin", "staff"]:
            return Response(
                {"error": "Only admin and staff can delete attendance records."},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Get the attendance record to delete
        try:
            # If attendance_id is provided in the URL, use it
            if attendance_id:
                attendance = Attendance.objects.get(id=attendance_id)
            else:
                # Otherwise, get it from request data
                emp_id = request.query_params.get('emp_id') or request.data.get('emp_id')
                date = request.query_params.get('date') or request.data.get('date')
                
                if not emp_id or not date:
                    return Response(
                        {"error": "Either attendance_id or both emp_id and date are required to delete an attendance record."},
                        status=status.HTTP_400_BAD_REQUEST
                    )
                
                attendance = Attendance.objects.get(emp_id=emp_id, date=date)
        except Attendance.DoesNotExist:
            return Response(
                {"error": "Attendance record not found."},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # For staff, check if the attendance record belongs to an intern assigned to them
        if role == "staff":
            # Check if the attendance record's emp_id is for an intern assigned to this staff
            is_assigned = UserData.objects.filter(
                emp_id=attendance.emp_id,
                reporting_manager=user
            ).exists()
            
            if not is_assigned:
                return Response(
                    {"error": "You can only delete attendance records for interns assigned to you."},
                    status=status.HTTP_403_FORBIDDEN
                )
        
        # Store information for logging before deletion
        attendance_id = attendance.id
        emp_id = attendance.emp_id.emp_id if attendance.emp_id else "Unknown"
        date = attendance.date
        
        # Delete the attendance record
        attendance.delete()
        
        # Log the activity
        log_activity(
            user,
            "Deleted",
            "Attendance",
            attendance_id,
            f"Deleted attendance record for employee {emp_id} on {date}",
            request
        )
        
        return Response(
            {"message": "Attendance record deleted successfully."},
            status=status.HTTP_204_NO_CONTENT
        )











#-------------------------------------------Changes----------------------------------------------------#
from django.db import transaction
from django.db.models import OuterRef, Subquery, Q, Sum, F, Count, Case, When, Value, IntegerField
from django.utils import timezone
from datetime import datetime, timedelta, date
import logging
import json
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import AttendanceClaim
from .serializers import AttendanceClaimSerializer, DocumentVersionSerializer
from Sims.permissions import StaffUserDataAccessPermission
from Sims.utils.email_utils import  send_offer_letter_reportlab 

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.db import transaction
from django.db.models import OuterRef, Subquery, F
from rest_framework.permissions import IsAuthenticated
from Sims.models import Document, DocumentVersion, Temp, UserData
from Sims.serializers import DocumentSerializer, DocumentVersionSerializer
from Sims.permissions import StaffUserDataAccessPermission


class DocumentView(APIView):
    permission_classes = [IsAuthenticated, StaffUserDataAccessPermission]

    def get(self, request, emp_id=None):
        try:
            user_temp = Temp.objects.get(user=request.user, is_deleted=False)
            user_role = user_temp.role.lower()
            
            if emp_id:
                # Get a specific document
                try:
                    document = Document.objects.get(pk=emp_id, is_deleted=False)
                    
                    # Authorization checks
                    if user_role == "intern" and document.uploader != user_temp:
                        return Response(
                            {"error": "You can only view your own documents"},
                            status=status.HTTP_403_FORBIDDEN
                        )
                    
                    if user_role == "staff":
                        try:
                            staff_dept = UserData.objects.get(user=request.user).department
                            if document.uploader.role.lower() == "intern":
                                intern_data = UserData.objects.get(emp_id=document.uploader)
                                if intern_data.department != staff_dept:
                                    return Response(
                                        {"error": "You can only view documents from interns in your department"},
                                        status=status.HTTP_403_FORBIDDEN
                                    )
                        except UserData.DoesNotExist:
                            pass
                    
                    if user_role == "admin" or user_role == "hr":
                        pass
                    
                    serializer = DocumentSerializer(document, context={'request': request})
                    return Response(serializer.data)
                    
                except Document.DoesNotExist:
                    return Response(
                        {"error": "Document not found"},
                        status=status.HTTP_404_NOT_FOUND
                    )
            
            # List documents with appropriate filtering based on user role
            documents = Document.objects.filter(is_deleted=False)
            
            if user_role == "intern":
                # Interns can only see their own uploaded documents
                documents = documents.filter(uploader=user_temp)
            elif user_role == "staff":
                # Staff can see documents from interns in their department
                try:
                    staff_dept = UserData.objects.get(user=request.user).department
                    intern_temps = Temp.objects.filter(
                        userdata__department=staff_dept,
                        role__iexact='intern'
                    )
                    documents = documents.filter(
                        uploader__in=intern_temps
                    ) | documents.filter(uploader=user_temp)
                except UserData.DoesNotExist:
                    documents = documents.filter(uploader=user_temp)
            # Admin and HR can see all documents (no additional filtering needed)
            
            # Apply additional filters from query parameters
            status_filter = request.query_params.get('status')
            if status_filter:
                documents = documents.filter(status=status_filter.upper())
                
            uploader_filter = request.query_params.get('uploader')
            if uploader_filter:
                documents = documents.filter(uploader__emp_id=uploader_filter)
                
            receiver_filter = request.query_params.get('receiver')
            if receiver_filter:
                documents = documents.filter(receiver__emp_id=receiver_filter)
            
            # Order by creation date (newest first)
            documents = documents.order_by('-created_at')
            
            # Pagination
            page = request.query_params.get('page', 1)
            page_size = request.query_params.get('page_size', 10)
            paginator = Paginator(documents, page_size)
            
            try:
                paginated_documents = paginator.page(page)
            except PageNotAnInteger:
                paginated_documents = paginator.page(1)
            except EmptyPage:
                paginated_documents = paginator.page(paginator.num_pages)
                
            serializer = DocumentSerializer(
                paginated_documents,
                many=True,
                context={'request': request}
            )
            
            return Response({
                'count': paginator.count,
                'next': paginated_documents.next_page_number() if paginated_documents.has_next() else None,
                'previous': paginated_documents.previous_page_number() if paginated_documents.has_previous() else None,
                'results': serializer.data
            })
            
        except Exception as e:
            return Response(
                {"error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def post(self, request):
        try:
            uploader_temp = Temp.objects.get(user=request.user)
            user_role = uploader_temp.role.lower()
        except Temp.DoesNotExist:
            return Response({"error": "User profile not found"}, status=status.HTTP_403_FORBIDDEN)

        if user_role not in ['admin', 'hr', 'staff', 'intern']:
            return Response({"error": "Unauthorized role"}, status=status.HTTP_403_FORBIDDEN)

        files = request.FILES.getlist('files')
        if not files:
            return Response({"error": "At least one file is required"}, status=status.HTTP_400_BAD_REQUEST)

        request_data = {
            'title': request.data.get('title', ''),
            'description': request.data.get('description', ''),
            'receiver': request.data.get('receiver', None),
            'uploader': uploader_temp.emp_id,
            'generate_offer_letter': request.data.get('generate_offer_letter', 'false').lower() == 'true'
        }

        # Validate and set department for staff/intern
        if user_role in ['admin', 'hr']:
            if not request_data['receiver']:
                return Response({"error": "Receiver is required for admin/HR uploads"}, status=status.HTTP_400_BAD_REQUEST)
            try:
                Temp.objects.get(emp_id=request_data['receiver'])
            except Temp.DoesNotExist:
                return Response({"error": "Receiver not found"}, status=status.HTTP_404_NOT_FOUND)

        elif user_role == 'staff':
            try:
                staff_data = UserData.objects.get(user=request.user)
                department = staff_data.department
                if request_data['receiver']:
                    receiver_temp = Temp.objects.get(emp_id=request_data['receiver'])
                    receiver_data = UserData.objects.get(emp_id=receiver_temp)
                    if receiver_data.department != department:
                        return Response({"error": "Receiver must be in your department"}, status=status.HTTP_403_FORBIDDEN)
                request_data['department'] = department.department
            except UserData.DoesNotExist:
                return Response({"error": "Staff department not configured"}, status=status.HTTP_400_BAD_REQUEST)

        elif user_role == 'intern':
            request_data['receiver'] = None
            try:
                intern_data = UserData.objects.get(emp_id=uploader_temp)
                request_data['department'] = intern_data.department.department
            except UserData.DoesNotExist:
                return Response({"error": "Intern profile not configured"}, status=status.HTTP_400_BAD_REQUEST)

        uploaded_docs = []
        for file in files:
            request_data_with_file = request_data.copy()
            request_data_with_file['file'] = file

            # Duplicate restriction removed: allow multiple documents with same title, uploader, and receiver
            doc_serializer = DocumentSerializer(data=request_data_with_file, context={'request': request})
            if not doc_serializer.is_valid():
                return Response(doc_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            document = doc_serializer.save()

            version_data = {
                'document': document.id,
                'file': file,
                'uploaded_by': uploader_temp.emp_id,
                'changes': request.data.get('changes', 'Initial version')
            }
            version_serializer = DocumentVersionSerializer(data=version_data)
            if version_serializer.is_valid():
                version_serializer.save()
                uploaded_docs.append({
                    'document': doc_serializer.data,
                    'version': version_serializer.data
                })
            else:
                document.delete()
                return Response(version_serializer.errors, status=status.HTTP_400_BAD_REQUEST)


        return Response(uploaded_docs, status=status.HTTP_201_CREATED)






    def patch(self, request, pk=None):
        try:
            user_temp = Temp.objects.get(user=request.user, is_deleted=False)
            document = Document.objects.get(pk=pk, is_deleted=False)
            role = user_temp.role.lower()
            
            # Authorization checks
            if role == "intern" and document.uploader != user_temp:
                return Response({"error": "You can only modify your own documents"}, 
                             status=status.HTTP_403_FORBIDDEN)
            
            if role == "staff":
                staff_dept = UserData.objects.get(user=request.user).department
                if document.uploader.role.lower() != "intern" or \
                   UserData.objects.get(emp_id=document.uploader).department != staff_dept:
                    return Response({"error": "Can only modify documents from interns in your department"}, 
                                 status=status.HTTP_403_FORBIDDEN)

            # Status transition validation
            current_status = document.status
            new_status = request.data.get('status', current_status)
            
            if role == "intern" and new_status != current_status and new_status != "SUBMITTED":
                return Response({"error": "Interns can only submit documents"}, 
                             status=status.HTTP_403_FORBIDDEN)
            
            # File update handling
            if 'file' in request.data and document.status == "APPROVED":
                return Response({"error": "Cannot modify approved documents"}, 
                             status=status.HTTP_403_FORBIDDEN)

            serializer = DocumentSerializer(document, data=request.data, partial=True)
            if serializer.is_valid():
                updated_doc = serializer.save()
                
                # Create new version if file changed
                if 'file' in request.data:
                    DocumentVersion.objects.create(
                        document=updated_doc,
                        file=updated_doc.file,
                        version_number=document.versions.count() + 1,
                        uploaded_by=user_temp
                    )
                
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        except Document.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete(self, request, pk=None):
        try:
            user_temp = Temp.objects.get(user=request.user, is_deleted=False)
            document = Document.objects.get(pk=pk)
            role = user_temp.role.lower()

            # Authorization checks
            if role == "intern" and document.uploader != user_temp:
                return Response({"error": "You can only delete your own documents"}, 
                             status=status.HTTP_403_FORBIDDEN)
            
            if role == "staff":
                staff_dept = UserData.objects.get(user=request.user).department
                if document.uploader.role.lower() != "intern" or \
                   UserData.objects.get(emp_id=document.uploader).department != staff_dept:
                    return Response({"error": "Can only delete documents from interns in your department"}, 
                                 status=status.HTTP_403_FORBIDDEN)

            # Soft deletion
            document.is_deleted = True
            document.status = "ARCHIVED"
            document.save()
            
            # Log deletion
            log_activity(
                request.user,
                "Deleted", 
                "Document",
                document.id,
                f"Soft-deleted document {document.declaration_number}",
                request
            )
            
            return Response({"message": "Document archived successfully"}, 
                          status=status.HTTP_204_NO_CONTENT)

        except Document.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class DocumentByEmpView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, emp_id=None):
        try:
            # Get the current user's temp profile
            try:
                user_temp = Temp.objects.get(user=request.user, is_deleted=False)
                user_role = user_temp.role.lower()
            except Temp.DoesNotExist:
                return Response(
                    {"error": "User profile not found"},
                    status=status.HTTP_403_FORBIDDEN
                )

            # Get the target employee
            try:
                target_emp = Temp.objects.get(emp_id=emp_id, is_deleted=False)
            except Temp.DoesNotExist:
                return Response(
                    {"error": "Employee not found"},
                    status=status.HTTP_404_NOT_FOUND
                )

            # Authorization checks
            if user_role == "intern" and user_temp.emp_id != emp_id:
                return Response(
                    {"error": "You can only view your own documents"},
                    status=status.HTTP_403_FORBIDDEN
                )

            if user_role == "staff":
                try:
                    staff_dept = UserData.objects.get(user=request.user).department
                    if target_emp.role.lower() == "intern":
                        intern_data = UserData.objects.get(emp_id=target_emp)
                        if intern_data.department != staff_dept:
                            return Response(
                                {"error": "You can only view documents from interns in your department"},
                                status=status.HTTP_403_FORBIDDEN
                            )
                except UserData.DoesNotExist:
                    # If staff has no department, they can only see their own docs
                    if user_temp.emp_id != emp_id:
                        return Response(
                            {"error": "You can only view your own documents"},
                            status=status.HTTP_403_FORBIDDEN
                        )

            # Get documents for the employee
            documents = Document.objects.filter(
                receiver=target_emp,
                is_deleted=False
            ).order_by('-created_at')

            # Apply status filter if provided
            status_filter = request.query_params.get('status')
            if status_filter:
                documents = documents.filter(status=status_filter.upper())

            # Pagination
            page = request.query_params.get('page', 1)
            page_size = request.query_params.get('page_size', 10)
            paginator = Paginator(documents, page_size)

            try:
                paginated_documents = paginator.page(page)
            except PageNotAnInteger:
                paginated_documents = paginator.page(1)
            except EmptyPage:
                paginated_documents = paginator.page(paginator.num_pages)

            serializer = DocumentSerializer(
                paginated_documents,
                many=True,
                context={'request': request}
            )

            return Response({
                'count': paginator.count,
                'next': paginated_documents.next_page_number() if paginated_documents.has_next() else None,
                'previous': paginated_documents.previous_page_number() if paginated_documents.has_previous() else None,
                'results': serializer.data
            })

        except Exception as e:
            return Response(
                {"error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    def patch(self, request, emp_id):
        try:
            document_type = request.data.get('document_type')
            file = request.FILES.get('file')
            
            if not document_type or not file:
                return Response(
                    {"error": "Both document_type and file are required"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Get the Temp instance for the employee
            try:
                temp_user = Temp.objects.get(emp_id=emp_id)
                uploader_temp = Temp.objects.get(user=request.user)
            except Temp.DoesNotExist:
                return Response(
                    {"error": "Employee not found"},
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Update or create the document
            document, created = Document.objects.update_or_create(
                receiver=temp_user,
                title=document_type,
                defaults={
                    'file': file,
                    'uploader': uploader_temp,
                    'status': 'SUBMITTED'  # Or whatever default status you want
                }
            )
            
            file_url = request.build_absolute_uri(document.file.url) if document.file else None
            
            return Response({
                "id": str(document.id),
                "emp_id": temp_user.emp_id,
                "title": document.title,
                "file": file_url,
                "uploaded_at": document.created_at,
                "uploaded_by": document.uploader.user.username if document.uploader else None
            }, status=status.HTTP_200_OK if not created else status.HTTP_201_CREATED)
            
        except Exception as e:
            return Response(
                {"error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_offer_letter_api(request):
    """API endpoint to generate and send an offer letter using Word template."""
    logger = logging.getLogger(__name__)
    error_details = {
        'step': 'Initialization',
        'error_type': None,
        'message': None
    }

    try:
        # 1. User Authentication and Permission Check
        error_details['step'] = 'Authentication'
        try:
            user_temp = Temp.objects.get(user=request.user)
            if user_temp.role.lower() not in ["admin", "hr", "staff"]:
                error_msg = "Only staff/admin/hr can generate offer letters"
                logger.warning(f"Permission denied for user: {request.user.username}")
                return Response(
                    {"success": False, "error": error_msg}, 
                    status=status.HTTP_403_FORBIDDEN
                )
        except Temp.DoesNotExist:
            error_msg = "User profile not found. Please complete your profile."
            logger.error(error_msg)
            return Response(
                {"success": False, "error": error_msg}, 
                status=status.HTTP_403_FORBIDDEN
            )

        # 2. Validate Request Data
        error_details['step'] = 'Request Validation'
        data = request.data
        intern_emp_id = data.get("emp_id")
        
        if not intern_emp_id:
            error_msg = "Intern emp_id is required"
            logger.error(error_msg)
            return Response(
                {"success": False, "error": error_msg}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        # 3. Get Intern Details
        error_details['step'] = 'Fetch Intern Details'
        try:
            intern_temp = Temp.objects.get(emp_id=intern_emp_id, role__iexact="intern")
            logger.info(f"Found intern: {intern_temp.user.get_full_name()} ({intern_temp.emp_id})")
        except Temp.DoesNotExist:
            error_msg = f"Intern with emp_id {intern_emp_id} not found"
            logger.error(error_msg)
            return Response(
                {"success": False, "error": "Intern not found"}, 
                status=status.HTTP_404_NOT_FOUND
            )

        # 4. Prepare Data for Offer Letter
        error_details['step'] = 'Prepare Offer Letter Data'
        try:
            # Get data from request with fallbacks
            offer_data = {
                'college_name': data.get("college_name", "").strip(),
                'start_date': data.get("start_date", ""),
                'end_date': data.get("end_date", ""),
                'position_title': data.get("position_title", "FullStack Intern").strip(),
                'domain': data.get("domain", "VDart Academy").strip(),
                'shift_time': data.get("shift_time", "9:00 AM to 1:00 PM IST").strip(),
                'shift_days': data.get("shift_days", "Monday to Friday").strip(),
                'work_location': data.get(
                    "work_location", 
                    "VDart, Global Capability Center, Mannarpuram"
                ).strip(),
                'reporting_to': data.get("reporting_to", "Derrick Alex").strip()
            }

            # Log the request details
            logger.info(f"Generating offer letter with data: {json.dumps(offer_data, indent=2)}")

            # 5. Generate and Send Offer Letter using Word template
            error_details['step'] = 'Generate and Send Offer Letter'
            
            # Import the send_offer_letter function from email_utils
            from .utils.email_utils import send_offer_letter
            
            success, message, error_info = send_offer_letter(
                user=intern_temp.user,
                emp_id=intern_temp.emp_id,
                **offer_data
            )

            if success:
                logger.info(f"Successfully sent offer letter to {intern_temp.user.email}")
                return Response(
                    {"success": True, "message": message},
                    status=status.HTTP_200_OK
                )
            else:
                error_msg = f"Failed to send offer letter: {message}"
                logger.error(error_msg)
                error_details.update({
                    'error_type': 'Offer Letter Error',
                    'message': message,
                    'details': error_info
                })
                return Response(
                    {"success": False, "error": message, "details": error_info},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

        except Exception as e:
            error_msg = f"Error processing offer letter data: {str(e)}"
            logger.error(error_msg, exc_info=True)
            error_details.update({
                'error_type': 'Data Processing Error',
                'message': str(e),
                'traceback': traceback.format_exc()
            })
            return Response(
                {"success": False, "error": "Error processing request", "details": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    except Exception as e:
        error_msg = f"Unexpected error in generate_offer_letter_api: {str(e)}"
        logger.error(error_msg, exc_info=True)
        error_details.update({
            'error_type': 'Unexpected Error',
            'message': str(e),
            'traceback': traceback.format_exc()
        })
        return Response(
            {"success": False, "error": "An unexpected error occurred", "details": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

# class LeaveHistoryView(APIView):
#     permission_classes = [IsAuthenticated]
# 
#     def get(self, request, emp_id=None):
# class LeaveHistoryView(APIView):
#     permission_classes = [IsAuthenticated]

#     def get(self, request, emp_id=None):
#         user = request.user

#         try:
#             user_temp = Temp.objects.get(user=user)
#             role = user_temp.role.lower()
#         except Temp.DoesNotExist:
#             return Response({"error": "User profile not found"}, status=404)

#         # For specific emp_id requests
#         if emp_id:
#             try:
#                 target_temp = Temp.objects.get(emp_id=emp_id)
#                 # Staff can only see their interns
#                 if role == 'staff':
#                     user_data = UserData.objects.get(user=target_temp.user)
#                     if user not in [user_data.reporting_manager, user_data.reporting_supervisor]:
#                         return Response({"error": "Not authorized"}, status=403)
#                 # Interns can only see their own
#                 elif role == 'intern' and user_temp.emp_id != emp_id:
#                     return Response({"error": "Not authorized"}, status=403)
#                 leave_requests = LeaveRequest.objects.filter(user=target_temp.user, status="APPROVED")
#             except Temp.DoesNotExist:
#                 return Response({"error": "Employee not found"}, status=404)
#         else:
#             if role == 'admin':
#                 leave_requests = LeaveRequest.objects.filter(status="APPROVED")
#             elif role == 'staff':
#                 interns = UserData.objects.filter(
#                     Q(reporting_manager=user) | Q(reporting_supervisor=user)
#                 ).values_list('user', flat=True)
#                 leave_requests = LeaveRequest.objects.filter(user__in=interns, status="APPROVED")
#             elif role == 'intern':
#                 leave_requests = LeaveRequest.objects.filter(user=user, status="APPROVED")
#             else:
#                 return Response({"error": "Invalid role"}, status=400)

#         # Format response
#         formatted_leaves = []
#         for leave in leave_requests.order_by('-from_date'):
#             formatted_leaves.append({
#                 "type": leave.leave_type,
#                 "dates": f"{leave.from_date.strftime('%d/%m/%Y')} - {leave.to_date.strftime('%d/%m/%Y')}",
#                 "status": leave.status,
#                 "reason": leave.request_reason
#             })
#         return Response(formatted_leaves)


class LeaveHistoryView(APIView):
    permission_classes = [IsAuthenticated, StaffAttendanceAccessPermission]
    def get(self, request):
        # Get staff's department
        user_temp = Temp.objects.get(user=request.user,is_deleted=False)
        role = user_temp.role.lower()
        if role == 'staff':
        
            try:
                staff_data = UserData.objects.get(user=request.user)
                department = staff_data.department
            except UserData.DoesNotExist:
                return Response({"error": "User data not found"}, status=404)
            # Get all interns in the same department
            interns_in_dept = UserData.objects.filter(
                department=department,
                emp_id__role='intern'
            ).values_list('user', flat=True)
            # Get leave requests for these interns
            leave_requests = LeaveRequest.objects.filter(
                user__in=interns_in_dept
            ).order_by('-created_at')
        elif role == 'intern':
            leave_requests=LeaveRequest.objects.filter(user=request.user).order_by('-created_at')
        elif role == 'admin':
            leave_requests= LeaveRequest.objects.all().order_by('-created_at')
        serializer = LeaveRequestSerializer(leave_requests, many=True)
        return Response(serializer.data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_leave_balance(request):
    user = request.user
    total = 5 # Default total leaves
    # Sum all approved leaves
    approved_leaves = LeaveRequest.objects.filter(user=user, status="APPROVED")
    leaves_taken = approved_leaves.aggregate(Sum('number_of_days'))['number_of_days__sum'] or 0
    used = leaves_taken
    remaining = max(total - used, 0)
    return Response({
        "used": used,
        "remaining": remaining,
        "total": total,
        "leaves_taken": leaves_taken
    })


class DeletedUsersView(APIView):
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        """Get all soft-deleted users (admin only)"""
        try:
            user_temp = Temp.objects.get(user=request.user)
            if user_temp.role.lower() not in ['admin', 'hr']:
                return Response({"error": "Only admin/HR can view deleted users"}, 
                              status=status.HTTP_403_FORBIDDEN)
            
            # Get all soft-deleted Temp records
            deleted_temps = Temp.objects.filter(is_deleted=True)
            
            deleted_users_data = []
            for temp in deleted_temps:
                try:
                    user_data = UserData.objects.get(emp_id=temp, is_deleted=True)
                    deleted_users_data.append({
                        'emp_id': temp.emp_id,
                        'username': temp.user.username,
                        'email': temp.user.email,
                        'role': temp.role,
                        'department': user_data.department.department if user_data.department else 'N/A',
                        'deleted_date': temp.updated_date,
                        'domain': user_data.domain.domain if user_data.domain else 'N/A',
                        'scheme': user_data.scheme or 'N/A'
                    })
                except UserData.DoesNotExist:
                    deleted_users_data.append({
                        'emp_id': temp.emp_id,
                        'username': temp.user.username,
                        'email': temp.user.email,
                        'role': temp.role,
                        'department': 'N/A',
                        'deleted_date': temp.updated_date,
                        'domain': 'N/A',
                        'scheme': 'N/A'
                    })
            
            return Response(deleted_users_data, status=status.HTTP_200_OK)
            
        except Temp.DoesNotExist:
            return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)





#-------------------------------------------------changes new 24-4 ------------------------------------#
class EmpEmailLookupView(APIView):
    permission_classes = [AllowAny]  # Allow unrestricted access

    def get(self, request, emp_id):
        try:
            temp = Temp.objects.get(emp_id=emp_id, is_deleted=False)
            user = temp.user
            return Response({"emp_id": emp_id, "email": user.email})
        except Temp.DoesNotExist:
            return Response({"error": "Employee not found."}, status=404)




class AssignedTaskHistoryView(APIView):
    permission_classes = [IsAuthenticated]
    def get(self, request):
        try:
            temp = Temp.objects.get(user=request.user, is_deleted=False)
        except Temp.DoesNotExist:
            return Response({"error": "User role not found"}, status=403)
        if temp.role.lower() != "intern":
            return Response({"error": "Only interns can view this history"}, status=403)
        tasks = Task.objects.filter(assigned_by=request.user, is_deleted=False)
        serializer = AssignedTaskHistorySerializer(tasks, many=True)
        return Response(serializer.data)
    def delete(self, request, pk):
        try:
            temp = Temp.objects.get(user=request.user, is_deleted=False)
        except Temp.DoesNotExist:
            return Response({"error": "User role not found"}, status=403)
        if temp.role.lower() != "intern":
            return Response({"error": "Only interns can delete assigned tasks"}, status=403)
        task = get_object_or_404(Task, pk=pk, is_deleted=False)
        if task.assigned_by != request.user:
            return Response({"error": "You can only delete tasks you assigned"}, status=403)
        task.delete()
        return Response({"message": "Task deleted successfully"}, status=204)
    







class StaffListByDepartmentView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Only allow admin users
        temp = Temp.objects.get(user=request.user)
        if temp.role.lower() != 'admin':
            return Response({"error": "Only admins can access this endpoint."}, status=403)
        
        department_name = request.query_params.get('department')
        if not department_name:
            return Response({"error": "Department parameter is required."}, status=400)
        
        try:
            department = Department.objects.get(department__iexact=department_name, is_deleted=False)
        except Department.DoesNotExist:
            return Response({"error": "Department not found."}, status=404)
        
        # Get all staff in the department
        staff_userdatas = UserData.objects.filter(department=department, emp_id__role='staff', is_deleted=False)
        staff_emps = Temp.objects.filter(emp_id__in=staff_userdatas.values_list('emp_id', flat=True), is_deleted=False)
        serializer = TempSerializer(staff_emps, many=True)
        return Response(serializer.data)
    



class DomainByDepartmentView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        temp = Temp.objects.get(user=request.user)
        if temp.role.lower() != 'admin':
            return Response({"error": "Only admins can access this endpoint."}, status=403)
        
        department_name = request.query_params.get('department')
        if not department_name:
            return Response({"error": "Department parameter is required."}, status=400)
        
        try:
            department = Department.objects.get(department__iexact=department_name, is_deleted=False)
        except Department.DoesNotExist:
            return Response({"error": "Department not found."}, status=404)
        
        domains = Domain.objects.filter(department=department, is_deleted=False)
        serializer = DomainSerializer(domains, many=True)
        return Response(serializer.data)





from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404


from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404
def download_document(request, pk):
    document = get_object_or_404(Document, pk=pk, is_deleted=False)
    if not document.file:
        raise Http404("File not found.")
    original_filename = document.file.name.split('/')[-1]
    response = FileResponse(document.file.open('rb'), as_attachment=True, filename=original_filename)
    return response




@api_view(["GET"])
@permission_classes([IsAuthenticated])
def interns_without_asset_in_department(request, department_name):
    try:
        # Get all active UserData for the given department where asset_code is NULL
        interns = UserData.objects.filter(
            department__department__iexact=department_name.strip(),
            emp_id__role="intern",
            asset_code__isnull=True,
            is_deleted=False
        ).select_related('user', 'emp_id')
        data = []
        for intern in interns:
            data.append({
                "id": intern.user.id,
                "emp_id": intern.emp_id.emp_id,
                "username": intern.user.username,
            })
        return Response(data, status=200)
    except Exception as e:
        return Response({"error": str(e)}, status=400)




from datetime import datetime, date

class AttendanceDateRangeView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Extract dates from JSON body
        start_date_str = request.data.get('start_date')
        end_date_str = request.data.get('end_date')

        # Use today's date if not provided
        today = date.today()
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date() if start_date_str else today
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date() if end_date_str else start_date
        except ValueError:
            return Response({"error": "Invalid date format. Use YYYY-MM-DD."}, status=400)

        if start_date > end_date:
            return Response({"error": "start_date cannot be after end_date"}, status=400)

        # Get current user's Temp (employee) object
        try:
            user_temp = Temp.objects.get(user=request.user, is_deleted=False)
        except Temp.DoesNotExist:
            return Response({"error": "User profile not found or deactivated"}, status=404)

        role = user_temp.role.lower()
        queryset = Attendance.objects.filter(date__range=[start_date, end_date], is_deleted=False)

        # Role-based filtering
        if role in ['admin', 'hr']:
            pass  # All records
        elif role == 'staff':
            try:
                staff_data = UserData.objects.get(user=request.user, is_deleted=False)
                interns_in_dept = UserData.objects.filter(
                    department=staff_data.department,
                    emp_id__role='intern',
                    is_deleted=False
                ).values_list('emp_id', flat=True)
                queryset = queryset.filter(emp_id__in=interns_in_dept)
            except UserData.DoesNotExist:
                return Response({"error": "Staff department not configured"}, status=400)
        elif role == 'intern':
            queryset = queryset.filter(emp_id=user_temp)
        else:
            return Response({"error": "Unauthorized role"}, status=403)

        serializer = AttendanceSerializer(queryset, many=True)
        return Response(serializer.data, status=200)
    



class UserDetailView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, emp_id=None):
        # Get requester's role
        requester_temp = Temp.objects.get(user=request.user)
        
        # Admin access
        if requester_temp.role.lower() == 'admin':
            if emp_id:
                users = User.objects.filter(profile__temp__emp_id=emp_id)
            else:
                users = User.objects.all()
        
        # Staff access
        elif requester_temp.role.lower() == 'staff':
            staff_dept = UserData.objects.get(user=request.user).department
            if emp_id:
                users = User.objects.filter(
                    profile__temp__emp_id=emp_id,
                    userdata__department=staff_dept
                )
            else:
                users = User.objects.filter(userdata__department=staff_dept)
        
        # Intern access
        else:
            users = User.objects.filter(id=request.user.id)

        serializer = UserDetailSerializer(users, many=True)
        return Response(serializer.data)





class AttendanceAnalysisView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        def format_timedelta(td):
            total_seconds = int(td.total_seconds()) if td else 0
            hours, remainder = divmod(abs(total_seconds), 3600)
            minutes, seconds = divmod(remainder, 60)
            return f"{hours:02d}:{minutes:02d}:{seconds:02d}"

        # Get current user's emp_id
        try:
            temp = Temp.objects.get(user=request.user)
            emp_id = temp.emp_id
        except Temp.DoesNotExist:
            return Response({"error": "Employee record not found"}, status=status.HTTP_404_NOT_FOUND)

        # Date handling with timezone
        date_str = request.query_params.get('date')
        try:
            if date_str:
                naive_date = datetime.strptime(date_str, '%Y-%m-%d').date()
            else:
                naive_date = timezone.localdate()
            user_tz = timezone.get_current_timezone()
            start_date = datetime.combine(naive_date, datetime.min.time(), tzinfo=user_tz)
            end_date = datetime.combine(naive_date, datetime.max.time(), tzinfo=user_tz)
        except ValueError:
            return Response({"error": "Invalid date format. Use YYYY-MM-DD"}, status=status.HTTP_400_BAD_REQUEST)

        # Get logs within date range
        logs = AttendanceLog.objects.filter(
            emp_id__emp_id=emp_id,
            time__range=(start_date, end_date)
        ).order_by('time')

        # Initialize trackers
        break_time = timedelta()
        lunch_time = timedelta()
        other_time = timedelta()
        total_current_time = timedelta()
        
        current_live_on = None
        current_break_start = None
        current_lunch_start = None
        current_other_start = None

        is_break_on = False
        is_lunch_on = False
        is_others_on = False
        latest_log_status = False
        latest_live_action = None

        if logs.exists():
            latest_log = logs.last()
            now = timezone.now()

            for log in logs:
                reason = (log.reason or "").upper().replace(" ", "_")

                # Live time calculation
                if reason == "LIVE_ON":
                    current_live_on = log.time
                    latest_live_action = 'LIVE_ON'
                elif reason == "LIVE_OFF":
                    if current_live_on:
                        total_current_time += log.time - current_live_on
                        current_live_on = None
                    latest_live_action = 'LIVE_OFF'

                # Break time calculation
                if reason == "START_BREAK":
                    current_break_start = log.time
                elif reason == "END_BREAK" and current_break_start:
                    break_time += log.time - current_break_start
                    current_break_start = None

                # Lunch time calculation
                if reason == "START_LUNCH":
                    current_lunch_start = log.time
                elif reason == "END_LUNCH" and current_lunch_start:
                    lunch_time += log.time - current_lunch_start
                    current_lunch_start = None

                # Other time calculation
                if reason == "START_OTHERS":
                    current_other_start = log.time
                elif reason in ("END_OTHERS", "END_OTHER") and current_other_start:
                    other_time += log.time - current_other_start
                    current_other_start = None

            # Handle ongoing sessions
            if current_live_on:
                total_current_time += now - current_live_on
            if current_break_start:
                break_time += now - current_break_start
                is_break_on = True
            if current_lunch_start:
                lunch_time += now - current_lunch_start
                is_lunch_on = True
            if current_other_start:
                other_time += now - current_other_start
                is_others_on = True

        # Determine current status
        if latest_live_action == 'LIVE_ON':
            latest_log_status = True
        elif latest_live_action == 'LIVE_OFF':
            latest_log_status = False
        elif logs.exists():
            latest_log_status = latest_log.is_in

        return Response({
            'date': naive_date.isoformat(),
            'break_time': format_timedelta(break_time),
            'lunch_time': format_timedelta(lunch_time),
            'other_time': format_timedelta(other_time),
            'total_current_time': format_timedelta(total_current_time),
            'is_currently_in': latest_log_status,
            'is_break_on': is_break_on,
            'is_lunch_on': is_lunch_on,
            'is_others_on': is_others_on,
            'logs_count': logs.count()
        }, status=status.HTTP_200_OK)




class SimpleAttendanceTimesView(APIView):
    permission_classes = [IsAuthenticated]


    def get(self, request):
        # Get current user's emp_id
        try:
            temp = Temp.objects.get(user=request.user, is_deleted=False)
            emp_id = temp.emp_id
        except Temp.DoesNotExist:
            return Response({"error": "Employee record not found"}, status=404)

        # Today's date range
        user_tz = timezone.get_current_timezone()
        today = timezone.localdate()
        start_date = datetime.combine(today, datetime.min.time(), tzinfo=user_tz)
        end_date = datetime.combine(today, datetime.max.time(), tzinfo=user_tz)

        # Get all logs for today
        logs = AttendanceLog.objects.filter(
            emp_id__emp_id=emp_id,
            time__range=(start_date, end_date)
        ).order_by('time')

        break_time = timedelta()
        lunch_time = timedelta()
        other_time = timedelta()
        total_current_time = timedelta()

        current_live_on = None
        current_break_start = None
        current_lunch_start = None
        current_other_start = None

        now = timezone.now()

        for log in logs:
            reason = (log.reason or "").upper().replace(" ", "_")
            # Live time
            if reason == "LIVE_ON":
                current_live_on = log.time
            elif reason == "LIVE_OFF" and current_live_on:
                total_current_time += log.time - current_live_on
                current_live_on = None
            # Break
            if reason == "START_BREAK":
                current_break_start = log.time
            elif reason == "END_BREAK" and current_break_start:
                break_time += log.time - current_break_start
                current_break_start = None
            # Lunch
            if reason == "START_LUNCH":
                current_lunch_start = log.time
            elif reason == "END_LUNCH" and current_lunch_start:
                lunch_time += log.time - current_lunch_start
                current_lunch_start = None
            # Others
            if reason == "START_OTHERS":
                current_other_start = log.time
            elif reason in ("END_OTHERS", "END_OTHER") and current_other_start:
                other_time += log.time - current_other_start
                current_other_start = None

        # Add ongoing sessions up to now
        if current_live_on:
            total_current_time += now - current_live_on
        if current_break_start:
            break_time += now - current_break_start
        if current_lunch_start:
            lunch_time += now - current_lunch_start
        if current_other_start:
            other_time += now - current_other_start

        def format_td(td):
            total_seconds = int(td.total_seconds())
            hours, remainder = divmod(abs(total_seconds), 3600)
            minutes, seconds = divmod(remainder, 60)
            return f"{hours:02d}:{minutes:02d}:{seconds:02d}"

        return Response({
            "date": today.isoformat(),
            "total_current_time": format_td(total_current_time),
            "break_time": format_td(break_time),
            "lunch_time": format_td(lunch_time),
            "other_time": format_td(other_time),
        })
    





class AssertAssignmentLogView(APIView):
    permission_classes = [IsAuthenticated, StaffAssertAccessPermission]

    def get(self, request):
        user = request.user
        try:
            user_temp = Temp.objects.get(user=user)
            
            if user_temp.role == 'admin':
                logs = AssertAssignmentLog.objects.filter(is_deleted=False)
            elif user_temp.role == 'staff':
                staff_dept = UserData.objects.get(user=user).department
                assets_in_dept = AssertStock.objects.filter(department=staff_dept)
                logs = AssertAssignmentLog.objects.filter(asset__in=assets_in_dept, is_deleted=False)
            else:  # Interns
                logs = AssertAssignmentLog.objects.filter(emp=user_temp, is_deleted=False)
            
            serializer = AssertAssignmentLogSerializer(logs, many=True)
            return Response(serializer.data)
        
        except Temp.DoesNotExist:
            return Response({"error": "User profile not found"}, status=404)



class AssertAllTimeHistoryView(APIView):
    permission_classes = [IsAuthenticated, StaffAssertAccessPermission]

    def get(self, request):
        user = request.user
        assert_id=request.data.get('assert_id')
        
        try:
            user_temp = Temp.objects.get(user=user)
            logs = AssertAssignmentLog.objects.filter(is_deleted=False)
            
            # Staff can only see their department's assets
            if user_temp.role == 'staff':
                staff_dept = UserData.objects.get(user=user).department
                logs = logs.filter(asset__department=staff_dept)
            
            # Filter by specific asset if provided
            if assert_id:
                asset = get_object_or_404(AssertStock, assert_id=assert_id)
                
                # Additional permission check for staff
                if user_temp.role == 'staff' and asset.department != staff_dept:
                    return Response(
                        {"error": "Not authorized to view this asset's history"},
                        status=status.HTTP_403_FORBIDDEN
                    )
                
                logs = logs.filter(asset=asset)
            
            serializer = AssertAssignmentLogSerializer(logs.order_by('-timestamp'), many=True)
            return Response(serializer.data)
            
        except Temp.DoesNotExist:
            return Response({"error": "User profile not found"}, status=status.HTTP_404_NOT_FOUND)
        except UserData.DoesNotExist:
            return Response({"error": "Staff department not configured"}, status=status.HTTP_400_BAD_REQUEST)


class UserAssertAllTimeHistoryView(APIView):
    permission_classes = [IsAuthenticated, StaffAssertAccessPermission]
    def get(self, request):
        user = request.user
        try:
            user_temp = Temp.objects.get(user=user)
            
            # Get the target emp_id from query params
            emp_id = request.data.get('emp_id')
            if not emp_id:
                return Response({"error": "emp_id parameter is required"}, status=400)
                
            target_temp = get_object_or_404(Temp, emp_id=emp_id)
            # Filter logs for specific events related to the user
            logs = AssertAssignmentLog.objects.filter(
                Q(event_type='ASSIGNED') | 
                Q(event_type='RETURNED') |
                Q(event_type='REPAIR') |
                Q(event_type='REPAIR_RETURNED'),
                emp=target_temp,
                is_deleted=False
            )
            # Staff can only see their department's assets
            if user_temp.role == 'staff':
                staff_dept = UserData.objects.get(user=user).department
                logs = logs.filter(asset__department=staff_dept)
            serializer = AssertAssignmentLogSerializer(logs.order_by('-timestamp'), many=True)
            return Response(serializer.data)
        except Temp.DoesNotExist:
            return Response({"error": "User profile not found"}, status=404)
        except UserData.DoesNotExist:
            return Response({"error": "Staff department not configured"}, status=400)







# In views.py
class DeletedUsersView(APIView):
    permission_classes = [IsAuthenticated, StaffUserDataAccessPermission]
    
    def get(self, request):
        # Get all soft-deleted users through Temp model
        deleted_users = Temp.objects.filter(is_deleted=True)
        
        response_data = []
        
        for temp in deleted_users:
            user_data = {
                "emp_id": temp.emp_id,
                "username": temp.user.username,
                "role": temp.role,
                "college_details": None,
                "personal_data": None,
                "user_data": None
            }
            
            # Get college details if exists
            try:
                college_details = CollegeDetails.objects.get(emp_id=temp)
                user_data["college_details"] = CollegeDetailsSerializer(college_details).data
            except CollegeDetails.DoesNotExist:
                pass
            
            # Get personal data if exists
            try:
                personal_data = PersonalData.objects.get(emp_id=temp)
                user_data["personal_data"] = PersonalDataSerializer(personal_data).data
            except PersonalData.DoesNotExist:
                pass
            
            # Get user data if exists
            try:
                user_data_record = UserData.objects.get(emp_id=temp)
                user_data["user_data"] = UserDataSerializer(user_data_record).data
            except UserData.DoesNotExist:
                pass
            
            response_data.append(user_data)
        
        return Response(response_data, status=status.HTTP_200_OK)


from django.http import HttpResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from datetime import datetime
from django.contrib.auth import get_user_model
from django.core.mail import EmailMessage
from io import BytesIO



User = get_user_model()

class AttendanceClaimView(APIView):
    """
    API View for handling attendance claims.
    - GET: List all claims (filtered by user role)
    - POST: Create a new attendance claim
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, claim_id=None):
        user = request.user
        
        # Check user role and filter claims accordingly
        if user.is_staff or user.is_superuser:
            # Staff/Admin can see all claims
            claims = AttendanceClaim.objects.all()
        else:
            # Regular users can only see their own claims
            claims = AttendanceClaim.objects.filter(user=user)
        
        # Filter by status if provided
        status_param = request.query_params.get('status')
        if status_param:
            claims = claims.filter(status=status_param.lower())
        
        # Filter by date range if provided
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        if start_date and end_date:
            claims = claims.filter(
                date__range=[start_date, end_date]
            )
        
        # Get specific claim if ID is provided
        if claim_id:
            claims = claims.filter(id=claim_id)
            if not claims.exists():
                return Response(
                    {"detail": "Claim not found or access denied"}, 
                    status=status.HTTP_404_NOT_FOUND
                )
        
        # Pagination
        page = request.query_params.get('page', 1)
        page_size = request.query_params.get('page_size', 10)
        paginator = Paginator(claims, page_size)
        
        try:
            claims_page = paginator.page(page)
        except (PageNotAnInteger, EmptyPage):
            claims_page = paginator.page(1)
        
        serializer = AttendanceClaimSerializer(claims_page, many=True)
        
        return Response({
            'count': paginator.count,
            'total_pages': paginator.num_pages,
            'current_page': claims_page.number,
            'results': serializer.data
        })
    
    def post(self, request):
        """
        Create a new attendance claim.
        Required fields: attendance_id, claim_date, reason
        """
        user = request.user
        data = request.data.copy()
        data['user'] = user.id
        
        # Validate attendance exists and belongs to user
        attendance_id = data.get('attendance')
        try:
            attendance = Attendance.objects.get(id=attendance_id)
            if attendance.user != user and not (user.is_staff or user.is_superuser):
                return Response(
                    {"detail": "You can only create claims for your own attendance"}, 
                    status=status.HTTP_403_FORBIDDEN
                )
        except Attendance.DoesNotExist:
            return Response(
                {"detail": "Attendance record not found"}, 
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Check for duplicate pending claims
        existing_claim = AttendanceClaim.objects.filter(
            attendance=attendance,
            status='pending'
        ).exists()
        
        if existing_claim:
            return Response(
                {"detail": "A pending claim already exists for this attendance"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        serializer = AttendanceClaimSerializer(data=data)
        if serializer.is_valid():
            claim = serializer.save()
            
            # Log the activity
            log_activity(
                user=user,
                action='create',
                entity_type='AttendanceClaim',
                entity_id=claim.id,
                description=f"Created attendance claim for {attendance.date}",
                request=request
            )
            
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class AttendanceClaimActionView(APIView):
    """
    API View for approving or rejecting attendance claims.
    PATCH /api/attendance-claims/<claim_id>/<action>/
    """
    permission_classes = [IsAuthenticated]
    
    def patch(self, request, claim_id, action):
        user = request.user
        
        # Validate action
        if action not in ['approve', 'reject']:
            return Response(
                {"detail": "Invalid action. Must be 'approve' or 'reject'"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get the claim
        try:
            claim = AttendanceClaim.objects.get(id=claim_id)
        except AttendanceClaim.DoesNotExist:
            return Response(
                {"detail": "Claim not found"}, 
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Check permissions
        if not (user.is_staff or user.is_superuser):
            return Response(
                {"detail": "You do not have permission to perform this action"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Check if claim is already processed
        if claim.status != 'pending':
            return Response(
                {"detail": f"This claim has already been {claim.status}"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Process the action
        if action == 'approve':
            claim.status = 'approved'
            
            # Get or create the attendance record
            from django.utils import timezone
            from datetime import datetime, time as datetime_time
            
            # Get the user's Temp record
            try:
                temp = Temp.objects.get(user=claim.user)
            except Temp.DoesNotExist:
                return Response(
                    {"detail": "User profile not found"}, 
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Create or update the attendance record
            attendance, created = Attendance.objects.get_or_create(
                user=claim.user,
                date=claim.date,
                defaults={
                    'check_in': claim.check_in or datetime_time(9, 0),  # Default check-in at 9 AM
                    'check_out': claim.check_out or datetime_time(18, 0),  # Default check-out at 6 PM
                    'status': 'present',
                    'created_by': claim.user,
                    'updated_by': claim.user,
                    'emp_id': temp
                }
            )
            
            # Update the attendance record with claimed values
            if claim.check_in:
                attendance.check_in = claim.check_in
            if claim.check_out:
                attendance.check_out = claim.check_out
            attendance.status = 'present'
            attendance.updated_by = user
            attendance.updated_at = timezone.now()
            attendance.save()
            
            # Create attendance log entries
            if claim.check_in:
                AttendanceLog.objects.create(
                    attendance=attendance,
                    emp_id=temp,
                    time=datetime.combine(claim.date, claim.check_in),
                    reason='CLAIM_APPROVED',
                    is_in=True,
                    ip_address=get_client_ip(request)
                )
                
            if claim.check_out:
                AttendanceLog.objects.create(
                    attendance=attendance,
                    emp_id=temp,
                    time=datetime.combine(claim.date, claim.check_out),
                    reason='CLAIM_APPROVED',
                    is_in=False,
                    ip_address=get_client_ip(request)
                )
            
            message = "Claim approved and attendance record updated successfully"
        else:  # reject
            claim.status = 'rejected'
            # Require a reason for rejection
            reason = request.data.get('reason')
            if not reason:
                return Response(
                    {"detail": "Reason is required when rejecting a claim"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            claim.review_notes = reason
            message = "Claim rejected"
        
        claim.reviewed_by = user
        claim.reviewed_at = timezone.now()
        claim.save()
        
        # Log the activity
        log_activity(
            user=user,
            action=action,
            entity_type='AttendanceClaim',
            entity_id=claim.id,
            description=f"{action.capitalize()}d attendance claim for {claim.attendance.date}",
            request=request
        )
        
        return Response({
            "message": message,
            "claim": AttendanceClaimSerializer(claim).data
        })


def generate_completed_certificate(request, emp_id):
    try:
        temp = Temp.objects.get(emp_id=emp_id)
        user = temp.user
        personalData = PersonalData.objects.get(emp_id=emp_id)
        pronouns_he_she = 'he' if personalData.gender == 'M' else 'she'
        pronouns_his_her = 'His' if personalData.gender == 'M' else 'Her'
        response = HttpResponse(content_type='application/pdf')

        # 🔔 Filename → Example: Saranya_P_Internship_Certificate.pdf
        filename = f"{user.first_name}_{user.last_name}_Internship_Certificate.pdf"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # ===== Title =====
        p.setFont("Helvetica-Bold", 24)
        p.drawCentredString(width / 2, height - 100, "CERTIFICATE")

        # ===== Body Text =====
        p.setFont("Helvetica", 13)
        text = p.beginText(80, height - 180)
        text.setLeading(20)
        text.textLine(f"This is to certify that {user.first_name} {user.last_name} from")
        # Get college name from CollegeDetails model
        try:
            college_details = CollegeDetails.objects.get(emp_id=temp)
            college_name = college_details.college_name
        except CollegeDetails.DoesNotExist:
            college_name = 'Dr. Umayal Ramanathan College for Women'
        text.textLine(f"{college_name}, Karaikudi")
        # Get domain from UserData model
        try:
            user_data = UserData.objects.get(emp_id=temp)
            domain = user_data.domain.domain if user_data.domain else 'Full Stack'
        except UserData.DoesNotExist:
            domain = 'Full Stack'
        text.textLine(f"has completed {pronouns_his_her} Project in {domain} at VDart Academy")
        # Get dates from UserData model
        try:
            user_data = UserData.objects.get(emp_id=temp)
            if user_data.start_date and user_data.end_date:
                text.textLine(f"from {user_data.start_date.strftime('%B %d, %Y')} to {user_data.end_date.strftime('%B %d, %Y')}.")
            else:
                text.textLine("during the internship period.")
        except UserData.DoesNotExist:
            text.textLine("during the internship period.")
        text.textLine("")
        text.textLine("During her project she demonstrated excellent interpersonal skills")
        text.textLine("with a self-motivated attitude to learn new things. Her performance")
        text.textLine("exceeded expectations, and she was able to complete the project successfully on time.")
        text.textLine("")
        text.textLine("We wish her all the best for her future endeavors.")
        p.drawText(text)

        # ===== Date =====
        p.setFont("Helvetica", 12)
        issue_date = datetime.today().strftime("%B %d, %Y")
        p.drawString(80, 100, f"Date: {issue_date}")

        # ===== Signature =====
        p.setFont("Helvetica-Bold", 12)
        p.drawRightString(width - 80, 100, "Naresh Kumar K")
        p.setFont("Helvetica", 11)
        p.drawRightString(width - 80, 85, "Regional Head - Corporate HR")
        p.drawRightString(width - 80, 70, "VDart - India")

        p.showPage()
        p.save()
        buffer.seek(0)
        email = EmailMessage(
            subject="Your Internship Certificate from VDart Group",
            body="Please find attached your internship certificate.",
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[user.email],
        )
        email.attach(f"VDart_certificate_{emp_id}.pdf", buffer.getvalue(), "application/pdf")
        email.send(fail_silently=False)
        return response

    except Temp.DoesNotExist:
        return HttpResponse("Intern not found", status=404)


class GenerateTaskCertificate(APIView):
    def post(self, request, format=None):
        serializer = TaskCertificateSerializer(data=request.data)
        if serializer.is_valid():
            try:
                intern = Temp.objects.get(emp_id=serializer.validated_data['emp_id'])
                tasks = Task.objects.filter(
                    assigned_to=intern.user,
                    status='completed',
                    committed_date__isnull=False
                ).order_by('committed_date')
                
                template_path = Path(settings.BASE_DIR) / 'templates' / 'certificates' / 'task_cert_temp.docx'
                doc = DocxDocument(template_path)
                
                # Replace placeholders in the document
                for paragraph in doc.paragraphs:
                    paragraph.text = paragraph.text.replace('{{INTERN_NAME}}', intern.user.get_full_name())
                    paragraph.text = paragraph.text.replace('{{CURRENT_DATE}}', datetime.now().strftime('%B %d, %Y'))
                    paragraph.text = paragraph.text.replace('{{TOTAL_TASKS}}', str(tasks.count()))
                    paragraph.text = paragraph.text.replace('{{PERFORMANCE_COMMENT}}', 
                                                          serializer.validated_data.get('performance_comment', 
                                                                                      'outstanding performance'))

                        # Add tasks table
                table = doc.add_table(rows=1, cols=3)
                
                # Add headers
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Task Title'
                hdr_cells[1].text = 'Description'
                hdr_cells[2].text = 'Completed On'
                
                # Add task rows
                for task in tasks:
                    row_cells = table.add_row().cells
                    row_cells[0].text = task.title
                    row_cells[1].text = task.description
                    row_cells[2].text = task.committed_date.strftime('%Y-%m-%d')
                
                # Save the modified document to a temporary file
                temp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                doc.save(temp_docx.name)
                temp_docx.close()

                # Create a temp PDF output file path
                temp_pdf_path = temp_docx.name.replace(".docx", ".pdf")

                # Convert DOCX to PDF (save directly to file path)
                pythoncom.CoInitialize() 
                convert(temp_docx.name, temp_pdf_path)

                # Now you can read the PDF into memory (optional)
                with open(temp_pdf_path, 'rb') as pdf_file:
                    pdf_bytes = pdf_file.read()

                pdf_buffer = BytesIO(pdf_bytes)
                pdf_buffer.seek(0)

                # Clean up temp files
                os.unlink(temp_docx.name)
                os.unlink(temp_pdf_path)
                
                # Send email with PDF attachment
                subject = f"Task Completion Certificate - {intern.user.get_full_name()}"
                message = f"Dear {intern.user.get_full_name()},\n\nPlease find attached your task completion certificate.\n\nBest regards,\nThe Management"
                to_email = [intern.user.email]
                
                send_email_with_attachment(
                    subject=subject,
                    message=message,
                    recipient_list=to_email,
                    attachment=pdf_buffer.getvalue(),
                    filename=f"Task_Certificate_{intern.emp_id}.pdf",
                    content_type='application/pdf'
                )
                
                return Response(
                    {"success": True, "message": "Task certificate sent successfully"},
                    status=status.HTTP_200_OK
                )
                
            except Temp.DoesNotExist:
                return Response(
                    {"success": False, "error": "Intern not found"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                return Response(
                    {"success": False, "error": str(e)},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class GenerateAttendanceCertificate(APIView):
    def post(self, request, format=None):
        serializer = AttendanceCertificateSerializer(data=request.data)
        if serializer.is_valid():
            try:
                intern = Temp.objects.get(emp_id=serializer.validated_data['emp_id'])
                attendance = Attendance.objects.filter(user=intern.user)
                
                
                # Load the Word template
                template_path = 'templates/certificates/attendance_certificate_template.docx'
                doc = Document(template_path)
                
                # Replace placeholders in the document
                for paragraph in doc.paragraphs:
                    paragraph.text = paragraph.text.replace('{{INTERN_NAME}}', intern.user.get_full_name())
                    paragraph.text = paragraph.text.replace('{{CURRENT_DATE}}', datetime.now().strftime('%B %d, %Y'))
                    paragraph.text = paragraph.text.replace('{{PERIOD}}', attendance.for_period)
                    paragraph.text = paragraph.text.replace('{{FROM_DATE}}', 
                                                          attendance.from_date.strftime('%B %d, %Y'))
                    paragraph.text = paragraph.text.replace('{{TO_DATE}}', 
                                                         attendance.to_date.strftime('%B %d, %Y'))
                    paragraph.text = paragraph.text.replace('{{ATTENDANCE_STATUS}}', 
                                                          'Full-time' if attendance.from_day_type == 'full' 
                                                          else 'Part-time')
                    # Add more placeholders as needed
                    paragraph.text = paragraph.text.replace('{{DOMAIN}}', intern.domain or 'N/A')
                    paragraph.text = paragraph.text.replace('{{DEPARTMENT}}', intern.department or 'N/A')
                    paragraph.text = paragraph.text.replace('{{TOTAL_DAYS}}', 
                                                         str((attendance.to_date - attendance.from_date).days + 1))
                
                # Save the modified document to a temporary file
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    doc.save(tmp.name)
                    tmp_path = tmp.name
                
                # Convert to PDF in memory
                pdf_buffer = BytesIO()
                convert(tmp_path, pdf_buffer)
                pdf_buffer.seek(0)
                
                # Clean up the temporary file
                os.unlink(tmp_path)
                
                # Send email with PDF attachment
                subject = f"Attendance Certificate - {intern.user.get_full_name()}"
                message = f"Dear {intern.user.get_full_name()},\n\nPlease find attached your attendance certificate for {attendance.for_period}.\n\nBest regards,\nThe Management"
                to_email = [intern.user.email]
                
                send_email_with_attachment(
                    subject=subject,
                    message=message,
                    recipient_list=to_email,
                    attachment=pdf_buffer,
                    filename=f"Attendance_Certificate_{intern.emp_id}.pdf",
                    content_type='application/pdf'
                )
                
                return Response(
                    {"success": True, "message": "Attendance certificate sent successfully"},
                    status=status.HTTP_200_OK
                )
                
            except (Temp.DoesNotExist, Attendance.DoesNotExist) as e:
                return Response(
                    {"success": False, "error": "Intern or attendance not found"},
                    status=status.HTTP_404_NOT_FOUND
                )
            except Exception as e:
                return Response(
                    {"success": False, "error": str(e)},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)