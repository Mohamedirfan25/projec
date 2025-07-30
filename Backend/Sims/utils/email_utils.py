import os
import traceback
import logging
from io import BytesIO
from datetime import datetime
from django.core.mail import EmailMessage, BadHeaderError
from django.conf import settings
from reportlab.lib.pagesizes import A4, letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from dateutil import parser
from docx import Document
import tempfile
import uuid
from datetime import date
from docx2pdf import convert

# Configure logging
logger = logging.getLogger(__name__)

def log_error(error_type, error_message, details=None):
    """Helper function to log errors with consistent formatting"""
    log_message = f"{error_type.upper()}: {error_message}"
    if details:
        log_message += f"\nDetails: {details}"
    logger.error(log_message, exc_info=True)
    return log_message

def generate_offer_letter_docx(user, emp_id, college_name, start_date, end_date,
                            position_title="FullStack Intern",
                            domain="VDart Academy",
                            shift_time="9:00 AM to 1:00 PM IST",
                            shift_days="Monday to Friday",
                            work_location="VDart, Global Capability Center, Mannarpuram",
                            reporting_to="Derrick Alex"):
    """
    Generate an offer letter using the Word template with all placeholders replaced.
    """
    error_details = {
        'step': 'Initialization',
        'error_type': None,
        'message': None
    }
    
    try:
        # Define the template path
        template_path = r"D:\Intern\projec\Backend\media\word docs\VDart_Offer_Letter_ACA030 (1).docx"
        
        # Check if template exists
        if not os.path.exists(template_path):
            error_msg = f"Template file not found at: {template_path}"
            logger.error(error_msg)
            error_details.update({
                'step': 'Template Validation',
                'error_type': 'FileNotFoundError',
                'message': error_msg,
                'template_path': template_path
            })
            return False, error_msg, None, error_details
            
        # Create a temporary file for the output
        output_path = os.path.join(tempfile.gettempdir(), f'offer_letter_{uuid.uuid4()}.docx')
        
        try:
            # Load the document
            doc = Document(template_path)
            
            # Helper function to format dates consistently
            def format_date(date_value):
                if not date_value:
                    return ""
                try:
                    # Handle string dates with timezone info (e.g., '2025-07-30T04:00:00')
                    if isinstance(date_value, str) and 'T' in date_value:
                        # Parse the date part only (before 'T')
                        date_part = date_value.split('T')[0]
                        date_obj = datetime.strptime(date_part, '%Y-%m-%d')
                        return date_obj.strftime('%d-%b-%Y')
                    # Handle datetime objects
                    elif hasattr(date_value, 'strftime'):
                        return date_value.strftime('%d-%b-%Y')
                    return str(date_value)
                except Exception as e:
                    logger.error(f"Error formatting date {date_value}: {str(e)}")
                    return str(date_value)
            
            # Format dates consistently
            formatted_start_date = format_date(start_date)
            formatted_end_date = format_date(end_date)
            
            # Debug gender information
            user_gender = getattr(user, 'gender', 'not_set')
            logger.info(f"User gender from database: '{user_gender}' (type: {type(user_gender)})")
            
            # Determine prefix based on gender
            gender_lower = str(user_gender).lower().strip() if user_gender else 'not_set'
            if gender_lower in ['male', 'm', '1', 'true', 'yes']:
                intern_prefix = 'Mr.'
            elif gender_lower in ['female', 'f', '2', 'false', 'no']:
                intern_prefix = 'Ms.'
            else:
                intern_prefix = 'Mr.'  # Default to Mr. if not specified
                logger.warning(f"Unexpected gender value: '{user_gender}'. Defaulting to 'Mr.'")
            
            logger.info(f"Using prefix: '{intern_prefix}' for user {getattr(user, 'id', 'unknown')}")
            
            # Prepare all possible replacements with proper formatting
            replacements = {
                # Basic information with gender prefix
                '{{intern_prefix}}': intern_prefix,
                '{{prefix}}': intern_prefix,  # Common alternative
                '{{title}}': intern_prefix,   # Another common alternative
                
                # Date fields with proper formatting
                '{{start_date}}': formatted_start_date,
                '{{end_date}}': formatted_end_date,
                '{{intern_start_date}}': formatted_start_date,
                '{{intern_end_date}}': formatted_end_date,
                '{{start_date_dd_mmm_yyyy}}': formatted_start_date,
                '{{end_date_dd_mmm_yyyy}}': formatted_end_date,
                '{{internship_start_date}}': formatted_start_date,
                '{{internship_end_date}}': formatted_start_date,
                
                # Other fields
                '{{intern_name}}': user.get_full_name(),
                '{{intern_id}}': str(emp_id),
                '{{college_name}}': college_name,
                '{{position_title}}': position_title,
                '{{domain}}': domain,
                '{{shift_time}}': shift_time,
                '{{shift_days}}': shift_days,
                '{{work_location}}': work_location,
                '{{reporting_to}}': reporting_to,
                '{{current_date}}': datetime.now().strftime('%d-%b-%Y'),
                
                # Additional variations of placeholders
                '{{intern_college}}': college_name,
                '{{intern_domain}}': domain,
                '{{intern_position}}': position_title,
                '{{intern_location}}': work_location,
                '{{intern_reporting_to}}': reporting_to,
                '{{intern_shift}}': f"{shift_time} ({shift_days})",
                '{{intern_period}}': f"from {formatted_start_date} to {formatted_end_date}",
                '{{internship_period}}': f"from {formatted_start_date} to {formatted_end_date}",
                '{{intern_shift_time}}': shift_time,
                '{{intern_shift_days}}': shift_days,
                '{{internship_shift_time}}': shift_time,
                '{{internship_shift_days}}': shift_days,
                '{{internship_shift}}': f"{shift_time} ({shift_days})",
                '{{internship_location}}': work_location,
                '{{internship_domain}}': domain,
                '{{internship_reporting_to}}': reporting_to,
            }
            
            # Log the replacements for debugging
            logger.info(f"Prepared replacements: {replacements}")
            
            # Function to recursively process all document elements
            def process_element(element):
                if hasattr(element, 'text'):
                    # Process paragraphs
                    for old_text, new_text in replacements.items():
                        if old_text in element.text:
                            # Store the original runs
                            original_runs = element.runs
                            element.text = element.text.replace(old_text, str(new_text))
                            
                            # Try to preserve formatting if there were runs
                            if original_runs and len(original_runs) > 0:
                                for run in element.runs:
                                    # Copy formatting from the first original run
                                    run.font.name = original_runs[0].font.name
                                    run.font.size = original_runs[0].font.size
                                    run.font.bold = original_runs[0].font.bold
                                    run.font.italic = original_runs[0].font.italic
                                    run.font.underline = original_runs[0].font.underline
                                    if original_runs[0].font.color.rgb:
                                        run.font.color.rgb = original_runs[0].font.color.rgb
                
                # Process tables
                if hasattr(element, 'tables'):
                    for table in element.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                process_element(cell)
                
                # Process paragraphs in the current element
                if hasattr(element, 'paragraphs'):
                    for paragraph in element.paragraphs:
                        process_element(paragraph)
            
            # Process the main document
            process_element(doc)
            
            # Save the document
            doc.save(output_path)
            logger.info(f"Successfully generated offer letter at: {output_path}")
            
            # Verify the content was replaced
            with open(output_path, 'rb') as f:
                content = f.read().decode('utf-8', errors='ignore')
                for placeholder in replacements.keys():
                    if placeholder in content:
                        logger.warning(f"Placeholder {placeholder} was not replaced in the document")
            
            return True, "Offer letter generated successfully", output_path, None
            
        except Exception as e:
            error_msg = f"Error generating offer letter: {str(e)}"
            logger.error(error_msg, exc_info=True)
            error_details.update({
                'step': 'Document Generation',
                'error_type': type(e).__name__,
                'message': str(e),
                'traceback': traceback.format_exc()
            })
            
            # Clean up the temp file if it was created
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except Exception as cleanup_error:
                    logger.warning(f"Error cleaning up temp file {output_path}: {str(cleanup_error)}")
            
            return False, error_msg, None, error_details
            
    except Exception as e:
        error_msg = f"Unexpected error in generate_offer_letter_docx: {str(e)}"
        logger.error(error_msg, exc_info=True)
        error_details.update({
            'step': 'Unexpected Error',
            'error_type': 'UnexpectedError',
            'message': str(e),
            'traceback': traceback.format_exc()
        })
        return False, error_msg, None, error_details

def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert a DOCX file to PDF using docx2pdf with enhanced error handling"""
    try:
        # Ensure the source file exists and is not empty
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Source DOCX file not found: {docx_path}")
            
        if os.path.getsize(docx_path) == 0:
            raise ValueError(f"Source DOCX file is empty: {docx_path}")
            
        logger.info(f"Starting PDF conversion: {docx_path} -> {pdf_path}")
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
        
        # Try direct conversion first
        try:
            convert(docx_path, pdf_path)
        except Exception as e:
            logger.warning(f"Direct conversion failed, trying with subprocess: {str(e)}")
            # Fallback to subprocess method
            import subprocess
            import sys
            
            # Use the full path to the docx2pdf executable if available
            if sys.platform == 'win32':
                # On Windows, try to use the installed docx2pdf executable
                try:
                    subprocess.run(['docx2pdf', docx_path, pdf_path], check=True)
                except (subprocess.CalledProcessError, FileNotFoundError):
                    # If that fails, try using python -m docx2pdf
                    python_exec = sys.executable or 'python'
                    subprocess.run([python_exec, '-m', 'docx2pdf', docx_path, os.path.dirname(pdf_path)], check=True)
            else:
                # On Unix-like systems
                subprocess.run(['docx2pdf', docx_path, os.path.dirname(pdf_path)], check=True)
        
        # Verify the PDF was created and is not empty
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file was not created at: {pdf_path}")
            
        if os.path.getsize(pdf_path) == 0:
            raise ValueError(f"Generated PDF file is empty: {pdf_path}")
            
        logger.info(f"Successfully converted DOCX to PDF: {pdf_path} (Size: {os.path.getsize(pdf_path)} bytes)")
        return True, "PDF generated successfully"
        
    except Exception as e:
        error_msg = f"Error converting DOCX to PDF: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return False, error_msg

def send_offer_letter(user, emp_id, college_name, start_date, end_date,
                    position_title="FullStack Intern",
                    domain="VDart Academy",
                    shift_time="9:00 AM to 1:00 PM IST",
                    shift_days="Monday to Friday",
                    work_location="VDart, Global Capability Center, Mannarpuram",
                    reporting_to="Derrick Alex"):
    """
    Main function to generate and send offer letter as PDF.
    """
    error_details = {
        'step': 'Initialization',
        'error_type': None,
        'message': None
    }
    
    # Create a temporary directory for the files
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            logger.info(f"Starting offer letter generation for {user.email} (ID: {emp_id})")
            
            # Step 1: Generate the DOCX from template
            success, message, docx_path, error_details = generate_offer_letter_docx(
                user=user,
                emp_id=emp_id,
                college_name=college_name,
                start_date=start_date,
                end_date=end_date,
                position_title=position_title,
                domain=domain,
                shift_time=shift_time,
                shift_days=shift_days,
                work_location=work_location,
                reporting_to=reporting_to
            )
            
            if not success:
                return False, message, error_details
            
            # Step 2: Convert DOCX to PDF
            pdf_filename = f"Offer_Letter_{emp_id}.pdf"
            pdf_path = os.path.join(temp_dir, pdf_filename)
            
            success, message = convert_docx_to_pdf(docx_path, pdf_path)
            
            # Clean up the temporary DOCX file
            try:
                if os.path.exists(docx_path):
                    os.remove(docx_path)
            except Exception as e:
                logger.warning(f"Failed to remove temporary DOCX file: {str(e)}")
            
            if not success:
                return False, message, {
                    'step': 'PDF Conversion',
                    'error_type': 'ConversionError',
                    'message': message
                }
            
            # Step 3: Send the email with PDF attachment
            email_subject = f"Offer Letter - {user.get_full_name() or 'New Hire'}"
            email_body = f"""
            Dear {user.get_full_name() or 'Candidate'},
            
            Please find attached your offer letter for the {position_title} position.
            
            Best regards,
            The Hiring Team
            """
            
            try:
                send_email_with_attachment(
                    subject=email_subject.strip(),
                    message=email_body.strip(),
                    recipient_list=[user.email],
                    attachment=pdf_path,
                    filename=pdf_filename,
                    content_type='application/pdf'
                )
                
                logger.info(f"Successfully sent offer letter to {user.email}")
                return True, "Offer letter sent successfully", None
                
            except Exception as email_error:
                error_msg = f"Failed to send email: {str(email_error)}"
                logger.error(error_msg, exc_info=True)
                return False, error_msg, {
                    'step': 'Email Sending',
                    'error_type': 'EmailError',
                    'message': str(email_error)
                }
            
        except Exception as e:
            error_msg = f"Unexpected error in send_offer_letter: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, error_msg, {
                'step': 'Unexpected Error',
                'error_type': 'UnexpectedError',
                'message': str(e),
                'traceback': traceback.format_exc()
            }

def send_email_with_attachment(
    subject,
    message,
    recipient_list,
    attachment=None,
    filename=None,
    content_type='application/octet-stream',
    from_email=None,
    fail_silently=False,
    **kwargs
):
    """
    Send an email with an attachment.
    
    Args:
        subject (str): Email subject
        message (str): Email body text
        recipient_list (list): List of recipient email addresses
        attachment: File-like object containing the attachment data
        filename (str): Name to give to the attachment
        content_type (str): MIME type of the attachment
        from_email (str): Sender email address (defaults to DEFAULT_FROM_EMAIL)
        fail_silently (bool): If True, exceptions will be caught and logged
        **kwargs: Additional arguments to pass to EmailMessage
    
    Returns:
        bool: True if the email was sent successfully, False otherwise
    """
    try:
        from_email = from_email or getattr(settings, 'DEFAULT_FROM_EMAIL')
        
        # Create the email message
        email = EmailMessage(
            subject=subject,
            body=message,
            from_email=from_email,
            to=recipient_list,
            **kwargs
        )
        
        # Add attachment if provided
        if attachment and filename:
            # For in-memory files
            if hasattr(attachment, 'read'):
                attachment.seek(0)  # Ensure we're at the start of the file
                email.attach(filename, attachment.read(), content_type)

            # For file paths
            elif isinstance(attachment, str):
                with open(attachment, 'rb') as file:
                    email.attach(filename, file.read(), content_type)

        
        # Send the email
        email.send(fail_silently=fail_silently)
        logger.info(f"Email sent successfully to {', '.join(recipient_list)}")
        return True
        
    except Exception as e:
        logger.error(f"Failed to send email: {str(e)}", exc_info=True)
        if not fail_silently:
            raise
        return False
